#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
WSU Manual Converter — DOCX → HTML (Preview & Downloads)

What it does
------------
- Preprocess DOCX (strip manual TOC pages, promote headings)
- Convert with Pandoc
- Post-process HTML:
    * Remove Pandoc's title block & stray <style>
    * Normalize whitespace
    * Convert typed lists (1./a./i.) to real <ol>/<li>
        - If a new list follows an item ending ':' → NEST inside that <li>
        - If outer list resumes afterward → MERGE back, preserve numbering
    * Strip typed heading numbers (DOM-aware):
        - Removes "Chapter 2 –", "Section I:", "I.A.2.", "1.2.3", etc.
        - Works even when the prefix is wrapped in <span>/<strong>/<em>, or split
- Outputs EXACT wrapper your site CSS/JS expects at the very top:
    <div class="manual-grid">
      <nav class="manual-toc">…</nav>
      <div class="manual"> …content… </div>
    </div>
- Keeps a simple preview and download buttons at http://localhost:5000/

"""

import os
import math
import re
import uuid
import tempfile
import subprocess
import json
import hashlib
import zipfile
import shutil
import io
import xml.etree.ElementTree as ET
import logging
from datetime import datetime
from pathlib import Path

from flask import Flask, request, render_template, render_template_string, send_file, redirect, url_for, flash
import markdown
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge

from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, Tag, NavigableString

# Load local configuration and session helper
from config import LOG_LEVEL, FLASK_SECRET_KEY, PERSIST_DIR, REFERENCE_DIR, SessionDir
from core.permalinks import (
    normalize_heading_signature,
    normalize_heading_ref,
    ensure_prefixed
)
from core.html_processor import (
    add_heading_ids,
    save_stable_heading_map,
    extract_body,
    strip_pandoc_styles,
    strip_images_and_figures,
    process_html_pipeline,
    apply_css_counter_numbering,
    strip_inline_formatting,
    sanitize_docx_ids_for_export,
    has_tables_in_html,
    normalize_spaces,
    normalize_typed_lists,
    strip_toc_sections_dom,
    infer_heading_levels_from_prefix,
    strip_heading_numbers_dom,
    apply_heading_edits,
    apply_list_classes_and_styles,
    format_manual_tables,
    apply_reference_edits,
    strip_html_assets,
    shift_heading_levels,
    extract_manual_fragment,
    parse_heading_id_map_json,
    build_manual_grid_block,
)
from core.pandoc_wrapper import run_pandoc
from core.docx_processor import (
    preprocess_docx,
    compute_sha256,
    extract_docx_hyperlinks,
    is_heading_style,
    get_outline_level,
    get_style_outline_level,
    get_heading_level,
    guess_heading_level,
    extract_sequence_map_from_doc,
    merge_sequence_maps,
    serialize_sequence_map,
    deserialize_sequence_map,
    inject_heading_list_numbering,
    apply_heading_levels_from_numbering,
    promote_headings,
    strip_manual_toc_paragraphs,
    build_numbering_crosswalk,
    generate_stable_ref_id,
    extract_heading_structure_and_references,
    has_tables_in_docx,
    sanitize_docx_bookmark_id,
    sanitize_docx_styles,
    fix_numbering_xml,
    relocate_body_level_bookmarks,
    extract_style_map_from_reference,
    build_clean_reference_doc,
    extract_reference_doc_summary,
)
from core.manual_structure import (
    scrape_heading_structure_from_html,
    auto_match_old_to_new_references,
    parse_heading_key,
    heading_sort_key,
    heading_dropdown_sort,
    build_display_text_from_heading,
    find_heading_by_full,
    lookup_heading_title,
    build_heading_crosswalk_from_map,
    convert_old_numbering_to_new,
    find_heading_order_violations,
    extract_heading_editor_rows,
    extract_heading_edits_from_form
)
from core.reference_linking import (
    extract_references_from_html,
    extract_external_links_from_html,
    extract_external_links_from_reference_text
)
from core.styling import (
    default_theme_settings,
    coerce_theme_settings,
    build_theme_css,
    contrast_ratio,
    resolve_reference_doc_path,
    get_reference_style_context,
    parse_reference_overrides_from_form,
    get_wp_css_text,
    get_wp_js_text,
)
from utils.helpers import (
    roman_to_int,
    normalize_hex_color,
    clamp_number,
    sanitize_theme_id,
    _int_to_roman,
    _int_to_letters,
    _format_number,
    _token_type_from_numfmt
)

# Configure logging
logging.basicConfig(
    level=LOG_LEVEL,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ------------------------------ Flask setup ------------------------------

app = Flask(__name__)
app.secret_key = FLASK_SECRET_KEY
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024
app.config["MAX_FORM_MEMORY_SIZE"] = 200 * 1024 * 1024
app.config["MAX_FORM_PARTS"] = 20000

@app.errorhandler(RequestEntityTooLarge)
def handle_request_too_large(error):
    flash("Request too large. Use a smaller file or increase the upload/form limits.")
    return redirect(url_for("index"))

# Print persist directory location on startup for debugging
import atexit
@atexit.register
def print_persist_dir():
    pass  # Will be printed in __main__

# Utility: compute SHA256 for a file

# --------------------------- DOCX preprocessing ---------------------------

# ----------------------------- Pandoc / HTML ------------------------------

# ---- HTML import helpers -------------------------------------------------

# ---- DOM-aware heading number scrubber ----------------------------------

# comprehensive prefix matcher (on normalized text)
# IMPORTANT: Must require whitespace after prefix to avoid consuming first letter of heading text
# The pattern ensures we don't match "I.A.3.Duties" as "I.A.3.D" - we need whitespace before actual text
_HEADING_PREFIX_RE = re.compile(
    r"^\s*(?:"
    # Section/Chapter with multi-level numbering: Section I.1 | Section I.A.2 | Chapter 1.2.3
    r"(?i:(?:Chapter|Section))\s+[IVXLCDM\d]+(?:\.[A-Z\d]+)*(?:\s*[:.\---])?\s+|"
    r"(?:\d+|[IVXLCDM]{1,6}|[A-Z]{1,3}|[a-z]{1,3})(?:[.\s]+(?:\d+|[IVXLCDM]{1,6}|[A-Z]{1,3}|[a-z]{1,3})){1,5}\.?\s+(?:[:.\---]\s*)?|"
    r"(?i:[IVXLCDM]+)\.(?:[A-Z]{1,3}|[a-z]{1,3})(?:\.\d+){0,3}\.?\s+(?:[:.\---]\s*)?|"  # I.A. | I.AA.2. | I.A.2.1. : (requires space after)
    r"(?i:[IVXLCDM]+)(?:\.\d+){0,3}\.?\s+(?:[:.\---]\s*)?|"         # II.3. | IV.2.1. (requires space after)
    r"(?:[A-Z]{1,3}|[a-z]{1,3})\.\s+(?:[:.\---]\s*)?|"                          # A. or AA. (requires space after)
    r"\d+(?:\.\d+){0,3}(?:[.)])?\s+(?:[:.\---]\s*)?"          # 1. | 1.2 | 1.2.3 (requires space after)
    r")"
)

_WHITESPACE_EQUIV = {
    '\u00a0': ' ',  # NBSP -> space
    '\u2009': ' ', '\u2002': ' ', '\u2003': ' ', '\u200a': ' ',
    '\u202f': ' ', '\u205f': ' ', '\u3000': ' '
}
_ZERO_WIDTH = {'\u200b', '\u200c', '\u200d'}

_HEADING_INFER_TOKEN = r'(?:\d+|[IVXLCDMivxlcdm]{1,6}|[A-Z]{1,3}|[a-z]{1,3})'
_HEADING_INFER_RE = re.compile(
    r'^\s*(?:(?i:Chapter|Section)\s+)?'
    r'(' + _HEADING_INFER_TOKEN + r'(?:\.\s*' + _HEADING_INFER_TOKEN + r'){1,5})\b'
)

# REMOVED: Automatic reference updating - now handled manually via review interface
# The update_internal_references function has been removed.
# References are now shown in the review interface for manual editing.

# ---------- List normalization (typed + existing <ol>) --------------------

# Helper function to convert old numbering format to new format and create anchor link
# Old format examples: "1.D.2.e", "1.4.2.5", "Section 1.D.2.e"

# ---------- List normalization (typed + existing <ol>) --------------------

_alpha = re.compile(r'^\s*([a-zA-Z])[.)]\s+')
_decimal = re.compile(r'^\s*(\d+)[.)]\s+')
_roman  = re.compile(r'^\s*((?i:i|ii|iii|iv|v|vi|vii|viii|ix|x|xi|xii|xiii|xiv|xv|xvi|xvii|xviii|xix|xx))[.)]\s+')

SPELLED_NUMS = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6,
    "seven": 7, "eight": 8, "nine": 9, "ten": 10, "eleven": 11, "twelve": 12,
    "thirteen": 13, "fourteen": 14, "fifteen": 15, "sixteen": 16, "seventeen": 17,
    "eighteen": 18, "nineteen": 19, "twenty": 20
}

@app.route("/heading_review/<session_id>", methods=["GET", "POST"])
def heading_review(session_id):
    """Review and edit heading crosswalk before detailed reference editing."""
    import json
    session = SessionDir(session_id)
    session_file = session.session_json
    if not session_file.exists():
        flash("Session expired or invalid.")
        return redirect(url_for("index"))

    session_data = json.loads(session_file.read_text(encoding='utf-8'))
    auto_crosswalk = session_data.get('approved_crosswalk') or session_data.get('auto_crosswalk', {})
    approved_crosswalk = session_data.get('approved_crosswalk', {})
    old_crosswalk = session_data.get('old_crosswalk', {})
    heading_map = session_data.get('heading_map', {})
    heading_order = session_data.get('heading_order', {})
    new_headings = session_data.get('new_headings', {})
    normalized_new_headings = False
    for heading_key, heading_data in new_headings.items():
        if heading_key.startswith(('H1:', 'H2:', 'H3:', 'H4:', 'H5:', 'H6:')):
            continue
        text = heading_data.get('text', '').strip()
        full = heading_data.get('full', '').strip()
        if text and full and not full.lower().startswith(heading_key.lower()):
            heading_data['full'] = f"{heading_key} - {text}"
            normalized_new_headings = True
        elif text and not full:
            heading_data['full'] = f"{heading_key} - {text}"
            normalized_new_headings = True
    if normalized_new_headings:
        session_data['new_headings'] = new_headings
        session_file.write_text(json.dumps(session_data, indent=2), encoding='utf-8')
    filename = session_data.get('filename', '')
    manual_type = session_data.get('manual_type', 'chapter')

    logger.debug(f"DEBUG [heading_review]: Session loaded with heading_map={len(heading_map)} entries")
    if heading_map:
        sample_keys = list(heading_map.keys())[:5]
        logger.debug(f"DEBUG [heading_review]: Sample heading_map keys: {sample_keys}")
    logger.debug(f"DEBUG [heading_review]: approved_crosswalk={len(approved_crosswalk)} entries")
    logger.debug(f"DEBUG [heading_review]: auto_crosswalk={len(auto_crosswalk)} entries")
    logger.debug(f"DEBUG [heading_review]: old_crosswalk={len(old_crosswalk)} entries")

    # If user chose to keep old headings, skip this step
    if session_data.get('mapping_mode', 'map_new') == 'keep_old':
        return redirect(url_for('review', session_id=session_id))

    # Auto-load heading_map.xlsx/csv from project root ONLY if no heading_map was extracted from DOCX
    # This prevents external files from overriding the extracted heading structure
    if not approved_crosswalk and not heading_map:
        for candidate in ["heading_map.xlsx", "heading_map.csv"]:
            file_path = Path(candidate)
            if file_path.exists():
                try:
                    with open(file_path, "rb") as fh:
                        class DummyFile:
                            def __init__(self, name, data):
                                self.filename = name
                                self._data = data
                                self.stream = io.BytesIO(data)
                            def read(self):
                                return self._data
                        dummy = DummyFile(file_path.name, fh.read())
                        loaded_map, _ = load_heading_map_file(dummy)
                        if loaded_map:
                            normalized_map = {}
                            tmp_order = {}
                            for idx, (old_ref, new_ref) in enumerate(loaded_map.items()):
                                norm_old = ensure_prefixed(normalize_heading_ref(old_ref), manual_type)
                                norm_new = ensure_prefixed(normalize_heading_ref(new_ref), manual_type)
                                if norm_old and norm_new:
                                    normalized_map[norm_old] = norm_new
                                    tmp_order[norm_old] = idx
                            approved_crosswalk = normalized_map
                            if tmp_order:
                                heading_order = tmp_order
                            session_data['approved_crosswalk'] = approved_crosswalk
                            if heading_order:
                                session_data['heading_order'] = heading_order
                            session_file.write_text(json.dumps(session_data, indent=2), encoding='utf-8')
                            flash(f"Loaded heading map from {file_path.name}.")
                            logger.debug(f"DEBUG [heading_review]: Auto-loaded {len(normalized_map)} entries from {file_path.name}")
                except Exception as e:
                    logger.error(f"DEBUG [heading_review]: Failed to auto-load heading map {file_path}: {e}")

    if request.method == 'POST':
        # Optional heading map upload
        uploaded_map = request.files.get('heading_map')
        updated = {}
        if uploaded_map and uploaded_map.filename:
            uploaded_mapping, _ = load_heading_map_file(uploaded_map)
            for idx, (old_ref, new_ref) in enumerate(uploaded_mapping.items()):
                norm_old = ensure_prefixed(normalize_heading_ref(old_ref), manual_type)
                norm_new = ensure_prefixed(normalize_heading_ref(new_ref), manual_type)
                if norm_old and norm_new:
                    updated[norm_old] = norm_new
                    heading_order[norm_old] = idx
        else:
            # Collect edits from form inputs with validation checkboxes
            # First, identify all heading IDs
            heading_ids = set()
            for key in request.form.keys():
                if key.startswith("valid_"):
                    heading_ids.add(key.replace("valid_", ""))

            # Now process each heading
            updated_titles = {}  # Track edited NEW titles {new_ref: title}

            for heading_id in heading_ids:
                # Check if this heading is marked as valid (checkbox checked)
                is_valid = f'valid_{heading_id}' in request.form

                if is_valid:
                    # Get the original old_ref, edited new_ref, and edited new_title
                    old_ref_raw = request.form.get(f'old_ref_{heading_id}', '').strip()
                    new_ref_raw = request.form.get(f'new_ref_{heading_id}', '').strip()
                    new_title_raw = request.form.get(f'new_title_{heading_id}', '').strip()

                    if old_ref_raw and new_ref_raw:
                        old_ref = ensure_prefixed(normalize_heading_ref(old_ref_raw), manual_type)
                        new_ref = ensure_prefixed(normalize_heading_ref(new_ref_raw), manual_type)
                        if old_ref and new_ref:
                            updated[old_ref] = new_ref

                            # If user edited the title, save it to update new_headings
                            if new_title_raw:
                                updated_titles[new_ref] = new_title_raw

                            logger.debug(f"DEBUG [heading_review POST]: Including valid entry: '{old_ref}' -> '{new_ref}'"
                                  f"{(' with edited title' if new_title_raw else '')}")        # If nothing provided, keep auto_crosswalk
        if not updated:
            fallback = {}
            if heading_map:
                fallback, heading_order = build_heading_crosswalk_from_map(heading_map, manual_type)
            elif old_crosswalk:
                for idx, (old_ref, new_ref) in enumerate(old_crosswalk.items()):
                    norm_old = ensure_prefixed(normalize_heading_ref(old_ref), manual_type)
                    norm_new = ensure_prefixed(normalize_heading_ref(new_ref), manual_type)
                    if norm_old and norm_new:
                        fallback[norm_old] = norm_new
                        heading_order[norm_old] = heading_order.get(norm_old, idx)
            else:
                fallback = auto_crosswalk
            updated = fallback

        session_data['approved_crosswalk'] = updated
        session_data['heading_order'] = heading_order

        # Update new_headings with edited titles
        if updated_titles:
            new_headings = session_data.get('new_headings', {})
            for new_ref, edited_title in updated_titles.items():
                # Update or create heading entry with edited title
                if new_ref in new_headings:
                    new_headings[new_ref]['text'] = edited_title
                    new_headings[new_ref]['full'] = f"{new_ref} - {edited_title}"
                else:
                    # Create new entry if doesn't exist
                    new_headings[new_ref] = {
                        'text': edited_title,
                        'full': f"{new_ref} - {edited_title}",
                        'level': 'unknown',
                        'id': ''
                    }
            session_data['new_headings'] = new_headings
            logger.debug(f"DEBUG [heading_review POST]: Updated {len(updated_titles)} NEW heading titles")

        session_file.write_text(json.dumps(session_data, indent=2), encoding='utf-8')

        # Standardized edit file path
        session = SessionDir(session_id)
        edit_file = session.edits_json
        edit_data = {}
        if edit_file.exists():
            try:
                edit_data = json.loads(edit_file.read_text(encoding='utf-8'))
            except Exception:
                edit_data = {}
        if not edit_data:
            edit_data = {
                'document': filename,
                'auto_crosswalk': auto_crosswalk,
                'approved_crosswalk': {},
                'reference_edits': {},
                'reference_validations': {},
                'reference_link_targets': {},
                'reference_ignored': {},
                'last_updated': str(datetime.now())
            }
        edit_data['document'] = filename
        edit_data['auto_crosswalk'] = auto_crosswalk
        edit_data['approved_crosswalk'] = updated
        edit_data['last_updated'] = str(datetime.now())
        edit_file.write_text(json.dumps(edit_data, indent=2), encoding='utf-8')

        flash(f"Heading crosswalk saved. {len(updated)} entries" + (f", {len(updated_titles)} titles edited." if updated_titles else "."))
        return redirect(url_for('review', session_id=session_id))

    # Build display list - PRIORITIZE approved_crosswalk so saved edits stay visible
    display_crosswalk = {}

    # Priority 1: Use approved_crosswalk if manually uploaded or saved
    if approved_crosswalk:
        logger.debug(f"DEBUG [heading_review]: Using approved_crosswalk ({len(approved_crosswalk)} entries)")
        display_crosswalk = approved_crosswalk.copy()
    # Priority 2: Use heading_map if available (extracted from DOCX)
    elif heading_map:
        logger.debug(f"DEBUG [heading_review]: Building display_crosswalk from heading_map ({len(heading_map)} entries)")
        display_crosswalk, heading_order = build_heading_crosswalk_from_map(heading_map, manual_type)
        logger.debug(f"DEBUG [heading_review]: Built display_crosswalk with {len(display_crosswalk)} entries from heading_map")
    # Priority 3: Fall back to old_crosswalk
    elif old_crosswalk:
        logger.debug(f"DEBUG [heading_review]: Falling back to old_crosswalk ({len(old_crosswalk)} entries)")
        tmp_order = {}
        for idx, (old_ref, new_ref) in enumerate(old_crosswalk.items()):
            norm_old = ensure_prefixed(normalize_heading_ref(old_ref), manual_type)
            norm_new = ensure_prefixed(normalize_heading_ref(new_ref), manual_type)
            if norm_old and norm_new:
                display_crosswalk[norm_old] = norm_new
                tmp_order[norm_old] = idx
        if tmp_order:
            heading_order = tmp_order
    # Priority 4: Use auto_crosswalk as last resort
    else:
        logger.debug(f"DEBUG [heading_review]: Falling back to auto_crosswalk ({len(auto_crosswalk)} entries)")
        display_crosswalk = auto_crosswalk

    # Persist heading_order for later sorting if we synthesized it
    if heading_order and session_data.get('heading_order', {}) != heading_order:
        session_data['heading_order'] = heading_order
        session_file.write_text(json.dumps(session_data, indent=2), encoding='utf-8')

    logger.debug(f"DEBUG [heading_review]: display_crosswalk before filtering: {len(display_crosswalk)} entries")
    if display_crosswalk:
        sample_items = list(display_crosswalk.items())[:5]
        logger.debug(f"DEBUG [heading_review]: Sample display_crosswalk items: {sample_items}")

    # Filter out blank/None keys and sort by parsed numeric key
    filtered_items = []
    for old_ref, new_ref in display_crosswalk.items():
        if not old_ref or str(old_ref).lower() == "none":
            continue
        if not new_ref or str(new_ref).lower() == "none":
            continue
        norm_old = ensure_prefixed(normalize_heading_ref(old_ref), manual_type)
        norm_new = ensure_prefixed(normalize_heading_ref(new_ref), manual_type)
        if not norm_old or not norm_new:
            continue
        filtered_items.append((norm_old, norm_new))

    logger.debug(f"DEBUG [heading_review]: After filtering: {len(filtered_items)} items remain")

    rows = sorted(
        filtered_items,
        key=lambda x: (
            heading_sort_key(x[0]),
            heading_order.get(x[0], heading_order.get(normalize_heading_ref(x[0]), 10**6))
        )
    )

    render_rows = []
    if heading_map:
        for old_ref, title in heading_map.items():
            is_synthetic = old_ref.startswith("_h")
            suggested_new_ref = ""
            if is_synthetic:
                match_key, _ = find_heading_by_full(title, new_headings)
                if match_key:
                    suggested_new_ref = match_key
            else:
                norm_old = ensure_prefixed(normalize_heading_ref(old_ref), manual_type)
                if not norm_old:
                    continue
                suggested_new_ref = display_crosswalk.get(norm_old) or display_crosswalk.get(old_ref) or old_ref
            render_rows.append({
                "old_ref": old_ref,
                "new_ref": suggested_new_ref or "",
                "old_title_override": title if is_synthetic else "",
                "is_synthetic": is_synthetic
            })
    else:
        for old_ref, new_ref in rows:
            render_rows.append({
                "old_ref": old_ref,
                "new_ref": new_ref,
                "old_title_override": "",
                "is_synthetic": False
            })

    logger.debug(f"DEBUG [heading_review]: After sorting: {len(rows)} rows to display")
    if rows:
        logger.debug(f"DEBUG [heading_review]: First 3 rows: {rows[:3]}")
        logger.debug(f"DEBUG [heading_review]: Last 3 rows: {rows[-3:]}")

    logger.debug(f"DEBUG [heading_review]: new_headings contains {len(new_headings)} entries")
    if new_headings:
        sample_new_keys = list(new_headings.keys())[:5]
        logger.debug(f"DEBUG [heading_review]: Sample new_headings keys: {sample_new_keys}")

    review_html = f"""
    <!doctype html>
    <html lang="en"><head>
    <meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
    <title>Heading Crosswalk Review</title>
    <style>
    body {{ font-family: system-ui, -apple-system, sans-serif; margin: 0; padding: 20px; background: #f7f7f8; }}
    .container {{ max-width: 1400px; margin: 0 auto; background: white; padding: 24px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
    table {{ width: 100%; border-collapse: collapse; margin: 12px 0; table-layout: auto; }}
    th, td {{ border: 1px solid #e5e7eb; padding: 8px; text-align: left; vertical-align: top; }}
    th:nth-child(n+2), td:nth-child(n+2) {{ text-align: center; }}
    td input[type=text] {{ text-align: center; }}
    th {{ background: #f3f4f6; font-weight: 600; }}
    /* Column 1 (Valid checkbox): narrow */
    td:nth-child(1), th:nth-child(1) {{ width: 40px; text-align: center; }}
    /* Column 2 (OLD Reference): wider to show full reference */
    td:nth-child(2) {{ white-space: nowrap; font-family: monospace; min-width: 150px; max-width: 250px; }}
    th:nth-child(2) {{ min-width: 150px; }}
    /* Column 3 (OLD Title): allow wrapping */
    td:nth-child(3) {{ word-wrap: break-word; word-break: break-word; max-width: 300px; }}
    /* Column 4 (NEW Reference input): wider to show full reference */
    td:nth-child(4) {{ min-width: 180px; max-width: 280px; }}
    th:nth-child(4) {{ min-width: 180px; }}
    /* Column 5 (NEW Title): allow wrapping */
    td:nth-child(5) {{ word-wrap: break-word; word-break: break-word; max-width: 300px; }}
    input[type=text] {{ width: 100%; padding: 6px; border: 1px solid #ddd; border-radius: 4px; font-family: monospace; box-sizing: border-box; }}
    .actions {{ margin-top: 16px; display: flex; gap: 12px; align-items: center; }}
    .btn {{ background: #981E32; color: #fff; border: 0; padding: 10px 16px; border-radius: 8px; cursor: pointer; text-decoration: none; }}
    .btn.secondary {{ background: #4b5563; }}
    .small {{ color: #555; font-size: 14px; }}
    </style>
    </head>
    <body>
    <div class="container">
    <h1>Heading Crosswalk Review</h1>
    <p class="small">Workflow: confirm heading numbering, save, then review references before export.</p>
    <p class="small"><strong>Review and validate the automatic heading mappings:</strong></p>
    <ul class="small" style="margin: 8px 0 16px 20px;">
      <li><strong>Uncheck "Valid?"</strong> to exclude invalid rows (TOC entries, malformed references, duplicates)</li>
      <li><strong>Edit "NEW Reference"</strong> to correct wrong conversions (e.g., "Chapter 5A.2.b" → "Chapter 5.1.2.2")</li>
      <li><strong>Edit "NEW Title"</strong> to clean up extracted titles (remove paragraph content, fix special characters, etc.)</li>
      <li><strong>Or upload a CSV/XLSX</strong> file to override all mappings</li>
      <li><em>Note: Un-numbered headings appear with a blank OLD Reference so you can assign a NEW Reference.</em></li>
    </ul>
    <form method="POST" enctype="multipart/form-data">
      <label style="font-weight:600;">Import heading map (optional, CSV/XLSX):</label>
      <input type="file" name="heading_map" accept=".csv,.xlsx" style="margin:6px 0 12px;">
      <div class="actions" style="margin-bottom:12px;">
        <button type="button" class="btn secondary" id="select_all_crosswalk">Select all</button>
        <button type="button" class="btn secondary" id="select_none_crosswalk">Select none</button>
      </div>
      <table>
        <thead><tr><th style="width:40px;">Valid?</th><th>OLD Reference</th><th>OLD Title</th><th>NEW Reference</th><th>NEW Title</th></tr></thead>
        <tbody>
    """
    title_lookup_failures = 0
    old_title_failures = 0
    for idx, row in enumerate(render_rows):
        old_ref = row["old_ref"]
        new_ref = row["new_ref"]
        is_synthetic = row["is_synthetic"]
        # OLD title: lookup from heading_map using the old reference
        # Enable debug for first 3 lookups
        old_title = row["old_title_override"] or lookup_heading_title(old_ref, heading_map, debug=(idx < 3))
        if not old_title:
            old_title_failures += 1

        # NEW title: The heading text doesn't change, only the numbering does
        # So use the same title from heading_map. If we're displaying this row,
        # it means we have an old_ref that maps to a heading, so use that same title.
        new_title = old_title  # Start with same title as old

        # But if new_headings has an actual entry (from HTML conversion), prefer that
        if new_headings:
            scraped_title = new_headings.get(new_ref, {}).get("text", "")
            if not scraped_title:
                alt_key = ensure_prefixed(normalize_heading_ref(new_ref), manual_type)
                scraped_title = new_headings.get(alt_key, {}).get("text", "")
            if scraped_title:
                new_title = scraped_title
            elif not new_title and idx < 3:  # Log first 3 lookup failures
                logger.error(f"DEBUG [heading_review]: NEW title lookup failed for '{new_ref}', tried alt_key '{alt_key}', using old_title as fallback")
                title_lookup_failures += 1
        if not old_title and new_title:
            old_title = new_title
        # Create a stable ID for this heading entry
        heading_id = f"head_{idx}"

        # Escape values for HTML attributes to prevent breaking the form
        escaped_new_ref = new_ref.replace("'", "&#39;").replace('"', "&quot;")
        escaped_new_title = new_title.replace("'", "&#39;").replace('"', "&quot;")
        escaped_old_ref = old_ref.replace("'", "&#39;").replace('"', "&quot;")
        display_old_ref = "" if is_synthetic else old_ref
        checked_attr = "checked" if not is_synthetic else ""

        review_html += (
            "<tr>"
            f"<td style='text-align:center;'><input type='checkbox' name='valid_{heading_id}' {checked_attr}></td>"
            f"<td>{display_old_ref}</td>"
            f"<td>{old_title}</td>"
            f"<td><input type='text' name='new_ref_{heading_id}' value='{escaped_new_ref}'></td>"
            f"<td><input type='text' name='new_title_{heading_id}' value='{escaped_new_title}' style='width:100%;'></td>"
            # Hidden field to store the original old_ref for lookup
            f"<input type='hidden' name='old_ref_{heading_id}' value='{escaped_old_ref}'>"
            "</tr>"
        )

    if old_title_failures > 0:
        logger.debug(f"DEBUG [heading_review]: Total OLD title lookup failures: {old_title_failures}")
    if title_lookup_failures > 0:
        logger.debug(f"DEBUG [heading_review]: Total NEW title lookup failures: {title_lookup_failures} (but used old_title as fallback)")

    review_html += """
        </tbody>
      </table>
      <div class="actions">
        <button type="submit" class="btn">Save & Continue</button>
        <a href="{{ url_for('index') }}" class="btn secondary">Cancel</a>
      </div>
    </form>
    <script>
      (function(){
        const allBtn = document.getElementById('select_all_crosswalk');
        const noneBtn = document.getElementById('select_none_crosswalk');
        const boxes = () => Array.from(document.querySelectorAll('input[type="checkbox"][name^="valid_"]'));
        if (allBtn) {
          allBtn.addEventListener('click', function(){
            boxes().forEach(b => { b.checked = true; });
          });
        }
        if (noneBtn) {
          noneBtn.addEventListener('click', function(){
            boxes().forEach(b => { b.checked = false; });
          });
        }
      })();
    </script>
    </div>
    </body>
    </html>
    """
    return review_html

def load_heading_id_map_from_request() -> tuple[dict, str]:
    """Parse signature-to-ID JSON from form or file."""
    raw_text = request.form.get("stable_heading_map_raw", "").strip()
    if not raw_text:
        f = request.files.get("stable_heading_map_file")
        if f and f.filename:
            raw_text = f.read().decode("utf-8", errors="ignore")
    
    return parse_heading_id_map_json(raw_text), raw_text

HOME_PAGE = """
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>WSU Manual Converter</title>
  {{ wordpress_css_tag|safe }}
  <style>
    body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; line-height: 1.5; color: #333; max-width: 1200px; margin: 0 auto; padding: 40px 20px; background: #f9f9f9; }
    h1 { color: #8d0a0a; margin-bottom: 30px; border-bottom: 2px solid #8d0a0a; padding-bottom: 10px; }
    .card { background: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.05); margin-bottom: 30px; }
    .form-group { margin-bottom: 20px; }
    label { display: block; margin-bottom: 8px; font-weight: 600; }
    input[type="file"] { display: block; width: 100%; padding: 10px; border: 1px dashed #ccc; border-radius: 4px; background: #fafafa; }
    button { background: #8d0a0a; color: white; border: none; padding: 12px 24px; border-radius: 6px; font-size: 16px; font-weight: 600; cursor: pointer; transition: background 0.2s; }
    button:hover { background: #6e0808; }
    .small { font-size: 13px; color: #666; }
    .checkbox { display: flex; align-items: flex-start; gap: 10px; margin-bottom: 15px; }
    .checkbox input { margin-top: 4px; }
    .checkbox label { font-weight: normal; margin-bottom: 0; }
    .alert { padding: 15px; border-radius: 6px; margin-bottom: 20px; background: #fee2e2; color: #991b1b; border: 1px solid #fecaca; }
    .section-title { font-size: 18px; font-weight: 600; margin-top: 30px; margin-bottom: 15px; color: #555; }
    details { margin-bottom: 15px; border: 1px solid #eee; border-radius: 4px; padding: 10px; }
    summary { font-weight: 600; cursor: pointer; }
    .preview-container { margin-top: 40px; border-top: 4px solid #8d0a0a; padding-top: 20px; }
    .export-actions { position: sticky; top: 0; background: white; padding: 15px; border-bottom: 1px solid #eee; z-index: 100; display: flex; gap: 10px; }
  </style>
</head>
<body>
  <h1>WSU Manual Converter</h1>

  {% if not show_preview %}
  <div class="card" style="padding: 18px 24px; margin-bottom: 24px; background: #fffbeb; border: 1px solid #fcd34d;">
    <p class="small" style="margin: 0; color: #78350f; font-size: 15px; line-height: 1.5;">
      <strong>How to use this tool:</strong>
      <a href="{{ url_for('instructions') }}" style="color: #8d0a0a; font-weight: 600;">Open full instructions</a>
      for setup, step-by-step conversion, WordPress deployment, heading maps, and troubleshooting.
    </p>
  </div>
  {% endif %}

  {% with messages = get_flashed_messages() %}
    {% if messages %}
      {% for message in messages %}
        <div class="alert">{{ message }}</div>
      {% endfor %}
    {% endif %}
  {% endwith %}

  {% if show_preview %}
    <div class="export-actions" style="display:flex; flex-wrap:wrap; gap:10px; align-items:center; margin-bottom:16px;">
      <a href="{{ url_for('download', token=token, kind='standalone') }}" class="btn" style="background:#4b5563; text-decoration:none; padding:12px 24px; color:white; border-radius:6px; font-weight:600; display:inline-block;">Download Standalone</a>
      <a href="{{ url_for('download', token=token, kind='fragment') }}" class="btn" style="background:#16a34a; text-decoration:none; padding:12px 24px; color:white; border-radius:6px; font-weight:600; display:inline-block;">Download Fragment</a>
      <a href="{{ url_for('download', token=token, kind='fragment_css') }}" class="btn" style="background:#059669; text-decoration:none; padding:12px 24px; color:white; border-radius:6px; font-weight:600; display:inline-block;">Fragment + CSS</a>
      <a href="{{ url_for('download', token=token, kind='docx') }}" class="btn" style="background:#7c3aed; text-decoration:none; padding:12px 24px; color:white; border-radius:6px; font-weight:600; display:inline-block;">Download DOCX</a>
      <a href="{{ url_for('download', token=token, kind='css') }}" class="btn" style="background:#0891b2; text-decoration:none; padding:12px 24px; color:white; border-radius:6px; font-weight:600; display:inline-block;">Download CSS</a>
      <a href="{{ url_for('download', token=token, kind='js') }}" class="btn" style="background:#d97706; text-decoration:none; padding:12px 24px; color:white; border-radius:6px; font-weight:600; display:inline-block;">Download JS</a>
      <a href="{{ url_for('download', token=token, kind='heading_map') }}" class="btn" style="background:#be185d; text-decoration:none; padding:12px 24px; color:white; border-radius:6px; font-weight:600; display:inline-block;">Download Heading Map</a>
      <form action="{{ url_for('export_session', session_id=session_id) }}" method="post" style="display:inline;">
        <button type="submit" style="background:#0284c7">Export Session Bundle (.zip)</button>
      </form>
      <a href="{{ url_for('index') }}" style="margin-left:auto; align-self:center;">Start Over</a>
    </div>
    <div class="preview-container">
      {{ body_html|safe }}
    </div>
    {{ wordpress_js_tag|safe }}
  {% else %}
    <div class="card">
      <form action="/convert" method="post" enctype="multipart/form-data">
        <div class="form-group">
          <label for="docx">Upload DOCX Manual:</label>
          <input type="file" id="docx" name="docx" accept=".docx" required>
          <p class="small">Standard Microsoft Word .docx file.</p>
        </div>

        <div class="section-title">Conversion Options</div>
        
        <div class="checkbox">
          <input id="edit_tables" name="edit_tables" type="checkbox">
          <label for="edit_tables">Open table review before export (if tables are detected)</label>
        </div>
        
        <div class="checkbox">
          <input id="preserve_numbers" name="preserve_numbers" type="checkbox">
          <label for="preserve_numbers">Keep heading numbers in text (e.g., "I.A. Section Title" instead of just "Section Title")</label>
        </div>
        <p class="small" style="margin-top:-8px; margin-left:24px; color:#666; margin-bottom:12px;">By default, the converter strips numbers like "Chapter 1" because the site CSS adds them automatically.</p>
        
        <div class="form-group">
          <label for="toc_depth">Table of Contents Depth:</label>
          <select id="toc_depth" name="toc_depth" style="padding: 6px 10px; border: 1px solid #ddd; border-radius: 6px; font-size: 14px; width: 200px;">
            <option value="1">Level 1 only (H1)</option>
            <option value="2" selected>Levels 1-2 (H1-H2)</option>
            <option value="3">Levels 1-3 (H1-H3)</option>
            <option value="4">Levels 1-4 (H1-H4)</option>
            <option value="5">Levels 1-5 (H1-H5)</option>
          </select>
        </div>

        <div class="form-group">
          <label for="mapping_mode">Heading mapping mode:</label>
          <select id="mapping_mode" name="mapping_mode" style="padding: 6px 10px; border: 1px solid #ddd; border-radius: 6px; font-size: 14px; width: 260px;">
            <option value="map_new" selected>Map to new numeric headings (recommended)</option>
            <option value="keep_old">Keep original headings/numbering</option>
          </select>
        </div>

        <details>
          <summary>Advanced: Permalink Continuity (Heading Map)</summary>
          <div style="padding: 10px 0;">
            <p class="small" style="margin-top:0;">Upload a previously exported <code>*.heading-map.json</code> file to reuse stable anchor IDs, or paste the JSON below. If both are provided, the pasted text wins.</p>
            <label for="stable_heading_map_file">Upload heading map JSON:</label>
            <input id="stable_heading_map_file" name="stable_heading_map_file" type="file" accept=".json,application/json" style="margin-bottom:12px;">
            <label for="stable_heading_map_raw">Or paste Signature-to-ID JSON:</label>
            <textarea id="stable_heading_map_raw" name="stable_heading_map_raw" rows="5" style="width:100%; font-family:monospace; font-size:12px;" placeholder='{"my heading": "fixed-id-1"}'></textarea>
          </div>
        </details>

        <div style="margin-top: 30px;">
          <button type="submit">Upload and Review</button>
        </div>
      </form>
    </div>

    <div class="section-title">Import & Restore</div>
    <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 20px;">
      <div class="card">
        <h3 style="margin-top:0">Session Bundle</h3>
        <form action="/import_bundle" method="post" enctype="multipart/form-data">
          <input type="file" name="bundle" accept=".zip" required style="margin-bottom:10px">
          <button type="submit" style="padding: 8px 16px; font-size:14px">Restore Session</button>
        </form>
      </div>
      <div class="card">
        <h3 style="margin-top:0">WordPress HTML</h3>
        <form action="/import_html" method="post" enctype="multipart/form-data">
          <input type="file" name="html_file" accept=".html,.htm" required style="margin-bottom:10px">
          <label for="stable_heading_map_file_html" style="font-size:13px;">Heading map JSON (optional):</label>
          <input id="stable_heading_map_file_html" type="file" name="stable_heading_map_file" accept=".json,application/json" style="margin-bottom:10px">
          <button type="submit" style="padding: 8px 16px; font-size:14px">Import HTML</button>
        </form>
      </div>
    </div>
  {% endif %}
</body>
</html>
"""

INSTRUCTIONS_PAGE = """
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Instructions — WSU Manual Converter</title>
  <style>
    body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; line-height: 1.55; color: #333; max-width: 900px; margin: 0 auto; padding: 40px 20px 60px; background: #f9f9f9; }
    h1 { color: #8d0a0a; margin-bottom: 8px; border-bottom: 2px solid #8d0a0a; padding-bottom: 10px; }
    .back { margin-bottom: 20px; }
    .back a { color: #8d0a0a; font-weight: 600; text-decoration: none; }
    .back a:hover { text-decoration: underline; }
    .card { background: white; padding: 28px 32px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.05); }
    .instructions-content :first-child { margin-top: 0; }
    .instructions-content h2 { color: #444; font-size: 1.25rem; margin-top: 1.75rem; padding-bottom: 6px; border-bottom: 1px solid #eee; }
    .instructions-content h3 { color: #555; font-size: 1.05rem; margin-top: 1.35rem; }
    .instructions-content hr { border: none; border-top: 1px solid #e5e5e5; margin: 2rem 0; }
    .instructions-content ul, .instructions-content ol { padding-left: 1.35rem; }
    .instructions-content li { margin-bottom: 0.35rem; }
    .instructions-content code { background: #f3f4f6; padding: 2px 7px; border-radius: 4px; font-size: 0.9em; }
    .instructions-content pre { background: #f3f4f6; padding: 14px 16px; border-radius: 6px; overflow-x: auto; font-size: 0.88rem; line-height: 1.45; }
    .instructions-content pre code { background: none; padding: 0; }
    .instructions-content table { border-collapse: collapse; width: 100%; margin: 1rem 0; font-size: 0.95rem; }
    .instructions-content th, .instructions-content td { border: 1px solid #ddd; padding: 10px 12px; text-align: left; vertical-align: top; }
    .instructions-content th { background: #f9fafb; font-weight: 600; }
    .instructions-content a { color: #8d0a0a; }
    .instructions-content blockquote { margin: 1rem 0; padding-left: 1rem; border-left: 4px solid #e5e7eb; color: #555; }
  </style>
</head>
<body>
  <p class="back"><a href="{{ url_for('index') }}">← Back to converter</a></p>
  <h1>Instructions</h1>
  <div class="card instructions-content">
    {{ content|safe }}
  </div>
</body>
</html>
"""

# --------------------------- Routes ---------------------------

@app.route("/instructions")
def instructions():
    instructions_path = Path(__file__).resolve().parent / "instructions.md"
    if not instructions_path.is_file():
        flash("Instructions are temporarily unavailable.")
        return redirect(url_for("index"))
    md_text = instructions_path.read_text(encoding="utf-8")
    content_html = markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "sane_lists", "nl2br"],
    )
    return render_template_string(INSTRUCTIONS_PAGE, content=content_html)

@app.route("/", methods=["GET"])
def index():
    # Load WordPress CSS and JS for preview
    wp_css_text = get_wp_css_text()
    wp_js_text = get_wp_js_text()
    wp_css = f"<style>{wp_css_text}</style>" if wp_css_text else ""
    wp_js = f"<script>{wp_js_text}</script>" if wp_js_text else ""

    return render_template_string(
        HOME_PAGE,
        show_preview=False,
        hide_upload=False,
        wordpress_css_tag=wp_css,
        wordpress_js_tag=wp_js
    )

@app.route("/convert", methods=["POST"])
def convert():
    f = request.files.get("docx")
    if not f or f.filename == "":
        flash("Please choose a .docx file.")
        return redirect(url_for("index"))

    # Phase 1: Initialize Session
    session_id = str(uuid.uuid4())
    session = SessionDir(session_id)
    
    preserve = bool(request.form.get("preserve_numbers"))
    mapping_mode = request.form.get("mapping_mode", "map_new")
    infer_heading_depth = bool(request.form.get("infer_heading_depth"))
    strip_docx_formatting = "strip_docx_formatting" in request.form
    toc_depth = request.form.get("toc_depth", "2")  # Default to 2 if not provided
    edit_tables = "edit_tables" in request.form
    stable_heading_map, stable_heading_map_raw = load_heading_id_map_from_request()
    try:
        toc_depth = int(toc_depth)
        if toc_depth < 1 or toc_depth > 5:
            toc_depth = 2  # Validate range
    except (ValueError, TypeError):
        toc_depth = 2  # Default if invalid
    
    filename = secure_filename(f.filename)
    if not filename.lower().endswith(".docx"):
        flash("Only .docx files are supported.")
        return redirect(url_for("index"))

    src = session.source_docx
    f.save(str(src))
    try:
        style_map = {}
        sequence_map = {}
        if mapping_mode == "map_new" and infer_heading_depth:
            style_map, sequence_map = extract_style_map_from_reference(src)

        pre = session.pre_docx
        temp_html = session.temp_html

        logger.debug("="*80)
        logger.info(f"SESSION: {session_id}")
        logger.info("PHASE 1: PREPROCESSING & EXTRACTING OLD REFERENCES")
        logger.debug("="*80)

        # Preprocess and extract OLD references from DOCX
        heading_map, old_crosswalk, references, manual_type = preprocess_docx(src, pre, style_map, sequence_map)
        docx_links_by_para = extract_docx_hyperlinks(Document(src))
        logger.info(f"Extracted {len(references)} OLD references from DOCX")

        logger.debug("="*80)
        logger.info("PHASE 2: PRE-CONVERSION TO GET NEW STRUCTURE")
        logger.debug("="*80)

        # Do preliminary conversion to HTML to get NEW numbering
        run_pandoc(pre, temp_html)
        logger.info(f"Generated preliminary HTML: {temp_html}")

        # Read and process the converted HTML (same as do_convert)
        temp_html_content = temp_html.read_text(encoding='utf-8', errors='ignore')
        temp_html_content = strip_pandoc_styles(normalize_spaces(temp_html_content))
        temp_html_content = strip_images_and_figures(temp_html_content)
        body = extract_body(temp_html_content)

        # Wrap body in .manual so the list normalizer has context
        wrapped = f'<div class="manual">{body}</div>'

        # Normalize lists
        normalized_html = normalize_typed_lists(wrapped)

        # Strip any TOC sections
        normalized_html = strip_toc_sections_dom(normalized_html)

        # Align preserve flag with mapping choice: keeping old headings implies keeping numbers
        if mapping_mode == "keep_old":
            preserve = True

        if mapping_mode == "map_new" and infer_heading_depth:
            if not style_map:
                flash("No style map found in the source DOCX; falling back to prefix-based inference.")
            normalized_html = infer_heading_levels_from_prefix(normalized_html, style_map if style_map else None)

        # Strip old heading numbers (so CSS counters apply correctly)
        if not preserve:
            normalized_html, _ = strip_heading_numbers_dom(normalized_html)

        # Apply numeric numbering to headings (mimics CSS counters)
        normalized_html = apply_css_counter_numbering(normalized_html, manual_type, preserve=preserve)

        # Add heading IDs for scraping
        normalized_html = add_heading_ids(normalized_html, stable_map=stable_heading_map)

        logger.debug("="*80)
        logger.info("PHASE 3: SCRAPING NEW HEADING STRUCTURE")
        logger.debug("="*80)

        # Scrape NEW headings from the converted HTML
        new_headings = scrape_heading_structure_from_html(normalized_html)
        logger.info(f"Scraped {len(new_headings)} NEW headings from converted HTML")

        logger.debug("="*80)
        logger.info("PHASE 4: AUTO-MATCHING OLD->NEW")
        logger.debug("="*80)

        # Auto-match OLD references to NEW headings
        # BUT: If keeping old numbering, create identity crosswalk (OLD->OLD) instead
        if mapping_mode == "keep_old":
            logger.info("Mode: keep_old - Building identity crosswalk (OLD->OLD)")
            auto_crosswalk = {}
            for ref in references:
                old_ref_text = ref[2]  # Extract the reference string
                # Skip if it doesn't look like a chapter/section reference
                if re.match(r'^(Chapter|Section)\s+[\dIVXLCDM]+', old_ref_text, re.IGNORECASE):
                    # Identity mapping: OLD stays as OLD
                    auto_crosswalk[old_ref_text] = old_ref_text
            logger.info(f"Built identity crosswalk with {len(auto_crosswalk)} OLD->OLD mappings")
        else:
            logger.info("Mode: map_new - Auto-matching OLD->NEW")
            auto_crosswalk = auto_match_old_to_new_references(references, new_headings, manual_type=manual_type)
            logger.info(f"Auto-matched {len(auto_crosswalk)} OLD->NEW mappings")

        # Store data in session for crosswalk editor
        import json
        heading_order = {
            ensure_prefixed(normalize_heading_ref(k), manual_type): idx
            for idx, k in enumerate(heading_map.keys())
            if normalize_heading_ref(k)
        }
        theme_settings, _ = coerce_theme_settings(None, manual_type)
        session_data = {
            'old_crosswalk': old_crosswalk,  # From Word doc (likely empty)
            'auto_crosswalk': auto_crosswalk,  # AUTO-matched OLD->NEW
            'approved_crosswalk': {},  # To be filled in heading review
            'new_headings': new_headings,  # NEW structure from HTML
            'references': references,  # OLD references found in text
            'heading_map': heading_map,  # OLD heading map (likely empty)
            'heading_order': heading_order,
            'manual_type': manual_type,
            'filename': filename,
            'src_path': str(src),
            'pre_path': str(pre),
            'temp_html_path': str(temp_html),
            'toc_depth': toc_depth,
            'preserve_numbers': preserve,
            'mapping_mode': mapping_mode,
            'strip_docx_formatting': strip_docx_formatting,
            'theme_settings': theme_settings,
            'heading_edits': {},
            'style_panels': {"doc": True, "toc": False, "heading": False},
            'infer_heading_depth': infer_heading_depth,
            'infer_style_map': style_map,
            'infer_sequence_map': serialize_sequence_map(sequence_map),
            'stable_heading_map': stable_heading_map,
            'stable_heading_map_raw': stable_heading_map_raw,
            'docx_links_by_para': docx_links_by_para,
            'edit_tables': edit_tables,
        }
        session_file = SessionDir(session_id).session_json
        session_file.write_text(json.dumps(session_data, indent=2, default=str), encoding='utf-8')

        logger.debug("="*80)
        logger.info(f"Session saved: {session_id}")
        logger.debug("="*80)

        # Redirect to heading crosswalk if mapping to new numbers; otherwise go to reference review
        if session_data['mapping_mode'] == "keep_old":
            return redirect(url_for('review', session_id=session_id))
        return redirect(url_for('heading_review', session_id=session_id))

    except subprocess.CalledProcessError as e:
        flash(f"Pandoc conversion failed. Is Pandoc installed and on PATH? ({e})")
        return redirect(url_for("index"))
    except Exception as e:
        flash(f"Conversion failed: {e}")
        return redirect(url_for("index"))

@app.route("/review/<session_id>", methods=["GET", "POST"])
def review(session_id):
    """Review crosswalk and references before conversion"""
    import json
    session = SessionDir(session_id)
    session_file = session.session_json
    if not session_file.exists():
        flash("Session expired or invalid.")
        return redirect(url_for("index"))

    session_data = json.loads(session_file.read_text(encoding='utf-8'))

    # NEW WORKFLOW: Load auto_crosswalk and new_headings from pre-conversion
    auto_crosswalk = session_data.get('approved_crosswalk') or session_data.get('auto_crosswalk', {})
    approved_crosswalk = session_data.get('approved_crosswalk', {})
    new_headings = session_data.get('new_headings', {})
    references = session_data.get('references', [])
    html_import = session_data.get('html_import', False)
    rebuild_links = session_data.get('rebuild_links', False)
    edit_tables = session_data.get('edit_tables', False)
    mapping_mode = session_data.get('mapping_mode', 'map_new')
    linked_ref_count = 0
    if html_import and references:
        for ref in references:
            if len(ref) > 5 and ref[5]:
                linked_ref_count += 1
    has_tables = False
    src_path = Path(session_data.get('src_path', ''))
    html_path = Path(session_data.get('html_path', ''))
    if html_import and html_path.exists():
        has_tables = has_tables_in_html(html_path)
    elif src_path.exists():
        has_tables = has_tables_in_docx(src_path)

    page_size = 50
    try:
        page = int(request.values.get('page', '1'))
    except Exception:
        page = 1
    # Group references by paragraph for display
    # Use stable IDs based on content, not sequential order
    ref_by_para = {}
    for ref in references:
        para_idx = ref[0]
        full_text = ref[1]
        old_ref = ref[2]
        start_pos = ref[3]
        # Generate stable ID based on location and content
        stable_id = generate_stable_ref_id(para_idx, start_pos, old_ref)
        if para_idx not in ref_by_para:
            ref_by_para[para_idx] = []
        ref_by_para[para_idx].append((ref, stable_id))
    para_keys = sorted(ref_by_para.keys())
    total_paras = len(para_keys)
    total_pages = max(1, math.ceil(total_paras / page_size))
    page = max(1, min(page, total_pages))
    start_idx = (page - 1) * page_size
    end_idx = start_idx + page_size
    page_para_keys = para_keys[start_idx:end_idx]
    showing_start = start_idx + 1 if total_paras else 0
    showing_end = min(end_idx, total_paras)
    old_crosswalk = session_data.get('old_crosswalk', {})  # Keep for compatibility
    manual_type = session_data.get('manual_type', 'chapter')
    filename = session_data.get('filename', '')

    logger.debug(f"DEBUG: Session loaded - new_headings has {len(new_headings)} entries, references has {len(references)} entries")
    logger.debug(f"DEBUG: Auto-matched {len(auto_crosswalk)} OLD->NEW references")
    if len(new_headings) == 0:
        logger.error("DEBUG: WARNING - new_headings is empty! Pre-conversion may have failed.")
    else:
        logger.debug(f"DEBUG: Sample new headings: {list(new_headings.keys())[:3]}")
    
    if request.method == 'POST':
        # CRITICAL DEBUG: Log ALL form data at the very start
        logger.debug("="*80)
        logger.debug("DEBUG: POST REQUEST RECEIVED")
        logger.debug("="*80)
        logger.debug(f"DEBUG: Request method: {request.method}")
        logger.debug(f"DEBUG: Form has {len(request.form)} total keys")
        logger.debug(f"DEBUG: All form keys: {list(request.form.keys())}")
        logger.debug(f"DEBUG: Raw form data (first 20 items): {dict(list(request.form.items())[:20])}")
        logger.debug(f"DEBUG: 'save_edits' in form? {('save_edits' in request.form)}")
        logger.debug(f"DEBUG: 'proceed' in form? {('proceed' in request.form)}")
        logger.debug("="*80)

        # Save on any POST so Enter submits don't drop changes
        if True:
            logger.debug(f"DEBUG: SAVE HANDLER EXECUTING - Save request received. Form keys: {list(request.form.keys())[:10]}...")

            # Save reference edits, validations, and link targets to persistent file
            edits = {}
            validations = {}
            link_targets = {}
            ignored = {}
            external_urls = {}
            if html_import:
                rebuild_links = 'rebuild_links' in request.form
                session_data['rebuild_links'] = rebuild_links
            # edit_tables was set during initial upload; preserve it from session
            edit_tables = session_data.get('edit_tables', False)
            edit_file = session.edits_json
            if edit_file.exists():
                try:
                    existing_data = json.loads(edit_file.read_text(encoding='utf-8'))
                    edits = existing_data.get('reference_edits', {}) or {}
                    validations = existing_data.get('reference_validations', {}) or {}
                    link_targets = existing_data.get('reference_link_targets', {}) or {}
                    ignored = existing_data.get('reference_ignored', {}) or {}
                    external_urls = existing_data.get('reference_external_urls', {}) or {}
                except Exception:
                    edits = {}
                    validations = {}
                    link_targets = {}
                    ignored = {}
                    external_urls = {}

            # Collect all reference IDs from the form
            all_ref_ids = set()
            for key in request.form.keys():
                if key.startswith('ref_edit_'):
                    ref_id = key.replace('ref_edit_', '')
                    all_ref_ids.add(ref_id)
                elif key.startswith('ref_valid_'):
                    ref_id = key.replace('ref_valid_', '')
                    all_ref_ids.add(ref_id)
                elif key.startswith('ref_target_'):
                    ref_id = key.replace('ref_target_', '')
                    all_ref_ids.add(ref_id)
                elif key.startswith('ref_ignore_'):
                    ref_id = key.replace('ref_ignore_', '')
                    all_ref_ids.add(ref_id)
                elif key.startswith('ref_external_'):
                    ref_id = key.replace('ref_external_', '')
                    all_ref_ids.add(ref_id)

            logger.debug(f"DEBUG: Collected {len(all_ref_ids)} reference IDs from form: {sorted(list(all_ref_ids))[:10]}")

            # Process edits, validations, and targets
            for ref_id in all_ref_ids:
                # ref_id is already a stable ID (e.g., "ref_42_123_a1b2c3d4"), use it directly
                edit_key = ref_id

                # Save edit text
                edit_value = request.form.get(f'ref_edit_{ref_id}', '').strip()
                if edit_value:
                    edits[edit_key] = edit_value
                elif edit_key in edits:
                    edits.pop(edit_key, None)

                # Save validation status (checkbox checked = valid)
                is_valid = f'ref_valid_{ref_id}' in request.form
                is_ignored = f'ref_ignore_{ref_id}' in request.form
                if is_ignored:
                    is_valid = False
                validations[edit_key] = is_valid
                if is_ignored:
                    ignored[edit_key] = True
                elif edit_key in ignored:
                    ignored.pop(edit_key, None)

                # Save link target (heading to link to)
                target_value = request.form.get(f'ref_target_{ref_id}', '').strip()
                if target_value:
                    link_targets[edit_key] = target_value
                elif edit_key in link_targets:
                    link_targets.pop(edit_key, None)

                # Save external URL (link to other manuals)
                external_value = request.form.get(f'ref_external_{ref_id}', '').strip()
                if external_value:
                    external_urls[edit_key] = external_value
                elif edit_key in external_urls:
                    external_urls.pop(edit_key, None)

                # Debug output
                logger.debug(f"DEBUG: Saving {ref_id}: valid={is_valid}, ignored={is_ignored}, edit='{edit_value[:30] if edit_value else '(none)'}...', target='{target_value[:30] if target_value else '(none)'}...', external='{external_value[:30] if external_value else '(none)'}...'")

            # Save to persistent edit file
            edit_data = {
                'document': filename,
                'auto_crosswalk': auto_crosswalk,
                'approved_crosswalk': approved_crosswalk,
                'reference_edits': edits,
                'reference_validations': validations,
                'reference_link_targets': link_targets,
                'reference_ignored': ignored,
                'reference_external_urls': external_urls,
                'last_updated': str(datetime.now())
            }
            edit_file.write_text(json.dumps(edit_data, indent=2), encoding='utf-8')
            session_data['approved_crosswalk'] = approved_crosswalk
            session_data['reference_edits'] = edits
            session_data['reference_validations'] = validations
            session_data['reference_link_targets'] = link_targets
            session_data['reference_ignored'] = ignored
            session_data['reference_external_urls'] = external_urls
            session_file.write_text(json.dumps(session_data, indent=2), encoding='utf-8')

            # Count valid vs invalid references
            valid_count = sum(1 for v in validations.values() if v)
            invalid_count = sum(1 for v in validations.values() if not v)

            # Debug: print save location
            logger.debug(f"DEBUG: Saved edits to: {edit_file}")
            logger.debug(f"DEBUG: Saved data - edits: {len(edits)} entries, validations: {len(validations)} entries ({valid_count} valid, {invalid_count} invalid), link_targets: {len(link_targets)} entries, external_urls: {len(external_urls)} entries")
            logger.debug(f"DEBUG: Sample validations: {dict(list(validations.items())[:5])}")

            flash(f"✓ Edits saved successfully! Found {valid_count} valid references, {invalid_count} skipped references. Saved to: {edit_file.name}. Export a session bundle to share or continue on another machine.")
            if 'proceed' in request.form:
                if edit_tables and has_tables:
                    return redirect(url_for('table_review', session_id=session_id))
                return redirect(url_for('do_convert', session_id=session_id))
            if 'next_page' in request.form:
                target_page = min(page + 1, total_pages)
                return redirect(url_for('review', session_id=session_id, page=target_page))
            if 'prev_page' in request.form:
                target_page = max(page - 1, 1)
                return redirect(url_for('review', session_id=session_id, page=target_page))
            return redirect(url_for('review', session_id=session_id, page=page))
        return redirect(url_for('review', session_id=session_id, page=page))
    
    # Use approved_crosswalk if provided, else auto_crosswalk for displaying OLD->NEW mappings
    display_crosswalk = approved_crosswalk if approved_crosswalk else auto_crosswalk

    logger.debug(f"DEBUG [review]: Using {'approved_crosswalk' if approved_crosswalk else 'auto_crosswalk'} with {len(display_crosswalk)} entries")
    if display_crosswalk:
        sample_items = list(display_crosswalk.items())[:5]
        logger.debug(f"DEBUG [review]: Sample display_crosswalk: {sample_items}")
    
    # Build review page HTML
    review_html = f"""
    <!doctype html>
    <html lang="en"><head>
    <meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
    <title>Review Crosswalk - WSU Manual Converter</title>
    <style>
    body {{ font-family: system-ui, -apple-system, sans-serif; margin: 0; padding: 20px; background: #f7f7f8; }}
    .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 24px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
    h1 {{ color: #981E32; margin-top: 0; }}
    h2 {{ color: #333; margin-top: 24px; border-bottom: 2px solid #ddd; padding-bottom: 8px; }}
    .crosswalk-table {{ width: 100%; border-collapse: collapse; margin: 16px 0; table-layout: auto; }}
    .crosswalk-table th, .crosswalk-table td {{ padding: 10px; text-align: left; border: 1px solid #ddd; vertical-align: top; }}
    .crosswalk-table th {{ background: #f5f5f5; font-weight: 600; }}
    .crosswalk-table tr:nth-child(even) {{ background: #fafafa; }}
    /* Reference columns: wider, no wrap, monospace */
    .crosswalk-table td:nth-child(1), .crosswalk-table td:nth-child(3) {{
        white-space: nowrap;
        font-family: monospace;
        min-width: 150px;
        max-width: 250px;
    }}
    .crosswalk-table th:nth-child(1), .crosswalk-table th:nth-child(3) {{
        min-width: 150px;
        max-width: 250px;
    }}
    /* Title columns: allow wrapping, normal font */
    .crosswalk-table td:nth-child(2), .crosswalk-table td:nth-child(4) {{
        word-wrap: break-word;
        word-break: break-word;
        max-width: 300px;
    }}
    .references-list {{ margin: 16px 0; }}
    .reference-item {{ padding: 12px; margin: 8px 0; background: #fff7ed; border-left: 4px solid #f59e0b; border-radius: 4px; }}
    .reference-item.invalid {{ background: #fee; border-left-color: #dc2626; opacity: 0.7; }}
    .reference-text {{ font-family: monospace; background: #fff; padding: 4px 8px; border-radius: 4px; display: inline-block; margin: 4px 0; }}
    .reference-review-item.ref-collapsed {{ background: #f1f5f9; border: 1px dashed #cbd5e1; padding: 6px !important; margin: 6px 0 !important; }}
    .reference-review-item.ref-collapsed .ref-details {{ display: none; }}
    .reference-review-item.ref-collapsed .ref-summary {{ margin-bottom: 0; }}
    .skip-badge {{ display: none; margin-left: 8px; font-size: 11px; font-weight: 600; letter-spacing: 0.02em; text-transform: uppercase; background: #64748b; color: #fff; padding: 2px 6px; border-radius: 999px; }}
    .reference-review-item.ref-collapsed .skip-badge {{ display: inline-flex; align-items: center; }}
    .highlight {{ background: yellow; padding: 2px 4px; font-weight: 600; }}
    .validation-checkbox {{ margin-bottom: 8px; padding: 8px; background: #f0fdf4; border: 1px solid #86efac; border-radius: 4px; }}
    .validation-checkbox.invalid {{ background: #fef2f2; border-color: #fca5a5; }}
    .validation-checkbox label {{ display: flex; align-items: center; gap: 8px; cursor: pointer; font-weight: 600; }}
    .validation-checkbox input[type="checkbox"] {{ width: 18px; height: 18px; cursor: pointer; }}
    .actions {{ margin-top: 24px; padding-top: 24px; border-top: 2px solid #ddd; }}
    button {{ background: #981E32; color: white; border: 0; padding: 12px 24px; border-radius: 8px; cursor: pointer; font-size: 16px; }}
    button:hover {{ background: #7a1728; }}
    .info {{ background: #e0f2fe; border-left: 4px solid #0284c7; padding: 12px; margin: 16px 0; border-radius: 4px; }}
    .flash-messages {{ margin: 16px 0; }}
    .flash-message {{ padding: 12px 16px; margin: 8px 0; border-radius: 8px; border-left: 4px solid; }}
    .flash-message.success {{ background: #d1fae5; border-left-color: #10b981; color: #065f46; }}
    .flash-message.error {{ background: #fee2e2; border-left-color: #ef4444; color: #991b1b; }}
    .flash-message.info {{ background: #dbeafe; border-left-color: #3b82f6; color: #1e40af; }}
    .debug-panel {{ background: #f3f4f6; border: 2px solid #9ca3af; padding: 12px; margin: 16px 0; border-radius: 8px; font-family: monospace; font-size: 12px; }}
    .debug-panel summary {{ cursor: pointer; font-weight: 600; color: #374151; }}
    .debug-panel pre {{ margin: 8px 0; overflow-x: auto; }}
    </style>
    </head>
    <body>
    <div class="container">
    <h1>Review Crosswalk and References</h1>
    <p class="small">Workflow: validate references, save, then proceed to export.</p>
    """

    # Load existing edits, validations, and link targets FIRST (before using them)
    # Standardized edit file path
    session = SessionDir(session_id)
    edit_file = session.edits_json
    existing_edits = {}
    existing_validations = {}
    existing_link_targets = {}
    existing_ignored = {}
    existing_external_urls = {}
    if edit_file.exists():
        try:
            edit_data = json.loads(edit_file.read_text(encoding='utf-8'))
            existing_edits = edit_data.get('reference_edits', {})
            existing_validations = edit_data.get('reference_validations', {})
            existing_link_targets = edit_data.get('reference_link_targets', {})
            existing_ignored = edit_data.get('reference_ignored', {})
            existing_external_urls = edit_data.get('reference_external_urls', {})
            valid_count_loaded = sum(1 for v in existing_validations.values() if v)
            logger.debug(f"DEBUG: Loaded {len(existing_edits)} edits, {len(existing_validations)} validations ({valid_count_loaded} valid) from: {edit_file}")
            logger.debug(f"DEBUG: Sample loaded validations: {dict(list(existing_validations.items())[:5])}")
        except Exception as e:
            logger.error(f"DEBUG: Failed to load edits from {edit_file}: {e}")
            pass
    else:
        logger.debug(f"DEBUG: No existing edits file found at: {edit_file}")

    # Build list of all NEW headings for dropdown (document order when available)
    dropdown_headings = []
    numbered_headings = []
    unnumbered_headings = []

    def looks_numbered_heading(text: str) -> bool:
        if not text:
            return False
        if re.match(r'^(Chapter|Section)\s+[\w\dIVXLCDM]+', text, re.IGNORECASE):
            return True
        if re.match(r'^\d+(?:\.\d+)+\s+', text):
            return True
        if re.match(r'^[IVXLCDM]+\.[A-Z0-9]+', text, re.IGNORECASE):
            return True
        if re.match(r'^[A-Z]\.\s+', text):
            return True
        return False
    
    # Prefer document order if available; fallback to numeric sort.
    for heading_key, heading_data in sorted(new_headings.items(), key=heading_dropdown_sort):
        full_text = heading_data.get('full', '')
        synthetic_key = heading_key.startswith(('H1:', 'H2:', 'H3:', 'H4:', 'H5:', 'H6:'))
        is_numbered = looks_numbered_heading(full_text)
        # Include numbered headings even if they came from H1/H2 synthetic keys.
        if not synthetic_key or is_numbered:
            if is_numbered:
                numbered_headings.append(heading_data)
            else:
                if len(heading_data['text'].split()) <= 15:
                    unnumbered_headings.append(heading_data)
        else:
            # For un-numbered headings, only include if they are reasonably short.
            # This filters out full paragraphs that Pandoc might have marked as headings.
            if len(heading_data['text'].split()) <= 15:
                unnumbered_headings.append(heading_data)

    unnumbered_headings.sort(key=lambda h: (h.get('text') or h.get('full') or '').lower())
    dropdown_headings = numbered_headings + unnumbered_headings

    logger.debug(f"DEBUG: Building dropdown with {len(dropdown_headings)} headings")

    # Flatten for debug/stats
    all_headings = [h['full'] for h in dropdown_headings]

    external_links_by_para = {}
    docx_links_by_para = session_data.get('docx_links_by_para', {}) or {}
    if html_import:
        html_path = session_data.get('html_path', '')
        if html_path and Path(html_path).exists():
            try:
                raw_html = Path(html_path).read_text(encoding='utf-8', errors='ignore')
                manual_html, _ = extract_manual_fragment(raw_html)
                external_links_by_para = extract_external_links_from_html(manual_html)
            except Exception as e:
                logger.error(f"DEBUG: Failed to extract external links from HTML import: {e}")
    if not external_links_by_para and references:
        external_links_by_para = extract_external_links_from_reference_text(references)

    # Helper function to generate dropdown options
    def build_dropdown_options(current_target):
        options = '<option value="">-- Select NEW heading to link to --</option>\n'
        for data in dropdown_headings:
            full = data['full']
            escaped = full.replace('&', '&amp;').replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')
            selected = 'selected' if full == current_target else ''
            options += f'<option value="{escaped}" {selected}>{escaped}</option>\n'
        return options

    # Add flash messages if any
    from flask import get_flashed_messages
    messages = get_flashed_messages()
    if messages:
        review_html += '<div class="flash-messages">'
        for message in messages:
            review_html += f'<div class="flash-message success">{message}</div>'
        review_html += '</div>'

    # Show debug summary (NOW we can safely use existing_validations)
    valid_loaded = sum(1 for v in existing_validations.values() if v)
    invalid_loaded = sum(1 for v in existing_validations.values() if not v)

    review_html += f"""
    <div class="info">
    <strong>File:</strong> {session_data.get('filename', 'Unknown')}<br>
    <strong>Manual Type:</strong> {manual_type.title()}<br>
    <strong>References Found:</strong> {len(references)}<br>
    <strong>Review Page:</strong> {page} of {total_pages} ({showing_start}-{showing_end} of {total_paras} paragraphs)<br>
    <strong>Auto-Matched OLD→NEW:</strong> {len(auto_crosswalk)}<br>
    <strong>NEW Headings Available:</strong> {len(all_headings)}<br>
    <strong>Previously Saved:</strong> {len(existing_validations)} validations ({valid_loaded} valid, {invalid_loaded} invalid)<br>
    <strong>HTML Import:</strong> {'Yes' if html_import else 'No'}<br>
    {f"<strong>Linked References Detected:</strong> {linked_ref_count}<br><strong>Rebuild Links:</strong> {'Yes' if rebuild_links else 'No'}" if html_import else ""}
    </div>

    <details class="debug-panel">
    <summary>🐛 Debug Info (Click to expand)</summary>
    <p><strong>Session ID:</strong> {session_id}</p>
    <p><strong>Edit File:</strong> {edit_file.name}</p>
    <p><strong>Edit File Exists:</strong> {'Yes' if edit_file.exists() else 'No'}</p>
    <p><strong>Loaded Validations Sample:</strong></p>
    <pre>{str(dict(list(existing_validations.items())[:5]))}</pre>
    <p><strong>Loaded Edits Sample:</strong></p>
    <pre>{str(dict(list(existing_edits.items())[:5]))}</pre>
    <p><strong>Headings Sample (first 5):</strong></p>
    <pre>{str(all_headings[:5])}</pre>
    <p><em>Open browser console (F12) and click "Save Edits" to see form data being submitted.</em></p>
    </details>
    
    <h2>Internal References Found</h2>
    <p>The following references to sections/chapters were found in the document body.</p>
    <div class="info" style="background: #fef3c7; border-left-color: #f59e0b;">
    <strong>Instructions:</strong>
    <ol style="margin: 8px 0; padding-left: 20px;">
    <li><strong>Review:</strong> Check the box for VALID chapter/section references (e.g., "Chapter 1.D" → "Chapter 1.4")</li>
    <li><strong>Skip non-links:</strong> Use "Skip linking" for acronyms or citations that should not link (e.g., "D.V.M", "Ph.D", "M.S.").</li>
    <li><strong>Edit if needed:</strong> Use Link text to correct the auto-match if needed</li>
    <li><strong>Save frequently:</strong> Click "Save Edits" to persist your progress - you can come back later!</li>
    <li><strong>Filter reviewed:</strong> Use the "Hide Reviewed Items" button below to focus on unreviewed references</li>
    <li><strong>External URLs:</strong> Links detected in the paragraph are listed for verification. Use the External URL field to add or correct links to other manuals.</li>
    </ol>
    Only validated references (with checkboxes checked) will create links in the final document.
    </div>

    <div style="margin: 16px 0; padding: 12px; background: #dbeafe; border-radius: 8px; border-left: 4px solid #3b82f6;">
    <strong>Progress:</strong> <span id="reviewProgress">Loading...</span><br>
    <button type="button" id="toggleReviewed" onclick="toggleReviewedItems()" style="margin-top: 8px; background: #3b82f6; color: white; border: 0; padding: 8px 16px; border-radius: 6px; cursor: pointer;">
    Hide Reviewed Items
    </button>
    <span id="filterStatus" style="margin-left: 12px; font-style: italic; color: #1e40af;"></span>
    </div>

    <form method="POST" id="reviewForm">
    <input type="hidden" name="page" value="{page}">
    {f'''<div class="info" style="background: #eef2ff; border-left-color: #6366f1;">
    <strong>HTML Import Options:</strong><br>
    <label style="display: flex; align-items: center; gap: 8px; margin-top: 6px;">
      <input type="checkbox" name="rebuild_links" id="rebuild_links" {"checked" if rebuild_links else ""}>
      <span>Rebuild internal links (remove existing section/chapter anchors before applying new links)</span>
    </label>
    <small style="color: #4b5563; display: block; margin-top: 6px;">When unchecked, existing internal links are preserved and shown read-only for verification.</small>
    </div>''' if html_import else ""}
    <div class="references-list">
    """

    for para_idx in page_para_keys:
        refs_with_ids = ref_by_para[para_idx]
        full_text = refs_with_ids[0][0][1] if refs_with_ids else ""
        # Highlight references in the text
        highlighted_text = full_text
        for ref, ref_id in refs_with_ids:
            old_ref = ref[2]
            # Escape for HTML
            old_ref_escaped = old_ref.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            highlighted_text = highlighted_text.replace(old_ref, f'<span class="highlight">{old_ref_escaped}</span>')
        
        review_html += f"""
        <div class="reference-item">
        <strong>Paragraph {para_idx + 1}:</strong><br>
        <div class="reference-text">{highlighted_text}</div>
        """
        external_links = external_links_by_para.get(para_idx, [])
        if external_links:
            links_list = ""
            for href in external_links:
                href_safe = href.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                links_list += f"<li><a href=\"{href_safe}\" target=\"_blank\" rel=\"noopener noreferrer\">{href_safe}</a></li>"
            review_html += f"""
            <div style="margin-top: 8px; padding: 8px; background: #eef2ff; border-left: 3px solid #6366f1; border-radius: 4px;">
                <strong style="color: #3730a3;">External URLs in this paragraph:</strong>
                <ul style="margin: 6px 0 0 18px;">{links_list}</ul>
            </div>
            """
        docx_links = docx_links_by_para.get(para_idx, [])
        if docx_links:
            links_list = ""
            for item in docx_links:
                href = (item.get("href") or "").replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                text = (item.get("text") or "").replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                bad = bool(item.get("bad"))
                badge = "<span style=\"margin-left:6px;color:#b91c1c;font-weight:600;\">Bad link</span>" if bad else ""
                label = text if text else href
                if href:
                    entry = f"<a href=\"{href}\" target=\"_blank\" rel=\"noopener noreferrer\">{label}</a>{badge}"
                else:
                    entry = f"<span>{label}</span>{badge}"
                links_list += f"<li>{entry}</li>"
            review_html += f"""
            <div style="margin-top: 8px; padding: 8px; background: #fef2f2; border-left: 3px solid #ef4444; border-radius: 4px;">
                <strong style="color: #991b1b;">DOCX links in this paragraph:</strong>
                <ul style="margin: 6px 0 0 18px;">{links_list}</ul>
            </div>
            """
        review_html += """
        <div style="margin-top: 12px;">
        """
        for ref, ref_id in refs_with_ids:
            old_ref = ref[2]

            # NEW WORKFLOW: Get auto-matched NEW reference
            auto_matched_new = auto_crosswalk.get(old_ref, "")
            # If auto-matched, find the full heading text from new_headings
            auto_target_full = ""
            auto_match_found = bool(auto_matched_new and auto_matched_new in new_headings)
            if auto_match_found:
                full_new_heading = new_headings[auto_matched_new]['full']
                auto_target_full = full_new_heading
            else:
                full_new_heading = "Not auto-matched"

            # Check if we have an existing manual edit (overrides auto-match)
            # ref_id is already a stable ID (e.g., "ref_42_123_a1b2c3d4"), use it directly
            edit_key = ref_id
            selected_target = existing_link_targets.get(edit_key, "")
            current_value = existing_edits.get(edit_key, "")
            if mapping_mode == "keep_old" and current_value and old_ref:
                target_label = (selected_target.split(' - ')[0] or '').strip() if selected_target else ''
                if not target_label and auto_target_full:
                    target_label = (auto_target_full.split(' - ')[0] or '').strip()
                if target_label:
                    current_norm = normalize_heading_ref(current_value)
                    target_norm = normalize_heading_ref(target_label)
                    old_norm = normalize_heading_ref(old_ref)
                    if current_norm and target_norm and old_norm and current_norm == target_norm and current_norm != old_norm:
                        current_value = old_ref
            if not current_value:
                if mapping_mode == "keep_old" and old_ref:
                    current_value = old_ref
                else:
                    if selected_target:
                        current_value = (selected_target.split(' - ')[0] or '').strip()
                    elif auto_target_full:
                        selected_target = auto_target_full
                        current_value = (auto_target_full.split(' - ')[0] or '').strip()
            external_value = existing_external_urls.get(edit_key, "")
            if not external_value:
                external_candidates = external_links_by_para.get(para_idx, [])
                if external_candidates:
                    external_value = external_candidates[0]
            if not external_value:
                docx_links = docx_links_by_para.get(para_idx, [])
                for item in docx_links:
                    href = item.get("href") or ""
                    if href and not item.get("bad"):
                        external_value = href
                        break

            # Check validation status (default to false - user must explicitly validate)
            has_saved_validation = edit_key in existing_validations
            is_valid = existing_validations.get(edit_key, False)
            is_ignored = bool(existing_ignored.get(edit_key, False))
            if is_ignored:
                is_valid = False
            checked = "checked" if is_valid else ""
            validation_class = "" if is_valid else "invalid"
            is_linked = bool(html_import and len(ref) > 5 and ref[5])
            is_read_only = bool(is_linked and not rebuild_links)
            if not has_saved_validation and auto_matched_new and auto_match_found:
                norm_old = normalize_heading_ref(old_ref)
                if norm_old:
                    norm_old = re.sub(r'^[\(\[]\s*', '', norm_old)
                    norm_old = re.sub(r'\s*[\)\]]$', '', norm_old)
                if mapping_mode == "keep_old":
                    norm_new = normalize_heading_ref(auto_matched_new)
                    if norm_new:
                        norm_new = re.sub(r'^[\(\[]\s*', '', norm_new)
                        norm_new = re.sub(r'\s*[\)\]]$', '', norm_new)
                    norm_old_no_prefix = re.sub(r'^(Chapter|Section)\s+', '', norm_old or '', flags=re.IGNORECASE)
                    norm_new_no_prefix = re.sub(r'^(Chapter|Section)\s+', '', norm_new or '', flags=re.IGNORECASE)
                    if norm_old and (norm_old == norm_new or norm_old_no_prefix == norm_new_no_prefix):
                        is_valid = True
                        checked = "checked"
                        validation_class = ""
                else:
                    is_valid = True
                    checked = "checked"
                    validation_class = ""
            if is_read_only:
                checked = ""
                validation_class = "invalid"

            # Determine if this item has been reviewed (exists in saved data)
            has_been_reviewed = (
                edit_key in existing_validations
                or edit_key in existing_edits
                or edit_key in existing_link_targets
                or edit_key in existing_ignored
                or edit_key in existing_external_urls
            )
            if is_read_only:
                has_been_reviewed = True
            review_status = "reviewed" if has_been_reviewed else "unreviewed"
            review_indicator = "✓ Reviewed" if has_been_reviewed else "○ Not reviewed yet"
            review_badge_color = "#10b981" if has_been_reviewed else "#9ca3af"

            # Debug: Print what we're rendering
            logger.debug(f"DEBUG: Rendering ref_{ref_id}: OLD='{old_ref}', AUTO-NEW='{auto_matched_new}', is_valid={is_valid}, reviewed={has_been_reviewed}, current_value='{current_value[:30] if current_value and len(current_value) > 30 else current_value}'")

            # Escape for HTML input value
            current_value_escaped = current_value.replace('&', '&amp;').replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')
            auto_target_escaped = auto_target_full.replace('&', '&amp;').replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')
            old_ref_escaped = old_ref.replace('&', '&amp;').replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')

            review_html += f"""
            <div style="margin: 8px 0; padding: 8px; background: #f9f9f9; border-radius: 4px;"
                 class="ref-item-{ref_id} reference-review-item{(' ref-collapsed' if is_ignored else '')}{(' ref-readonly' if is_read_only else '')}"
                 data-ref-id="{ref_id}"
                 data-auto-target="{auto_target_escaped}"
                 data-old-ref="{old_ref_escaped}"
                 data-is-ignored="{str(is_ignored).lower()}"
                 data-review-status="{review_status}"
                 data-is-valid="{str(is_valid and not is_ignored).lower()}">
                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px;">
                    <span style="font-size: 12px; color: #6b7280;">Reference #{ref_id}</span>
                    <span style="display: inline-flex; align-items: center; gap: 6px;">
                        <span class="skip-badge" id="skip_badge_{ref_id}">Skipped</span>
                        <span style="font-size: 12px; font-weight: 600; color: {review_badge_color};">{review_indicator}</span>
                    </span>
                </div>
                <div class="ref-summary" style="margin-bottom: 8px; padding: 8px; background: #fff; border-left: 3px solid #f59e0b; border-radius: 4px;">
                    <div>
                        <strong style="color: #92400e;">Old reference (from text):</strong><br>
                        <code style="background: #fef3c7; padding: 2px 6px; border-radius: 3px;">{old_ref}</code>
                    </div>
                    {f"<div class='small' style='margin-top:6px;color:#475569;'>Linked in HTML (read-only unless Rebuild is enabled).</div>" if is_read_only else ""}
                    {"" if auto_match_found else "<div class='small' style='margin-top:6px;color:#b45309;'>No internal match found. Add an external URL or skip linking.</div>"}
                </div>

                <div class="validation-checkbox" style="margin-top: 6px; background: #eef2ff; border-color: #c7d2fe;">
                    <label>
                        <input type="checkbox"
                               name="ref_ignore_{ref_id}"
                               id="ref_ignore_{ref_id}"
                               value="1"
                               {"checked" if is_ignored else ""} {"disabled" if is_read_only else ""}>
                        <span>Skip linking (not a real reference)</span>
                    </label>
                </div>
                <div class="ref-details">
                    <div style="margin-bottom: 8px; padding: 8px; background: #fff; border-left: 3px solid #10b981; border-radius: 4px;">
                        <strong style="color: #065f46;">Auto-matched new:</strong><br>
                        <code style="background: #d1fae5; padding: 2px 6px; border-radius: 3px;">{full_new_heading}</code>
                    </div>
                    <div class="validation-checkbox {validation_class}">
                        <label>
                            <input type="checkbox"
                                   name="ref_valid_{ref_id}"
                                   id="ref_valid_{ref_id}"
                                   value="1"
                                   {checked} {"disabled" if is_read_only else ""}>
                            <span>This is a valid chapter/section reference</span>
                        </label>
                    </div>
                <div style="margin-bottom: 8px;">
                    <label for="ref_target_{ref_id}" style="display: block; margin-bottom: 4px; font-weight: 600;">Link target (select new heading to link to):</label>
                    <select id="ref_target_{ref_id}"
                            name="ref_target_{ref_id}"
                            {"disabled" if is_read_only else ""}
                            style="width: 100%; padding: 6px; border: 1px solid #ddd; border-radius: 4px; font-family: monospace; font-size: 13px;">
                        {build_dropdown_options(selected_target)}
                    </select>
                </div>

                <div>
                    <label for="ref_edit_{ref_id}" style="display: block; margin-bottom: 4px; font-weight: 600;">Link text (optional):</label>
                    <input type="text"
                           id="ref_edit_{ref_id}"
                           name="ref_edit_{ref_id}"
                           value="{current_value_escaped}"
                           placeholder="Leave blank to use the selected link text, or enter custom link text"
                           style="width: 100%; padding: 6px; border: 1px solid #ddd; border-radius: 4px; font-family: monospace;"
                           {"readonly" if is_read_only else ""}>
                    <small style="color: #666; display: block; margin-top: 4px;">If left blank, the selected link text (or the auto-match) will be used.</small>
                </div>
                <div style="margin-top: 10px;">
                    <label for="ref_external_{ref_id}" style="display: block; margin-bottom: 4px; font-weight: 600;">External URL (optional):</label>
                    <input type="text"
                           id="ref_external_{ref_id}"
                           name="ref_external_{ref_id}"
                           value="{external_value.replace('&', '&amp;').replace('\"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')}"
                           placeholder="https://example.com/manual"
                           style="width: 100%; padding: 6px; border: 1px solid #ddd; border-radius: 4px; font-family: monospace;"
                           {"readonly" if is_read_only else ""}>
                    <small style="color: #666; display: block; margin-top: 4px;">If set, this external link is used instead of the internal link target.</small>
                </div>
                <script>
                document.getElementById('ref_valid_{ref_id}').addEventListener('change', function() {{
                    const parent = document.querySelector('.ref-item-{ref_id}').parentElement;
                    const vbox = this.closest('.validation-checkbox');
                    if (this.checked) {{
                        parent.classList.remove('invalid');
                        vbox.classList.remove('invalid');
                    }} else {{
                        parent.classList.add('invalid');
                        vbox.classList.add('invalid');
                    }}
                }});
                (function() {{
                    const ignoreBox = document.getElementById('ref_ignore_{ref_id}');
                    const validBox = document.getElementById('ref_valid_{ref_id}');
                    const targetSelect = document.getElementById('ref_target_{ref_id}');
                    const editInput = document.getElementById('ref_edit_{ref_id}');
                    const externalInput = document.getElementById('ref_external_{ref_id}');
                    const container = document.querySelector('.ref-item-{ref_id}');
                    if (ignoreBox && ignoreBox.hasAttribute('disabled')) {{
                        return;
                    }}
                    function targetLabel() {{
                        if (!targetSelect) return '';
                        if (!targetSelect.value) return '';
                        const option = targetSelect.options[targetSelect.selectedIndex];
                        const text = option ? option.text : targetSelect.value;
                        return (text.split(' - ')[0] || '').trim();
                    }}
                    function syncTargetToOverride(force) {{
                        if (!targetSelect || !editInput) return;
                        if (ignoreBox && ignoreBox.checked) return;
                        const selected = targetSelect.value || '';
                        if (!force && editInput.value.trim()) {{
                            return;
                        }}
                        if (!selected) {{
                            editInput.value = '';
                            return;
                        }}
                        editInput.value = targetLabel();
                    }}
                    function syncIgnoreState() {{
                        if (container) {{
                            container.classList.toggle('ref-collapsed', ignoreBox.checked);
                        }}
                        if (ignoreBox.checked) {{
                            validBox.checked = false;
                            targetSelect.style.pointerEvents = 'none';
                            targetSelect.style.opacity = '0.6';
                            editInput.readOnly = true;
                            editInput.style.opacity = '0.6';
                            externalInput.readOnly = true;
                            externalInput.style.opacity = '0.6';
                            targetSelect.setAttribute('aria-disabled', 'true');
                            editInput.setAttribute('aria-disabled', 'true');
                            externalInput.setAttribute('aria-disabled', 'true');
                        }} else {{
                            targetSelect.style.pointerEvents = '';
                            targetSelect.style.opacity = '';
                            editInput.readOnly = false;
                            editInput.style.opacity = '';
                            externalInput.readOnly = false;
                            externalInput.style.opacity = '';
                            targetSelect.removeAttribute('aria-disabled');
                            editInput.removeAttribute('aria-disabled');
                            externalInput.removeAttribute('aria-disabled');
                        }}
                    }}
                    function applyAutoMatch() {{
                        if (!validBox.checked) return;
                        if (targetSelect && !targetSelect.value) {{
                            const autoTarget = container ? container.getAttribute('data-auto-target') : '';
                            if (autoTarget) {{
                                targetSelect.value = autoTarget;
                            }}
                        }}
                        if (editInput && !editInput.value.trim()) {{
                            editInput.value = targetLabel();
                        }}
                    }}
                    ignoreBox.addEventListener('change', syncIgnoreState);
                    validBox.addEventListener('change', applyAutoMatch);
                    targetSelect.addEventListener('change', function() {{ syncTargetToOverride(true); }});
                    syncIgnoreState();
                    syncTargetToOverride(false);
                    applyAutoMatch();
                }})();
                </script>
                </div>
            </div>
            """
        review_html += "</div></div>"
    
    if total_pages > 1:
        review_html += f"<p><em>Page {page} of {total_pages}. Use Previous/Next to move between pages.</em></p>"
    
    review_html += """
    </div>

    """

    prev_disabled = "disabled" if page <= 1 else ""
    next_disabled = "disabled" if page >= total_pages else ""

    review_html += f"""
    <div class="actions" style="display: flex; gap: 12px; align-items: center;">
    <button type="submit" name="save_edits" style="background: #0284c7;" onclick="debugForm(event, 'save_edits')">Save Edits</button>
    <button type="submit" name="prev_page" style="background: #64748b;" {prev_disabled}>Save & Previous</button>
    <button type="submit" name="next_page" style="background: #64748b;" {next_disabled}>Save & Next</button>
    <button type="submit" name="proceed" style="background: #981E32;" onclick="debugForm(event, 'proceed')">Proceed with Conversion</button>
    </div>
    </form>
    """

    review_html += """
    <script>
    function debugForm(event, action) {
        const form = document.getElementById('reviewForm');
        const formData = new FormData(form);
        console.log('=== FORM SUBMISSION DEBUG ===');
        console.log('Action:', action);

        let checkedCount = 0;
        let uncheckedCount = 0;
        let editCount = 0;
        let targetCount = 0;

        for (let [key, value] of formData.entries()) {
            if (key.startsWith('ref_valid_')) {
                checkedCount++;
                console.log('✓ Checked:', key, '=', value);
            } else if (key.startsWith('ref_edit_')) {
                if (value.trim()) {
                    editCount++;
                    console.log('✏️ Edit:', key, '=', value.substring(0, 30));
                }
            } else if (key.startsWith('ref_target_')) {
                if (value.trim()) {
                    targetCount++;
                    console.log('🔗 Target:', key, '=', value.substring(0, 30));
                }
            }
        }

        // Count all checkboxes to find unchecked ones
        const allCheckboxes = document.querySelectorAll('input[type="checkbox"][name^="ref_valid_"]');
        uncheckedCount = allCheckboxes.length - checkedCount;

        console.log('Summary: ' + checkedCount + ' checked, ' + uncheckedCount + ' unchecked, ' + editCount + ' edits, ' + targetCount + ' targets');
        console.log('=== END DEBUG ===');

        // Don't prevent submission
        return true;
    }
    </script>
    <script>
    // Debug: Show checkbox status on page load
    window.addEventListener('DOMContentLoaded', function() {
        console.log('=== PAGE LOAD DEBUG ===');
        const allCheckboxes = document.querySelectorAll('input[type="checkbox"][name^="ref_valid_"]');
        const checkedCheckboxes = document.querySelectorAll('input[type="checkbox"][name^="ref_valid_"]:checked');
        console.log('Total checkboxes on page:', allCheckboxes.length);
        console.log('Pre-checked checkboxes:', checkedCheckboxes.length);

        checkedCheckboxes.forEach(cb => {
            console.log('  - ' + cb.name + ' is checked');
        });

        if (allCheckboxes.length === 0) {
            console.warn('WARNING: No checkboxes found on page! References may not be rendering.');
        }

        // Add change listeners to all checkboxes for debugging
        allCheckboxes.forEach(cb => {
            cb.addEventListener('change', function(e) {
                console.log('Checkbox changed:', e.target.name, 'checked:', e.target.checked);
            });
        });

        console.log('=== END PAGE LOAD DEBUG ===');

        // Update progress counter on page load
        updateProgressCounter();
    });

    // Toggle hiding reviewed items (only those marked valid/checked)
    let showingReviewedItems = true;
    function toggleReviewedItems() {
        showingReviewedItems = !showingReviewedItems;
        const reviewedItems = document.querySelectorAll('.reference-review-item[data-is-valid="true"]');
        const button = document.getElementById('toggleReviewed');
        const statusSpan = document.getElementById('filterStatus');

        reviewedItems.forEach(item => {
            item.style.display = showingReviewedItems ? '' : 'none';
        });

        // Hide entire paragraph blocks if all their valid refs are hidden
        const paragraphBlocks = document.querySelectorAll('.reference-item');
        paragraphBlocks.forEach(block => {
            const targetRefs = Array.from(block.querySelectorAll('.reference-review-item[data-is-valid="true"]'));
            if (targetRefs.length === 0) {
                block.style.display = ''; // nothing to hide in this block
                return;
            }
            const anyVisible = targetRefs.some(r => r.style.display !== 'none');
            block.style.display = anyVisible || showingReviewedItems ? '' : 'none';
        });

        if (showingReviewedItems) {
            button.textContent = 'Hide Reviewed Items';
            statusSpan.textContent = '';
        } else {
            button.textContent = 'Show All Items';
            statusSpan.textContent = '(Showing only unreviewed items)';
        }

        updateProgressCounter();
    }

    // Update progress counter
    function updateProgressCounter() {
        const allItems = document.querySelectorAll('.reference-review-item');
        const reviewedItems = document.querySelectorAll('.reference-review-item[data-review-status="reviewed"]');
        const validItems = document.querySelectorAll('.reference-review-item[data-is-valid="true"]');
        const skippedItems = document.querySelectorAll('.reference-review-item[data-is-ignored="true"]');
        const unreviewedItems = document.querySelectorAll('.reference-review-item[data-review-status="unreviewed"]');
        const autoMatchedItems = Array.from(allItems).filter(item => {
            const target = item.getAttribute('data-auto-target') || '';
            return target.trim().length > 0;
        });

        const total = allItems.length;
        const reviewed = reviewedItems.length;
        const valid = validItems.length;
        const skipped = skippedItems.length;
        const autoMatched = autoMatchedItems.length;
        const unreviewed = unreviewedItems.length;

        const progressSpan = document.getElementById('reviewProgress');
        progressSpan.innerHTML = `
            <strong>${reviewed} of ${total} reviewed</strong>
            (${valid} valid references, ${skipped} skipped, ${autoMatched} auto-matched, ${unreviewed} not yet reviewed)
        `;

        console.log('Progress:', {total, reviewed, valid, skipped, autoMatched, unreviewed});
    }
    </script>
    <p style="margin-top: 12px; color: #666; font-size: 14px;">
    <strong>Note:</strong> Click "Save Edits" frequently to save your progress. Your corrections persist and will be applied automatically on future conversions of this document.
    </p>
    </div>
    </div>
    </body>
    </html>
    """
    
    return review_html

@app.route("/table_review/<session_id>", methods=["GET", "POST"])
def table_review(session_id):
    """Review table formatting options before conversion."""
    session = SessionDir(session_id)
    session_file = session.session_json
    if not session_file.exists():
        flash("Session expired or invalid.")
        return redirect(url_for("index"))

    session_data = json.loads(session_file.read_text(encoding='utf-8'))
    html_import = session_data.get('html_import', False)
    src_path = Path(session_data.get('src_path', ''))
    html_path = Path(session_data.get('html_path', ''))
    has_tables = False
    if html_import and html_path.exists():
        has_tables = has_tables_in_html(html_path)
    elif src_path.exists():
        has_tables = has_tables_in_docx(src_path)
    if not has_tables:
        return redirect(url_for("do_convert", session_id=session_id))

    manual_type = session_data.get('manual_type', 'chapter')
    theme_settings, warnings = coerce_theme_settings(session_data.get('theme_settings'), manual_type)

    if request.method == "POST":
        updates = session_data.get('theme_settings', {}).copy()
        updates.update(request.form.to_dict())
        theme_settings, warnings = coerce_theme_settings(updates, manual_type)
        session_data['theme_settings'] = theme_settings
        session_file.write_text(json.dumps(session_data, indent=2, default=str), encoding='utf-8')
        if 'back' in request.form:
            return redirect(url_for('review', session_id=session_id))
        return redirect(url_for("do_convert", session_id=session_id))

    review_html = f"""
    <!doctype html>
    <html lang="en"><head>
    <meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
    <title>Table Review</title>
    <style>
    body {{ font-family: system-ui, -apple-system, sans-serif; margin: 0; padding: 20px; background: #f7f7f8; }}
    .container {{ max-width: 900px; margin: 0 auto; background: white; padding: 24px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
    h1 {{ color: #981E32; margin-top: 0; }}
    label {{ font-weight: 600; display: block; margin: 12px 0 6px; }}
    select, input[type="number"], input[type="color"] {{ padding: 6px 8px; border: 1px solid #ddd; border-radius: 6px; }}
    .row {{ display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }}
    .checkbox {{ display: flex; align-items: center; gap: 8px; margin: 8px 0; }}
    .actions {{ margin-top: 16px; display: flex; gap: 12px; align-items: center; }}
    .btn {{ background: #981E32; color: #fff; border: 0; padding: 10px 16px; border-radius: 8px; cursor: pointer; text-decoration: none; }}
    .btn.secondary {{ background: #4b5563; }}
    .small {{ color: #555; font-size: 14px; }}
    </style>
    </head>
    <body>
    <div class="container">
      <h1>Table Review</h1>
      <p class="small">Choose common table formatting rules before conversion.</p>
      <form method="POST">
        <div class="row">
          <div>
            <label for="table_col1_align">Column 1 alignment</label>
            <select id="table_col1_align" name="table_col1_align">
              <option value="left" {"selected" if theme_settings.get("table_col1_align") == "left" else ""}>Left</option>
              <option value="center" {"selected" if theme_settings.get("table_col1_align") == "center" else ""}>Center</option>
              <option value="right" {"selected" if theme_settings.get("table_col1_align") == "right" else ""}>Right</option>
            </select>
          </div>
          <div>
            <label for="table_coln_align">Columns 2+ alignment</label>
            <select id="table_coln_align" name="table_coln_align">
              <option value="left" {"selected" if theme_settings.get("table_coln_align") == "left" else ""}>Left</option>
              <option value="center" {"selected" if theme_settings.get("table_coln_align") == "center" else ""}>Center</option>
              <option value="right" {"selected" if theme_settings.get("table_coln_align") == "right" else ""}>Right</option>
            </select>
          </div>
          <div>
            <label for="table_header_align">Header row alignment</label>
            <select id="table_header_align" name="table_header_align">
              <option value="left" {"selected" if theme_settings.get("table_header_align") == "left" else ""}>Left</option>
              <option value="center" {"selected" if theme_settings.get("table_header_align") == "center" else ""}>Center</option>
              <option value="right" {"selected" if theme_settings.get("table_header_align") == "right" else ""}>Right</option>
            </select>
          </div>
          <div>
            <label for="table_layout_mode">Column width mode</label>
            <select id="table_layout_mode" name="table_layout_mode">
              <option value="auto" {"selected" if theme_settings.get("table_layout_mode") == "auto" else ""}>Auto</option>
              <option value="fixed" {"selected" if theme_settings.get("table_layout_mode") == "fixed" else ""}>Fixed</option>
            </select>
          </div>
        </div>
        <div class="row" style="margin-top:8px;">
          <div>
            <label for="table_header_bg">Header background</label>
            <input id="table_header_bg" name="table_header_bg" type="color" value="{theme_settings.get("table_header_bg")}">
          </div>
          <div>
            <label for="table_header_color">Header text color</label>
            <input id="table_header_color" name="table_header_color" type="color" value="{theme_settings.get("table_header_color")}">
          </div>
          <div class="checkbox" style="align-items:flex-end;">
            <input id="table_header_bold" name="table_header_bold" type="checkbox" {"checked" if theme_settings.get("table_header_bold") else ""}>
            <label for="table_header_bold">Header bold</label>
          </div>
        </div>
        <div class="row" style="margin-top:8px;">
          <div>
            <label for="table_border_color">Border color</label>
            <input id="table_border_color" name="table_border_color" type="color" value="{theme_settings.get("table_border_color")}">
          </div>
          <div>
            <label for="table_border_width">Border width (px)</label>
            <input id="table_border_width" name="table_border_width" type="number" step="0.5" value="{theme_settings.get("table_border_width")}">
          </div>
          <div>
            <label for="table_border_style">Border style</label>
            <select id="table_border_style" name="table_border_style">
              <option value="solid" {"selected" if theme_settings.get("table_border_style") == "solid" else ""}>Solid</option>
              <option value="dashed" {"selected" if theme_settings.get("table_border_style") == "dashed" else ""}>Dashed</option>
              <option value="dotted" {"selected" if theme_settings.get("table_border_style") == "dotted" else ""}>Dotted</option>
              <option value="none" {"selected" if theme_settings.get("table_border_style") == "none" else ""}>None</option>
            </select>
          </div>
          <div>
            <label for="table_cell_padding">Cell padding (px)</label>
            <input id="table_cell_padding" name="table_cell_padding" type="number" step="1" value="{theme_settings.get("table_cell_padding")}">
          </div>
          <div>
            <label for="table_row_stripe_color">Stripe color</label>
            <input id="table_row_stripe_color" name="table_row_stripe_color" type="color" value="{theme_settings.get("table_row_stripe_color")}">
          </div>
          <div class="checkbox" style="align-items:flex-end;">
            <input id="table_row_stripe" name="table_row_stripe" type="checkbox" {"checked" if theme_settings.get("table_row_stripe") else ""}>
            <label for="table_row_stripe">Row striping</label>
          </div>
        </div>
        <div class="actions">
          <button type="submit" class="btn">Continue to conversion</button>
          <button type="submit" name="back" value="1" class="btn secondary">Back to references</button>
        </div>
      </form>
    </div>
    </body>
    </html>
    """
    return review_html

@app.route("/convert/<session_id>", methods=["GET"])
def do_convert(session_id):
    """Perform the actual conversion after review"""
    import json
    
    session = SessionDir(session_id)
    session_file = session.session_json
    if not session_file.exists():
        flash("Session expired or invalid.")
        return redirect(url_for("index"))
    
    session_data = json.loads(session_file.read_text(encoding='utf-8'))
    html_import = session_data.get('html_import', False)
    src_path = Path(session_data.get('src_path', '')) if session_data.get('src_path') else session.source_docx
    pre_path = Path(session_data.get('pre_path', '')) if session_data.get('pre_path') else session.pre_docx
    filename = session_data.get('filename', src_path.name)
    toc_depth = session_data.get('toc_depth', 2)
    preserve = session_data.get('preserve_numbers', False)
    manual_type = session_data.get('manual_type', 'chapter')
    mapping_mode = session_data.get('mapping_mode', 'map_new')
    numbering_mode = session_data.get('numbering_mode', "preserve" if preserve else "css-counters")
    rebuild_links = session_data.get('rebuild_links', False)
    strip_docx_formatting = session_data.get('strip_docx_formatting', False)
    infer_heading_depth = session_data.get('infer_heading_depth', False)
    infer_style_map = session_data.get('infer_style_map') or {}
    infer_sequence_map = deserialize_sequence_map(session_data.get('infer_sequence_map'))
    stable_heading_map = session_data.get('stable_heading_map') or {}
    heading_edits = session_data.get('heading_edits', {})
    theme_settings, theme_warnings = coerce_theme_settings(session_data.get('theme_settings'), manual_type)
    theme_id = theme_settings.get("theme_id", "manual")

    # Align preserve flag with mapping choice: keeping old headings implies keeping numbers
    # BUT: Respect the user's checkbox - only force preserve=True for keep_old, don't override checkbox for map_new
    if mapping_mode == "keep_old":
        preserve = True

    logger.debug(f"DEBUG [do_convert]: mapping_mode={mapping_mode}, preserve={preserve} (from checkbox: {session_data.get('preserve_numbers', False)})")

    # NEW WORKFLOW: Load auto_crosswalk and new_headings
    auto_crosswalk = session_data.get('approved_crosswalk') or session_data.get('auto_crosswalk', {})
    new_headings = session_data.get('new_headings', {})
    original_references = session_data.get('references', [])

    # Load saved reference edits, validations, and link targets if they exist
    edit_file = session.edits_json
    reference_edits = {}
    reference_validations = {}
    reference_link_targets = {}
    reference_ignored = {}
    reference_external_urls = {}
    if edit_file.exists():
        try:
            edit_data = json.loads(edit_file.read_text(encoding='utf-8'))
            reference_edits = edit_data.get('reference_edits', {})
            reference_validations = edit_data.get('reference_validations', {})
            reference_link_targets = edit_data.get('reference_link_targets', {})
            reference_ignored = edit_data.get('reference_ignored', {})
            reference_external_urls = edit_data.get('reference_external_urls', {})
            approved_from_edits = edit_data.get('approved_crosswalk')
            if approved_from_edits:
                auto_crosswalk = approved_from_edits
            logger.debug("="*80)
            logger.debug("DEBUG: Loaded reference edits from file")
            logger.debug(f"DEBUG: Edit file: {edit_file}")
            logger.debug(f"DEBUG: Loaded {len(reference_edits)} edits, {len(reference_validations)} validations, {len(reference_link_targets)} link targets")
            logger.debug(f"DEBUG: Ignored references: {len(reference_ignored)}")
            logger.debug(f"DEBUG: External URL references: {len(reference_external_urls)}")
            valid_count = sum(1 for v in reference_validations.values() if v)
            logger.debug(f"DEBUG: Valid references: {valid_count}, Invalid/False positives: {len(reference_validations) - valid_count}")
            if reference_validations:
                logger.debug(f"DEBUG: Sample validations (first 5): {dict(list(reference_validations.items())[:5])}")
        except Exception as e:
            logger.error(f"ERROR: Could not load edit file: {e}")
            import traceback
            traceback.print_exc()
    else:
        logger.debug(f"DEBUG: No edit file found at: {edit_file}")
    
    # Now do the conversion
    try:
        # 1. Source Acquisition
        if not html_import:
            out = session.export_html
            run_pandoc(pre_path, out)
            html_content = out.read_text(encoding="utf-8", errors="ignore")
        else:
            html_content = (session.root / "import.html").read_text(encoding="utf-8", errors="ignore")

        # 2. Unified Processing Pipeline (Single-Pass BeautifulSoup)
        pipeline_config = {
            'toc_depth': toc_depth,
            'mapping_mode': mapping_mode,
            'preserve_numbers': preserve,
            'infer_heading_depth': infer_heading_depth,
            'infer_style_map': infer_style_map,
            'stable_heading_map': stable_heading_map,
            'heading_edits': heading_edits,
            'table_align_mode': theme_settings.get('table_align_mode', 'auto'),
            'table_col1_align': theme_settings.get('table_col1_align'),
            'table_coln_align': theme_settings.get('table_coln_align'),
            'table_header_align': theme_settings.get('table_header_align'),
            'references': original_references,
            'reference_edits': reference_edits,
            'reference_validations': reference_validations,
            'reference_link_targets': reference_link_targets,
            'reference_ignored': reference_ignored,
            'reference_external_urls': reference_external_urls,
            'auto_crosswalk': auto_crosswalk,
            'new_headings': new_headings,
            'skip_linked_text': (html_import and not rebuild_links),
            'rebuild_links': bool(html_import and rebuild_links)
        }
        
        final_html, toc_html = process_html_pipeline(html_content, session_id, pipeline_config)
        
        # Save processed body
        out = session.export_html
        out.write_text(final_html, encoding='utf-8')
        normalized = final_html
        
        # 3. Build Preview Wrappers
        # Set numbering mode: "preserve" keeps original numbers, "css-counters" applies CSS auto-numbering
        numbering_mode = session_data.get('numbering_mode') if html_import else ("preserve" if preserve else "css-counters")
        
        # Use our server-side generated TOC instead of an empty placeholder
        theme_attr = f' data-theme="{sanitize_theme_id(theme_id, "manual")}"'
        manual_grid_block = (
            '<!-- ACCESSIBILITY: Skip navigation link for keyboard users (WCAG 2.4.1) -->\n'
            '<a href="#main-content" class="skip-to-main">Skip to main content</a>\n'
            f'<div class="manual-grid" data-toc-depth="{toc_depth}" data-manual-type="{manual_type}" data-numbering-mode="{numbering_mode}"{theme_attr}>\n'
            '  <nav class="manual-toc" role="navigation" aria-label="Table of Contents">\n'
            '    <h2 id="toc-heading">Table of Contents</h2>\n'
            '    <div class="manual-search">\n'
            '      <input type="text" class="manual-search-input" placeholder="Search headings and content..." aria-label="Search table of contents" aria-describedby="search-help" role="searchbox">\n'
            '      <button type="button" class="manual-search-clear" aria-label="Clear search">X</button>\n'
            '    </div>\n'
            '    <span id="search-help" class="sr-only">Type to filter headings and content</span>\n'
            f'    {toc_html}\n'
            '  </nav>\n'
            f'  <main class="manual" id="main-content" role="main" tabindex="-1">\n'
            f'    {final_html}\n'
            '  </main>\n'
            '</div>\n'
        )

        # 4. Export Artifacts
        token = str(uuid.uuid4())
        fragment_path = session.root / f"{token}_fragment.html"
        standalone_path = session.root / f"{token}_standalone.html"
        docx_html_path = session.root / f"{token}_docx_source.html"
        docx_path = session.root / f"{token}_{src_path.stem}_numbered.docx"
        css_path = session.root / f"{token}_wordpress.css"
        manual_content_path = session.root / f"{token}_manual.html"
        manual_content_raw_path = session.root / f"{token}_manual_raw.html"
        
        # Save manual content
        manual_content_path.write_text(final_html, encoding='utf-8')
        manual_content_raw_path.write_text(final_html, encoding='utf-8')

        # Build combined CSS
        wp_css = get_wp_css_text()
        wp_js = get_wp_js_text()
        theme_css = build_theme_css(theme_settings)
        combined_css = f"{wp_css}\n{theme_css}"
        css_path.write_text(combined_css, encoding='utf-8')

        # Standalone HTML
        standalone_html = f'<!doctype html><html lang="en"><head><meta charset="utf-8"><title>{filename}</title><style>{combined_css}</style></head><body>{manual_grid_block}<script>{wp_js}</script></body></html>'
        standalone_path.write_text(standalone_html, encoding='utf-8')

        # DOCX Generation
        try:
            # Re-generate numbered HTML for DOCX specifically (bakes in numbers)
            docx_source_html = final_html
            if strip_docx_formatting:
                docx_source_html = strip_inline_formatting(docx_source_html)
            
            # Note: apply_css_counter_numbering is still in the file for now
            numbered_html = apply_css_counter_numbering(docx_source_html, manual_type, preserve=preserve)
            numbered_html = sanitize_docx_ids_for_export(numbered_html)
            docx_html_path.write_text(f'<!doctype html><html><body>{numbered_html}</body></html>', encoding='utf-8')
            
            subprocess.run(["pandoc", str(docx_html_path), "-f", "html", "-t", "docx", "-o", str(docx_path)], check=True)
            fix_numbering_xml(docx_path)
            sanitize_docx_styles(docx_path)
            relocate_body_level_bookmarks(docx_path)
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")

        # Metadata for download
        meta = {
            "session_id": session_id,
            "filename": filename,
            "manual_type": manual_type,
            "toc_depth": toc_depth,
            "numbering_mode": numbering_mode,
            "theme_settings": theme_settings,
            "theme_id": theme_id,
            "docx_path": str(docx_path),
            "docx_html_path": str(docx_html_path),
            "manual_content_path": str(manual_content_path),
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M")
        }
        (session.root / f"{token}_meta.json").write_text(json.dumps(meta, indent=2), encoding='utf-8')

        session_data.update({
            'token': token,
            'manual_content_path': str(manual_content_path),
            'manual_content_raw_path': str(manual_content_raw_path),
            'css_path': str(css_path),
            'numbering_mode': numbering_mode,
            'theme_settings': theme_settings
        })
        session_file.write_text(json.dumps(session_data, indent=2, default=str), encoding='utf-8')

        # Persist the stable map for future revisions
        save_stable_heading_map(session_id, final_html)

        return render_template_string(HOME_PAGE, 
                                    show_preview=True, 
                                    hide_upload=True, 
                                    body_html=BeautifulSoup(final_html, 'html.parser'), 
                                    token=token, 
                                    session_id=session_id,
                                    page_title=filename.replace('.docx', '').replace('_', ' ').title(),
                                    manual_type=manual_type,
                                    toc_depth=toc_depth,
                                    numbering_mode=numbering_mode,
                                    wordpress_css_tag=f"<style>{combined_css}</style>",
                                    wordpress_js_tag=f"<script>{wp_js}</script>",
                                    theme_settings=theme_settings,
                                    theme_id=theme_id)

    except Exception as e:
        logger.error(f"Conversion failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        flash(f"Conversion failed: {e}")
        return redirect(url_for("index"))

@app.route("/download/<token>/<kind>", methods=["GET"])
def download(token, kind):
    # Retrieve metadata to find the session_id
    # Tokens are stored in the root of the session they belong to
    # We have to find which session directory contains this token_meta.json
    session_id = None
    meta = {}
    for session_path in PERSIST_DIR.iterdir():
        if not session_path.is_dir(): continue
        meta_file = session_path / f"{token}_meta.json"
        if meta_file.exists():
            session_id = session_path.name
            try:
                meta = json.loads(meta_file.read_text(encoding='utf-8'))
            except Exception:
                pass
            break
    
    if not session_id:
        flash("Download not found or expired.")
        return redirect(url_for("index"))

    session = SessionDir(session_id)
    
    if kind == "css":
        css_file = session.root / f"{token}_wordpress.css"
        if css_file.exists():
            return send_file(str(css_file), as_attachment=True, download_name="wordpress.css")
    
    if kind == "js":
        js_path = Path(__file__).parent / "wordpress.js"
        if js_path.exists():
            return send_file(str(js_path), as_attachment=True, download_name="wordpress.js")
            
    if kind in ("fragment", "fragment_css", "standalone"):
        manual_content_path = Path(meta.get("manual_content_path", "") or "")
        if manual_content_path.exists():
            normalized = manual_content_path.read_text(encoding='utf-8', errors='ignore')
            manual_type = meta.get("manual_type") or "chapter"
            toc_depth = meta.get("toc_depth") or 2
            numbering_mode = meta.get("numbering_mode") or "css-counters"
            theme_settings = meta.get("theme_settings") or default_theme_settings(manual_type)
            theme_id = meta.get("theme_id") or "manual"

            manual_grid_block = build_manual_grid_block(normalized, toc_depth, manual_type, numbering_mode, theme_id=theme_id)

            if kind == "fragment":
                fragment_body = shift_heading_levels(normalized, 1)
                fragment_html = build_manual_grid_block(fragment_body, toc_depth, manual_type, numbering_mode, heading_offset=1, theme_id=theme_id)
                return send_file(io.BytesIO(fragment_html.encode('utf-8')), as_attachment=True, download_name="manual_fragment.html", mimetype="text/html")
            elif kind == "fragment_css":
                wp_css = get_wp_css_text()
                theme_css = build_theme_css(theme_settings)
                wp_js = get_wp_js_text()
                combined_css = f"{wp_css}\n{theme_css}"
                fragment_body = shift_heading_levels(normalized, 1)
                fragment_html = build_manual_grid_block(fragment_body, toc_depth, manual_type, numbering_mode, heading_offset=1, theme_id=theme_id)
                styled_fragment = f"<style>\n{combined_css}\n</style>\n{fragment_html}\n<script>\n{wp_js}\n</script>"
                return send_file(io.BytesIO(styled_fragment.encode('utf-8')), as_attachment=True, download_name="manual_fragment_styled.html", mimetype="text/html")
            else:
                wp_js = get_wp_js_text()
                wp_css = get_wp_css_text()
                theme_css = build_theme_css(theme_settings)
                combined_css = f"{wp_css}\n{theme_css}"
                standalone_html = f'<!doctype html><html lang="en"><head><meta charset="utf-8"><title>{meta.get("filename", "manual")}</title><style>{combined_css}</style></head><body>{manual_grid_block}<script>{wp_js}</script></body></html>'
                return send_file(io.BytesIO(standalone_html.encode('utf-8')), as_attachment=True, download_name="manual_standalone.html", mimetype="text/html")

    if kind == "docx":
        docx_path = Path(meta.get("docx_path", ""))
        if docx_path.exists():
            return send_file(str(docx_path), as_attachment=True, download_name=f"{meta.get('filename', 'document')}_numbered.docx")

    if kind == "heading_map":
        stable_map_file = session.stable_map_json
        if stable_map_file.exists():
            doc_stem = Path(meta.get("filename", "document")).stem
            return send_file(str(stable_map_file), as_attachment=True, download_name=f"{doc_stem}.heading-map.json", mimetype="application/json")

    flash("Download type not supported or file missing.")
    return redirect(url_for("index"))

@app.route("/update_theme", methods=["POST"])
def update_theme():
    session_id = request.form.get("session_id", "")
    token = request.form.get("token", "")
    if not session_id:
        flash("Missing session information.")
        return redirect(url_for("index"))

    session = SessionDir(session_id)
    session_file = session.session_json
    if not session_file.exists():
        flash("Session expired or invalid.")
        return redirect(url_for("index"))

    session_data = json.loads(session_file.read_text(encoding='utf-8'))
    manual_type = session_data.get('manual_type', 'chapter')
    toc_depth = session_data.get('toc_depth', 2)
    numbering_mode = session_data.get('numbering_mode', 'css-counters')
    
    # Update style panels state
    style_panels = session_data.get('style_panels', {"doc": True, "toc": False, "heading": False})
    style_panels["doc"] = request.form.get("doc_panel_open") == "1"
    style_panels["toc"] = request.form.get("toc_panel_open") == "1"
    style_panels["heading"] = request.form.get("heading_panel_open") == "1"
    session_data['style_panels'] = style_panels

    # Handle theme reset or update
    if 'reset_theme' in request.form:
        theme_settings, warnings = coerce_theme_settings(None, manual_type)
    else:
        theme_settings, warnings = coerce_theme_settings(request.form.to_dict(), manual_type)
    
    session_data['theme_settings'] = theme_settings
    session_file.write_text(json.dumps(session_data, indent=2), encoding='utf-8')

    flash("Styling updated.")
    return redirect(url_for('do_convert', session_id=session_id))

@app.route("/export/<session_id>", methods=["POST"])
def export_session(session_id):
    """Export session bundle (DOCX + edits + manifest)"""
    session = SessionDir(session_id)
    session_file = session.session_json
    if not session_file.exists():
        flash("Session not found.")
        return redirect(url_for("index"))

    session_data = json.loads(session_file.read_text(encoding='utf-8'))
    filename = session_data.get('filename', 'document.docx')
    src_path = Path(session_data.get('src_path', ''))
    
    # Use standardized edit path
    edit_file = session.edits_json
    if not src_path.exists() or not edit_file.exists():
        flash("Missing source DOCX or edits file; cannot export session.")
        return redirect(url_for("review", session_id=session_id))

    bundle_name = f"{Path(filename).stem}_{session_id[:8]}_session.zip"
    bundle_path = session.root / bundle_name

    manifest = {
        "document": filename,
        "doc_hash": compute_sha256(src_path),
        "manual_type": session_data.get('manual_type', 'chapter'),
        "toc_depth": session_data.get('toc_depth', 2),
        "mapping_mode": session_data.get('mapping_mode', 'map_new'),
        "theme_settings": session_data.get('theme_settings', {}),
        "heading_edits": session_data.get('heading_edits', {}),
        "stable_heading_map": session_data.get('stable_heading_map', {}),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "files": {
            "docx": src_path.name,
            "edits": edit_file.name
        }
    }

    with zipfile.ZipFile(bundle_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(src_path, arcname=src_path.name)
        zf.write(edit_file, arcname=edit_file.name)
        # Include the permalink continuity artifact if it exists
        stable_map_file = session.stable_map_json
        if stable_map_file.exists():
            zf.write(stable_map_file, arcname=stable_map_file.name)
        zf.writestr("manifest.json", json.dumps(manifest, indent=2))

    return send_file(str(bundle_path), as_attachment=True, download_name=bundle_name)

@app.route("/import_html", methods=["POST"])
def import_html():
    """Import HTML (standalone/fragment/WordPress) and start a review session"""
    f = request.files.get("html_file")
    if not f or f.filename == "":
        flash("Please upload an HTML file.")
        return redirect(url_for("index"))

    # Phase 1: Initialize Session
    session_id = str(uuid.uuid4())
    session = SessionDir(session_id)

    filename = secure_filename(f.filename)
    if not filename.lower().endswith((".html", ".htm")):
        flash("Only .html or .htm files are supported.")
        return redirect(url_for("index"))

    # Save to session-isolated root
    html_path = session.root / "import.html"
    f.save(str(html_path))

    strip_docx_formatting = "strip_docx_formatting" in request.form
    stable_heading_map, stable_heading_map_raw = load_heading_id_map_from_request()

    try:
        raw_html = html_path.read_text(encoding='utf-8', errors='ignore')
        cleaned_html = strip_html_assets(raw_html)
        html_path.write_text(cleaned_html, encoding='utf-8')
        manual_html, meta = extract_manual_fragment(cleaned_html)
        heading_offset = 0
        try:
            heading_offset = int(meta.get('heading_offset') or 0)
        except (ValueError, TypeError):
            heading_offset = 0
        if heading_offset:
            manual_html = shift_heading_levels(manual_html, -heading_offset)

        manual_type = meta.get('manual_type') or "chapter"
        toc_depth = meta.get('toc_depth') or "2"
        numbering_mode = meta.get('numbering_mode') or "css-counters"
        theme_id = meta.get('theme_id') or ""

        try:
            toc_depth = int(toc_depth)
            if toc_depth < 1 or toc_depth > 5:
                toc_depth = 2
        except (ValueError, TypeError):
            toc_depth = 2
        mapping_mode = "keep_old" if numbering_mode == "preserve" else "map_new"
        preserve_numbers = (numbering_mode == "preserve")

        manual_with_ids = add_heading_ids(
            manual_html,
            overwrite_existing=False,
            stable_map=stable_heading_map
        )
        scrape_html = manual_with_ids
        if numbering_mode == "css-counters":
            scrape_html = apply_css_counter_numbering(scrape_html, manual_type, preserve=False)
        new_headings = scrape_heading_structure_from_html(scrape_html)

        references = extract_references_from_html(manual_with_ids)

        if mapping_mode == "keep_old":
            auto_crosswalk = {ref[2]: ref[2] for ref in references}
        else:
            auto_crosswalk = auto_match_old_to_new_references(references, new_headings, manual_type=manual_type)

        theme_settings, _ = coerce_theme_settings({"theme_id": theme_id} if theme_id else None, manual_type)
        session_data = {
            'old_crosswalk': {},
            'auto_crosswalk': auto_crosswalk,
            'approved_crosswalk': {},
            'new_headings': new_headings,
            'references': references,
            'heading_map': {},
            'heading_order': {},
            'manual_type': manual_type,
            'filename': filename,
            'src_path': str(html_path),
            'pre_path': "",
            'temp_html_path': "",
            'toc_depth': toc_depth,
            'preserve_numbers': preserve_numbers,
            'mapping_mode': mapping_mode,
            'numbering_mode': numbering_mode,
            'html_import': True,
            'html_path': str(html_path),
            'rebuild_links': False,
            'strip_docx_formatting': strip_docx_formatting,
            'theme_settings': theme_settings,
            'heading_edits': {},
            'style_panels': {"doc": True, "toc": False, "heading": False},
            'infer_heading_depth': False,
            'infer_style_map': {},
            'infer_sequence_map': {},
            'stable_heading_map': stable_heading_map,
            'stable_heading_map_raw': stable_heading_map_raw,
        }
        session_file = session.session_json
        session_file.write_text(json.dumps(session_data, indent=2, default=str), encoding='utf-8')

        return redirect(url_for('review', session_id=session_id))
    except Exception as e:
        flash(f"HTML import failed: {e}")
        return redirect(url_for("index"))

@app.route("/import_bundle", methods=["POST"])
def import_bundle():
    """Import a session bundle zip and copy contents into isolated session workspace"""
    f = request.files.get("bundle")
    skip_review = 'skip_review' in request.form
    if not f or not f.filename.lower().endswith(".zip"):
        flash("Please upload a .zip session bundle.")
        return redirect(url_for("index"))

    # Phase 1: Initialize Session
    session_id = str(uuid.uuid4())
    session = SessionDir(session_id)
    
    # Save ZIP to session root
    bundle_path = session.root / "import_bundle.zip"
    f.save(str(bundle_path))

    try:
        with zipfile.ZipFile(bundle_path, 'r') as zf:
            # Safety check: Prevent path traversal.
            # Normalize separators and resolve each entry's final path so that
            # backslash variants ("..\\foo") and URL-encoded tricks are caught.
            safe_root = session.root.resolve()
            safe_root_str = str(safe_root)
            if not safe_root_str.endswith(os.sep):
                safe_root_str += os.sep
                
            for name in zf.namelist():
                # Reject obvious absolute paths and dot-dot segments
                normalized_name = name.replace('\\', '/')
                if normalized_name.startswith('/') or '..' in normalized_name.split('/'):
                    flash("Security error: Malicious bundle detected.")
                    return redirect(url_for("index"))
                # Resolve-and-confirm: ensure the entry stays inside the session root
                target = (safe_root / name).resolve()
                if not str(target).startswith(safe_root_str):
                    flash("Security error: Malicious bundle detected.")
                    return redirect(url_for("index"))
            zf.extractall(session.root)
            
        manifest_path = session.manifest_json
        if not manifest_path.exists():
            flash("Invalid bundle: manifest.json missing.")
            return redirect(url_for("index"))
            
        manifest = json.loads(manifest_path.read_text(encoding='utf-8'))
        doc_name = manifest["files"].get("docx")
        edits_name = manifest["files"].get("edits")
        
        doc_path = session.root / doc_name
        edits_path = session.root / edits_name

        if not doc_path.exists() or not edits_path.exists():
            flash("Invalid bundle: missing DOCX or edits file.")
            return redirect(url_for("index"))

        # Verify hash
        expected_hash = manifest.get("doc_hash", "")
        actual_hash = compute_sha256(doc_path)
        if expected_hash and expected_hash != actual_hash:
            flash("Warning: DOCX hash does not match manifest. Proceeding to import anyway.")

        # Standardize filenames in the isolated session
        dest_doc = session.source_docx
        dest_edits = session.edits_json
        
        # If the filenames in zip weren't already standardized, move them
        if doc_path != dest_doc:
            shutil.move(str(doc_path), str(dest_doc))
        if edits_path != dest_edits:
            shutil.move(str(edits_path), str(dest_edits))

        flash(f"Session imported for {doc_name}.")

        # Immediately start a review session from the imported doc/edits
        try:
            pre = session.pre_docx
            temp_html = session.temp_html

            logger.debug("="*80)
            logger.info("IMPORT: PREPROCESSING & EXTRACTING OLD REFERENCES")
            logger.debug("="*80)

            style_map = {}
            sequence_map = {}
            if manifest.get("mapping_mode", "map_new") == "map_new" and manifest.get("infer_heading_depth", False):
                style_map = manifest.get("infer_style_map", {}) or {}
                sequence_map = deserialize_sequence_map(manifest.get("infer_sequence_map", {}) or {})

            heading_map, old_crosswalk, references, manual_type = preprocess_docx(session.source_docx, session.pre_docx, style_map, sequence_map)
            docx_links_by_para = extract_docx_hyperlinks(Document(session.source_docx))

            logger.debug("="*80)
            logger.info("IMPORT: PRE-CONVERSION TO GET NEW STRUCTURE")
            logger.debug("="*80)

            run_pandoc(session.pre_docx, session.temp_html)
            temp_html_content = session.temp_html.read_text(encoding='utf-8', errors='ignore')
            temp_html_content = strip_pandoc_styles(normalize_spaces(temp_html_content))
            temp_html_content = strip_images_and_figures(temp_html_content)
            body = extract_body(temp_html_content)
            wrapped = f'<div class="manual">{body}</div>'
            normalized_html = normalize_typed_lists(wrapped)
            normalized_html = strip_toc_sections_dom(normalized_html)
            if manifest.get("mapping_mode", "map_new") == "map_new" and manifest.get("infer_heading_depth", False):
                style_map = manifest.get("infer_style_map", {}) or {}
                normalized_html = infer_heading_levels_from_prefix(normalized_html, style_map if style_map else None)
            normalized_html, _ = strip_heading_numbers_dom(normalized_html)
            normalized_html = apply_css_counter_numbering(normalized_html, manual_type, preserve=False)
            stable_heading_map = manifest.get("stable_heading_map", {}) or {}
            normalized_html = add_heading_ids(normalized_html, stable_map=stable_heading_map)

            logger.debug("="*80)
            logger.info("IMPORT: SCRAPING NEW HEADING STRUCTURE")
            logger.debug("="*80)
            new_headings = scrape_heading_structure_from_html(normalized_html)

            logger.debug("="*80)
            logger.info("IMPORT: AUTO-MATCHING OLD->NEW")
            logger.debug("="*80)
            auto_crosswalk = auto_match_old_to_new_references(references, new_headings, manual_type=manual_type)

            # Store session
            heading_order = {
                ensure_prefixed(normalize_heading_ref(k), manual_type): idx
                for idx, k in enumerate(heading_map.keys())
                if normalize_heading_ref(k)
            }
            theme_settings, _ = coerce_theme_settings(manifest.get("theme_settings"), manual_type)
            session_data = {
                'old_crosswalk': old_crosswalk,
                'auto_crosswalk': auto_crosswalk,
                'approved_crosswalk': {},
                'new_headings': new_headings,
                'references': references,
                'heading_map': heading_map,
                'heading_order': heading_order,
                'manual_type': manual_type,
                'filename': session.source_docx.name,
                'src_path': str(session.source_docx),
                'pre_path': str(session.pre_docx),
                'temp_html_path': str(session.temp_html),
                'toc_depth': manifest.get("toc_depth", 2),
                'preserve_numbers': manifest.get("preserve_numbers", False),
                'mapping_mode': manifest.get("mapping_mode", 'map_new'),
                'theme_settings': theme_settings,
                'infer_heading_depth': manifest.get("infer_heading_depth", False),
                'infer_style_map': manifest.get("infer_style_map", {}),
                'infer_sequence_map': manifest.get("infer_sequence_map", {}),
                'heading_edits': manifest.get("heading_edits", {}),
                'stable_heading_map': manifest.get("stable_heading_map", {}) or {},
                'stable_heading_map_raw': manifest.get("stable_heading_map_raw", "") or "",
                'strip_docx_formatting': manifest.get("strip_docx_formatting", False),
                'docx_links_by_para': docx_links_by_para,
            }
            # If edits file exists in bundle, seed approved_crosswalk from it
            if session.edits_json.exists():
                try:
                    edits_data = json.loads(session.edits_json.read_text(encoding='utf-8'))
                    appr = edits_data.get('approved_crosswalk', {})
                    if appr:
                        session_data['approved_crosswalk'] = appr
                except Exception as e:
                    logger.error(f"DEBUG: Failed to seed approved_crosswalk from edits: {e}")
            session_file = session.session_json
            session_file.write_text(json.dumps(session_data, indent=2, default=str), encoding='utf-8')

            if skip_review:
                return redirect(url_for('do_convert', session_id=session_id))
            return redirect(url_for('heading_review', session_id=session_id))
        except Exception as e:
            flash(f"Import succeeded but could not start review automatically: {e}")
            return redirect(url_for("index"))
    except Exception as e:
        flash(f"Import failed: {e}")
        return redirect(url_for("index"))

# --------------------------------- Main ------------------------------------

if __name__ == "__main__":
    try:
        subprocess.run(["pandoc", "-v"], check=True, capture_output=True)
    except Exception:
        logger.info("Pandoc is not installed or not on PATH. Install from https://pandoc.org/installing.html")
        raise
    logger.info(f"Persist directory for edits: {PERSIST_DIR}")

    port = int(os.environ.get("PORT", 5000))
    is_local = not os.environ.get("PORT")

    if is_local:
        logger.info(f"Starting on http://127.0.0.1:{port}")
        import webbrowser, threading
        threading.Timer(1.0, webbrowser.open, args=[f"http://127.0.0.1:{port}"]).start()
    else:
        logger.info(f"Starting on port {port}")

    app.run(host="0.0.0.0", port=port, debug=is_local)
