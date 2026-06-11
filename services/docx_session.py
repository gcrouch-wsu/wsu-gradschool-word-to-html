"""Shared DOCX session-start pipeline.

/convert and /import_bundle both run the same Phase 1-3 sequence (preprocess
DOCX -> Pandoc -> normalize HTML) and persist the same session_data shape.
Keeping both here means the two paths cannot drift apart again.
"""
import logging
import re

from docx import Document

from core.docx_processor import preprocess_docx, extract_docx_hyperlinks
from core.html_processor import (
    add_heading_ids,
    extract_body,
    normalize_spaces,
    normalize_typed_lists,
    strip_images_and_figures,
    strip_pandoc_styles,
    strip_toc_sections_dom,
)
from core.manual_structure import scrape_heading_structure_from_html
from core.pandoc_wrapper import run_pandoc
from core.permalinks import ensure_prefixed, normalize_heading_ref
from config import SessionDir

logger = logging.getLogger(__name__)


def run_docx_prepipeline(session: SessionDir, style_map: dict, sequence_map: dict):
    """Phases 1-2: preprocess the uploaded DOCX, convert with Pandoc, and
    normalize the raw HTML up to (but not including) heading-number handling,
    which differs per caller.

    Returns (heading_map, old_crosswalk, references, manual_type,
    docx_links_by_para, normalized_html).
    """
    heading_map, old_crosswalk, references, manual_type = preprocess_docx(
        session.source_docx, session.pre_docx, style_map, sequence_map
    )
    docx_links_by_para = extract_docx_hyperlinks(Document(session.source_docx))
    logger.info(f"Extracted {len(references)} OLD references from DOCX")

    run_pandoc(session.pre_docx, session.temp_html)
    logger.info(f"Generated preliminary HTML: {session.temp_html}")

    temp_html_content = session.temp_html.read_text(encoding='utf-8', errors='ignore')
    temp_html_content = strip_pandoc_styles(normalize_spaces(temp_html_content))
    temp_html_content = strip_images_and_figures(temp_html_content)
    body = extract_body(temp_html_content)

    # Wrap body in .manual so the list normalizer has context
    wrapped = f'<div class="manual">{body}</div>'
    normalized_html = normalize_typed_lists(wrapped)
    normalized_html = strip_toc_sections_dom(normalized_html)
    return heading_map, old_crosswalk, references, manual_type, docx_links_by_para, normalized_html


def scrape_new_structure(normalized_html: str, stable_heading_map: dict) -> tuple[str, dict]:
    """Phase 3: assign heading IDs (honoring the stable map) and scrape the
    NEW heading structure from the converted HTML."""
    normalized_html = add_heading_ids(normalized_html, stable_map=stable_heading_map)
    new_headings = scrape_heading_structure_from_html(normalized_html)
    logger.info(f"Scraped {len(new_headings)} NEW headings from converted HTML")
    return normalized_html, new_headings


def build_identity_crosswalk(references: list) -> dict:
    """keep_old mode: every reference maps to itself (OLD -> OLD)."""
    auto_crosswalk = {}
    for ref in references:
        old_ref_text = ref[2]
        # Skip if it doesn't look like a chapter/section reference
        if re.match(r'^(Chapter|Section)\s+[\dIVXLCDM]+', old_ref_text, re.IGNORECASE):
            auto_crosswalk[old_ref_text] = old_ref_text
    return auto_crosswalk


def build_heading_order(heading_map: dict, manual_type: str) -> dict:
    return {
        ensure_prefixed(normalize_heading_ref(k), manual_type): idx
        for idx, k in enumerate(heading_map.keys())
        if normalize_heading_ref(k)
    }


def build_session_data(
    *,
    manual_type: str,
    filename: str,
    src_path,
    references: list,
    new_headings: dict,
    auto_crosswalk: dict,
    heading_map: dict | None = None,
    old_crosswalk: dict | None = None,
    heading_order: dict | None = None,
    approved_crosswalk: dict | None = None,
    pre_path="",
    temp_html_path="",
    toc_depth: int = 2,
    preserve_numbers: bool = False,
    numbering_mode: str | None = None,
    mapping_mode: str = "map_new",
    strip_docx_formatting: bool = False,
    theme_settings: dict | None = None,
    heading_edits: dict | None = None,
    infer_heading_depth: bool = False,
    infer_style_map: dict | None = None,
    infer_sequence_map=None,
    stable_heading_map: dict | None = None,
    stable_heading_map_raw: str = "",
    docx_links_by_para: dict | None = None,
    edit_tables: bool = False,
    html_import: bool = False,
    html_path="",
    rebuild_links: bool = False,
) -> dict:
    """The single canonical session_data shape used by every session-creating
    route (/convert, /import_html, /import_bundle)."""
    return {
        'old_crosswalk': old_crosswalk or {},
        'auto_crosswalk': auto_crosswalk,
        'approved_crosswalk': approved_crosswalk or {},
        'new_headings': new_headings,
        'references': references,
        'heading_map': heading_map or {},
        'heading_order': heading_order or {},
        'manual_type': manual_type,
        'filename': filename,
        'src_path': str(src_path),
        'pre_path': str(pre_path) if pre_path else "",
        'temp_html_path': str(temp_html_path) if temp_html_path else "",
        'toc_depth': toc_depth,
        'preserve_numbers': preserve_numbers,
        'numbering_mode': numbering_mode or ("preserve" if preserve_numbers else "css-counters"),
        'mapping_mode': mapping_mode,
        'html_import': html_import,
        'html_path': str(html_path) if html_path else "",
        'rebuild_links': rebuild_links,
        'strip_docx_formatting': strip_docx_formatting,
        'theme_settings': theme_settings or {},
        'heading_edits': heading_edits or {},
        'style_panels': {"doc": True, "toc": False, "heading": False},
        'infer_heading_depth': infer_heading_depth,
        'infer_style_map': infer_style_map or {},
        'infer_sequence_map': infer_sequence_map or {},
        'stable_heading_map': stable_heading_map or {},
        'stable_heading_map_raw': stable_heading_map_raw or "",
        'docx_links_by_para': docx_links_by_para or {},
        'edit_tables': edit_tables,
    }
