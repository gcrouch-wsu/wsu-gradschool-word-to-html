import docx
import sys

def read_sections(docx_path, start_para, end_para):
    doc = docx.Document(docx_path)
    for i in range(start_para, min(end_para, len(doc.paragraphs))):
        print(f"[{i+1}] {doc.paragraphs[i].text}")

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python scripts/read_manual.py <path_to_docx> <start_para> <end_para>")
        sys.exit(1)
    read_sections(sys.argv[1], int(sys.argv[2]), int(sys.argv[3]))
