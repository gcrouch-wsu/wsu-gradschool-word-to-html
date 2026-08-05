#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX Configuration Generator

Standalone Flask app for analyzing DOCX files and generating JSON configuration manifests.
This app allows users to configure styles, headings, and other settings that will be
used by the main conversion app.
"""

import os
import re
import json
import tempfile
import uuid
from datetime import datetime
from pathlib import Path

from flask import Flask, abort, request, render_template_string, send_file, redirect, url_for, flash
from werkzeug.exceptions import RequestEntityTooLarge

from config import is_valid_session_id

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Shared helpers — single canonical implementations live in core/ and utils/
from utils.helpers import _int_to_roman, _int_to_letters
from core.docx_processor import (
    is_heading_style,
    serialize_sequence_map,
    _extract_numbering_defs,
    _extract_heading_prefix_tokens,
    _classify_heading_token,
)
from core.html_processor import _normalized

# ------------------------------ Flask setup ------------------------------

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "config-generator-secret")
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024
app.config["MAX_FORM_MEMORY_SIZE"] = 200 * 1024 * 1024

PERSIST_DIR = Path(tempfile.gettempdir()) / "docx_config_generator"
PERSIST_DIR.mkdir(parents=True, exist_ok=True)


def _session_paths(session_id: str):
    """Resolve a session's files, or raise 404 if the id is not one we issued.

    Session ids are always ``str(uuid.uuid4())`` generated here, so anything else
    arriving in a URL is hostile or stale. Interpolating one straight into a path
    let ``/export/..%5Cwhatever`` read a file outside PERSIST_DIR on Windows, and
    ``/example/..%5Cwhatever`` write one — the backslash form survives Werkzeug's
    routing, which only collapses forward slashes. `config.is_valid_session_id` is
    the same check the main app applies for the same reason.
    """
    if not is_valid_session_id(session_id):
        abort(404)
    return (
        PERSIST_DIR / f"{session_id}_config.json",
        PERSIST_DIR / f"{session_id}_analysis.json",
    )

@app.errorhandler(RequestEntityTooLarge)
def handle_request_too_large(error):
    flash("File too large. Please use a smaller DOCX file.")
    return redirect(url_for("index"))

# ------------------------------ Utility Functions ------------------------------

def rgb_to_hex(rgb_color):
    """Convert RGBColor to hex string."""
    if rgb_color is None:
        return None
    if isinstance(rgb_color, RGBColor):
        return f"#{rgb_color.rgb:06x}"
    return None

def pt_to_px(pt_value):
    """Convert points to pixels (approximate: 1pt = 1.33px)."""
    if pt_value is None:
        return None
    try:
        # Handle Length objects from python-docx
        if hasattr(pt_value, 'pt'):
            pt_value = pt_value.pt
        return round(float(pt_value) * 1.33, 1)
    except (TypeError, ValueError, AttributeError):
        return None

def length_to_pt(value):
    """Convert a Word length to points."""
    if value is None:
        return None
    try:
        if hasattr(value, 'pt'):
            return value.pt
        return float(value)
    except (TypeError, ValueError, AttributeError):
        return None

def extract_paragraph_format(paragraph_format):
    """Extract paragraph formatting details."""
    if paragraph_format is None:
        return {}
    alignment = None
    if paragraph_format.alignment is not None:
        alignment = getattr(paragraph_format.alignment, "name", None)
        if alignment:
            alignment = alignment.lower()
    line_spacing_rule = None
    if paragraph_format.line_spacing_rule is not None:
        line_spacing_rule = getattr(paragraph_format.line_spacing_rule, "name", None)
        if line_spacing_rule:
            line_spacing_rule = line_spacing_rule.lower()
    line_spacing = paragraph_format.line_spacing
    if line_spacing is not None:
        try:
            if hasattr(line_spacing, 'pt'):
                line_spacing = line_spacing.pt
            else:
                line_spacing = float(line_spacing)
        except (TypeError, ValueError):
            line_spacing = None
    return {
        "alignment": alignment,
        "left_indent_pt": length_to_pt(paragraph_format.left_indent),
        "right_indent_pt": length_to_pt(paragraph_format.right_indent),
        "first_line_indent_pt": length_to_pt(paragraph_format.first_line_indent),
        "space_before_pt": length_to_pt(paragraph_format.space_before),
        "space_after_pt": length_to_pt(paragraph_format.space_after),
        "line_spacing": line_spacing,
        "line_spacing_rule": line_spacing_rule,
    }

def extract_style_info(style, include_none=False):
    """Extract style information from a Word style."""
    info = {
        "name": style.name,
        "font_name": None,
        "font_size": None,
        "font_size_pt": None,
        "color": None,
        "bold": None if include_none else False,
        "italic": None if include_none else False,
        "underline": None if include_none else False,
        "paragraph_format": extract_paragraph_format(getattr(style, "paragraph_format", None)),
    }
    
    try:
        font = style.font
        if font.name:
            info["font_name"] = font.name
        if font.size:
            # font.size is a Length object, get value in points
            try:
                size_pt = font.size.pt if hasattr(font.size, 'pt') else float(font.size)
                info["font_size_pt"] = size_pt
                info["font_size"] = pt_to_px(size_pt)
            except (AttributeError, ValueError, TypeError):
                try:
                    # Try direct conversion
                    info["font_size"] = float(font.size)
                    info["font_size_pt"] = float(font.size)
                except:
                    pass
        if font.color and font.color.rgb:
            info["color"] = rgb_to_hex(font.color)
        if font.bold is not None:
            info["bold"] = bool(font.bold)
        elif not include_none:
            info["bold"] = False
        if font.italic is not None:
            info["italic"] = bool(font.italic)
        elif not include_none:
            info["italic"] = False
        if font.underline is not None:
            info["underline"] = bool(font.underline)
        elif not include_none:
            info["underline"] = False
    except Exception as e:
        print(f"Error extracting style info for {style.name}: {e}")
    
    return info

def resolve_style_info(style):
    """Resolve inherited style values from the base style chain."""
    chain = []
    visited = set()
    current = style
    while current is not None and current.name not in visited:
        chain.append(current)
        visited.add(current.name)
        current = current.base_style

    resolved = {
        "name": style.name,
        "font_name": None,
        "font_size": None,
        "font_size_pt": None,
        "color": None,
        "bold": None,
        "italic": None,
        "underline": None,
        "paragraph_format": {
            "alignment": None,
            "left_indent_pt": None,
            "right_indent_pt": None,
            "first_line_indent_pt": None,
            "space_before_pt": None,
            "space_after_pt": None,
            "line_spacing": None,
            "line_spacing_rule": None,
        },
        "style_chain": [s.name for s in chain],
    }

    for current_style in chain:
        info = extract_style_info(current_style, include_none=True)
        for key in ["font_name", "font_size", "font_size_pt", "color", "bold", "italic", "underline"]:
            if resolved[key] is None and info.get(key) is not None:
                resolved[key] = info.get(key)
        paragraph_format = info.get("paragraph_format", {})
        for key in resolved["paragraph_format"]:
            if resolved["paragraph_format"][key] is None and paragraph_format.get(key) is not None:
                resolved["paragraph_format"][key] = paragraph_format.get(key)

    return resolved

def infer_bullet_type(lvl_text):
    """Infer a bullet style from Word list text."""
    if not lvl_text:
        return "disc"
    if any(sym in lvl_text for sym in ["\u25E6", "\u25CB", "o"]):
        return "circle"
    if any(sym in lvl_text for sym in ["\u25AA", "\u25A0", "\u25A1", "\u25AB", "\u25FC", "\u25FD"]):
        return "square"
    if any(sym in lvl_text for sym in ["\u2013", "\u2014", "-"]):
        return "dash"
    return "disc"

def bullet_glyph(bullet_type):
    """Return a glyph for the requested bullet type."""
    mapping = {
        "disc": "\u2022",
        "circle": "\u25E6",
        "square": "\u25AA",
        "dash": "\u2013",
    }
    return mapping.get(bullet_type, "\u2022")

def hex_to_rgb(hex_color):
    """Convert a hex color string to RGBColor."""
    if not hex_color:
        return None
    value = hex_color.lstrip("#")
    if len(value) != 6:
        return None
    try:
        return RGBColor.from_string(value.upper())
    except Exception:
        return None

def config_font_size_pt(style_info):
    """Get font size in points from a config style dict."""
    if not style_info:
        return None
    if style_info.get("font_size_pt"):
        return style_info.get("font_size_pt")
    size_pt = style_info.get("font_size")
    if size_pt:
        try:
            return float(size_pt)
        except (TypeError, ValueError):
            return None
    return None

def apply_run_style(run, style_info):
    """Apply font styling to a run."""
    if not style_info:
        return
    font = run.font
    font_name = style_info.get("font_name")
    font_size_pt = config_font_size_pt(style_info)
    font_color = hex_to_rgb(style_info.get("color"))
    if font_name:
        font.name = font_name
    if font_size_pt:
        font.size = Pt(font_size_pt)
    if font_color:
        font.color.rgb = font_color
    if style_info.get("bold") is not None:
        run.bold = bool(style_info.get("bold"))
    if style_info.get("italic") is not None:
        run.italic = bool(style_info.get("italic"))
    if style_info.get("underline") is not None:
        run.underline = bool(style_info.get("underline"))

def apply_paragraph_format(paragraph, format_info):
    """Apply paragraph spacing/indentation rules."""
    if not format_info:
        return
    pf = paragraph.paragraph_format
    alignment = format_info.get("alignment")
    if alignment:
        alignment_map = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        pf.alignment = alignment_map.get(alignment, pf.alignment)
    if format_info.get("left_indent_pt") is not None:
        pf.left_indent = Pt(format_info["left_indent_pt"])
    if format_info.get("right_indent_pt") is not None:
        pf.right_indent = Pt(format_info["right_indent_pt"])
    if format_info.get("first_line_indent_pt") is not None:
        pf.first_line_indent = Pt(format_info["first_line_indent_pt"])
    if format_info.get("space_before_pt") is not None:
        pf.space_before = Pt(format_info["space_before_pt"])
    if format_info.get("space_after_pt") is not None:
        pf.space_after = Pt(format_info["space_after_pt"])
    if format_info.get("line_spacing") is not None:
        pf.line_spacing = format_info["line_spacing"]

def round_half(value):
    """Round a numeric value to the nearest 0.5."""
    try:
        return round(float(value) * 2) / 2
    except (TypeError, ValueError):
        return value

def normalize_format_name(name):
    """Normalize list format names for consistency."""
    if not name:
        return ""
    name = str(name)
    if name.lower() == "decimalzero":
        return "decimalZero"
    return name

def default_ordered_sequence(first_format):
    """Return a reasonable default ordered list sequence based on the first format."""
    fmt = normalize_format_name(first_format)
    if fmt in {"lowerLetter", "upperLetter"}:
        roman = "lowerRoman" if fmt == "lowerLetter" else "upperRoman"
        return [fmt, roman, "decimal"]
    if fmt in {"lowerRoman", "upperRoman"}:
        letter = "lowerLetter" if fmt == "lowerRoman" else "upperLetter"
        return [fmt, "decimal", letter]
    if fmt == "decimalZero":
        return ["decimalZero", "lowerLetter", "lowerRoman"]
    return ["decimal", "lowerLetter", "lowerRoman"]

def ensure_ordered_format_levels(formats, max_level=5):
    """Ensure ordered list formats exist through the requested level."""
    if formats is None:
        formats = {}
    first_format = None
    for level_key in sorted(formats.keys(), key=lambda x: int(x) if str(x).isdigit() else 0):
        info = formats.get(level_key) or {}
        if isinstance(info, dict) and info.get("disabled"):
            continue
        fmt_value = info.get("format") if isinstance(info, dict) else None
        if fmt_value:
            first_format = fmt_value
            break
    if not first_format:
        first_format = "decimal"
    sequence = default_ordered_sequence(first_format)
    for level in range(max_level + 1):
        key = str(level)
        if key not in formats:
            fmt = sequence[level % len(sequence)]
            formats[key] = {
                "format": fmt,
                "lvl_text": f"%{level + 1}." if fmt else "",
                "start": 1,
            }
        else:
            if formats[key].get("disabled"):
                formats[key]["format"] = ""
                formats[key]["lvl_text"] = ""
                continue
            formats[key]["format"] = normalize_format_name(formats[key].get("format"))
            if formats[key]["format"]:
                if not formats[key].get("lvl_text"):
                    formats[key]["lvl_text"] = f"%{level + 1}."
                formats[key].setdefault("start", 1)
            else:
                fmt = sequence[level % len(sequence)]
                formats[key]["format"] = fmt
                formats[key]["lvl_text"] = formats[key].get("lvl_text") or f"%{level + 1}."
                formats[key].setdefault("start", 1)
    return formats

def ensure_unordered_levels(formats, max_level=5):
    """Ensure unordered list bullet types exist through the requested level."""
    if formats is None:
        formats = {}
    sequence = ["disc", "circle", "square", "dash", "disc"]
    for level in range(max_level + 1):
        key = str(level)
        if key not in formats:
            formats[key] = sequence[level % len(sequence)]
        else:
            value = formats.get(key)
            if isinstance(value, dict):
                if value.get("disabled"):
                    value["type"] = ""
                    formats[key] = value
                else:
                    value.setdefault("type", value.get("type") or "")
                    formats[key] = value
    return formats

def format_list_marker(format_name, lvl_text, value, placeholder_index=None):
    """Format a list marker using Word-style format names and templates."""
    fmt = (format_name or "decimal").lower()
    if fmt == "decimalzero":
        formatted = f"{value:02d}"
    elif fmt in {"upperletter", "upperalpha"}:
        formatted = _int_to_letters(value, upper=True)
    elif fmt in {"lowerletter", "loweralpha"}:
        formatted = _int_to_letters(value, upper=False)
    elif fmt == "upperroman":
        formatted = _int_to_roman(value)
    elif fmt == "lowerroman":
        formatted = _int_to_roman(value).lower()
    else:
        formatted = str(value)

    template = lvl_text or "%1."
    if placeholder_index is None:
        placeholders = [int(n) for n in re.findall(r"%(\d+)", template)]
        placeholder_index = max(placeholders) if placeholders else 1
    def repl(match):
        idx = int(match.group(1))
        if idx == placeholder_index:
            return formatted
        return ""
    marker = re.sub(r"%(\d+)", repl, template).strip()
    marker = marker.lstrip(".-–— ").strip()
    marker = re.sub(r"\.{2,}", ".", marker).strip()
    marker = re.sub(r"\s+", " ", marker).strip()
    return marker or formatted

def build_example_docx(config, output_path):
    """Build a sample DOCX showing resolved styles and list formats."""
    doc = Document()
    body_style = config.get("styles", {}).get("body", {})
    heading_styles = config.get("styles", {}).get("headings", {})
    lists_cfg = config.get("lists", {})
    ordered_formats = lists_cfg.get("multilevel_formats", {}) or {}
    unordered_formats = lists_cfg.get("unordered_formats", {}) or {}

    def add_heading(text, level_key):
        style_info = heading_styles.get(level_key, {})
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, style_info.get("paragraph_format", {}))
        run = paragraph.add_run(text)
        apply_run_style(run, style_info)
        return paragraph

    def add_body(text):
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, body_style.get("paragraph_format", {}))
        run = paragraph.add_run(text)
        apply_run_style(run, body_style)
        return paragraph

    add_heading("Style Samples", "h1")
    add_body("This document shows the full set of heading and list styles used for the DOCX export.")

    for level in range(1, 7):
        add_heading(f"Heading Level {level} Sample", f"h{level}")
        add_body("Sample paragraph text to show body styling under this heading level.")

    add_heading("Ordered List Samples", "h2")
    if ordered_formats:
        for level_key in sorted(ordered_formats.keys(), key=lambda x: int(x)):
            fmt_info = ordered_formats[level_key]
            if isinstance(fmt_info, dict) and fmt_info.get("disabled"):
                continue
            fmt_name = fmt_info.get("format") if isinstance(fmt_info, dict) else fmt_info
            if not fmt_name:
                continue
            lvl_text = fmt_info.get("lvl_text") if isinstance(fmt_info, dict) else None
            for value in (1, 2):
                marker = format_list_marker(fmt_name, lvl_text, value, int(level_key) + 1)
                paragraph = doc.add_paragraph()
                apply_paragraph_format(paragraph, body_style.get("paragraph_format", {}))
                indent_base = body_style.get("paragraph_format", {}).get("left_indent_pt") or 0
                paragraph.paragraph_format.left_indent = Pt(indent_base + (int(level_key) * 18))
                run_marker = paragraph.add_run(f"{marker} ")
                apply_run_style(run_marker, body_style)
                run_text = paragraph.add_run(f"Sample ordered item {value} (level {level_key})")
                apply_run_style(run_text, body_style)
    else:
        add_body("No ordered list formats detected.")

    add_heading("Unordered List Samples", "h2")
    if unordered_formats:
        for level_key in sorted(unordered_formats.keys(), key=lambda x: int(x)):
            bullet_type = unordered_formats[level_key]
            if isinstance(bullet_type, dict):
                if bullet_type.get("disabled"):
                    continue
                bullet_type = bullet_type.get("type") or infer_bullet_type(bullet_type.get("lvl_text", ""))
            if not bullet_type:
                continue
            glyph = bullet_glyph(bullet_type)
            for value in (1, 2):
                paragraph = doc.add_paragraph()
                apply_paragraph_format(paragraph, body_style.get("paragraph_format", {}))
                indent_base = body_style.get("paragraph_format", {}).get("left_indent_pt") or 0
                paragraph.paragraph_format.left_indent = Pt(indent_base + (int(level_key) * 18))
                run_marker = paragraph.add_run(f"{glyph} ")
                apply_run_style(run_marker, body_style)
                run_text = paragraph.add_run(f"Sample unordered item {value} (level {level_key}, {bullet_type})")
                apply_run_style(run_text, body_style)
    else:
        add_body("No unordered list formats detected.")

    add_heading("Notes", "h2")
    add_body("This document is a visual sample of resolved styles. List markers are shown as text for clarity.")

    doc.save(str(output_path))

def strip_heading_prefix(text):
    """Remove common numbering prefixes from a heading."""
    return _strip_heading_prefix_for_preview(text)

def numeric_prefix(level):
    """Generate a simple numeric prefix for preview purposes."""
    try:
        level_int = int(level)
    except (TypeError, ValueError):
        level_int = 1
    return ".".join(["1"] * max(level_int, 1))

def build_preview_style(style_info, body_style):
    """Build inline CSS for preview text."""
    font_family = (style_info or {}).get("font_name") or body_style.get("font") or "Calibri"
    font_family = font_family.replace('"', '').replace("'", "")
    font_size = (style_info or {}).get("font_size_pt") or (style_info or {}).get("font_size") or body_style.get("size") or 12
    try:
        font_size = float(font_size)
    except (TypeError, ValueError):
        font_size = 12.0
    color = (style_info or {}).get("color") or body_style.get("color") or "#111111"
    bold = (style_info or {}).get("bold")
    italic = (style_info or {}).get("italic")
    underline = (style_info or {}).get("underline")
    line_height = (style_info or {}).get("paragraph_format", {}).get("line_spacing")
    if line_height is None:
        line_height = body_style.get("line_height") or 1.2
    try:
        line_height = float(line_height)
    except (TypeError, ValueError):
        line_height = 1.2
    weight = 700 if bold else 400
    font_style = "italic" if italic else "normal"
    text_decoration = "underline" if underline else "none"
    return (
        f"font-family: '{font_family}', sans-serif; "
        f"font-size: {font_size}pt; "
        f"color: {color}; "
        f"font-weight: {weight}; "
        f"font-style: {font_style}; "
        f"text-decoration: {text_decoration}; "
        f"line-height: {line_height};"
    )


def build_preview_data(config, analysis):
    """Build preview data from config and optional analysis."""
    styles = config.get("styles", {})
    body_style = styles.get("body", {})
    heading_styles = styles.get("headings", {})
    list_cfg = config.get("lists", {})
    ordered_formats = list_cfg.get("multilevel_formats", {}) or {}
    unordered_formats = list_cfg.get("unordered_formats", {}) or {}
    body_paragraph = body_style.get("paragraph_format", {}) or {}
    base_indent_pt = body_paragraph.get("left_indent_pt") or 0

    analysis = analysis or {}
    style_samples = analysis.get("style_samples", {})
    heading_samples = style_samples.get("headings", {}) or {}
    body_sample = style_samples.get("body", {})
    original_list_formats = analysis.get("list_formats", {})
    original_ordered_formats = original_list_formats.get("ordered", {}) or {}
    original_unordered_formats = original_list_formats.get("unordered", {}) or {}

    body_text = body_sample.get("text") if body_sample else ""
    if not body_text:
        body_text = "Body text preview. Adjust font, size, color, and spacing to match the final DOCX export."

    body_preview_style = build_preview_style({}, body_style)

    heading_previews = []
    for level in range(1, 7):
        level_key = str(level)
        sample = heading_samples.get(level_key, {})
        text = sample.get("text") if isinstance(sample, dict) else ""
        if not text:
            text = f"Heading level {level} preview"
        style_info = heading_styles.get(f"h{level}", {})
        heading_previews.append({
            "level": level,
            "text": text,
            "style": build_preview_style(style_info, body_style),
        })

    mapping_preview = []
    for level in range(1, 7):
        level_key = str(level)
        sample = heading_samples.get(level_key, {})
        keep_text = sample.get("text") if isinstance(sample, dict) else ""
        if not keep_text:
            keep_text = f"Heading level {level}"
        base_text = strip_heading_prefix(keep_text) or keep_text
        map_text = f"{numeric_prefix(level)} {base_text}".strip()
        style_info = heading_styles.get(f"h{level}", {})
        mapping_preview.append({
            "level": level,
            "keep_text": keep_text,
            "map_text": map_text,
            "output_text": map_text if config.get("conversion", {}).get("mapping_mode") != "keep_old" else keep_text,
            "style": build_preview_style(style_info, body_style),
        })

    mapping_mode = config.get("conversion", {}).get("mapping_mode", "map_new")
    heading_examples = []
    source_headings = analysis.get("headings", []) or []
    seen_levels = set()
    for heading in source_headings:
        level = heading.get("level") or 1
        if level in seen_levels:
            continue
        keep_text = heading.get("text") or f"Heading level {level}"
        base_text = _strip_heading_prefix_for_preview(keep_text) or keep_text
        map_text = f"{numeric_prefix(level)} {base_text}".strip()
        style_info = heading_styles.get(f"h{level}", {})
        heading_examples.append({
            "level": level,
            "keep_text": keep_text,
            "output_text": keep_text if mapping_mode == "keep_old" else map_text,
            "style": build_preview_style(style_info, body_style),
        })
        seen_levels.add(level)
        if len(seen_levels) >= 6:
            break
    if not heading_examples:
        heading_examples = [
            {
                "level": item.get("level"),
                "keep_text": item.get("keep_text"),
                "output_text": item.get("output_text"),
                "style": item.get("style"),
            }
            for item in mapping_preview
        ]

    def build_ordered_preview(formats):
        preview = []
        for level_key in sorted(formats.keys(), key=lambda x: int(x)):
            fmt_info = formats[level_key] or {}
            if isinstance(fmt_info, dict) and fmt_info.get("disabled"):
                continue
            fmt_name = normalize_format_name(fmt_info.get("format") if isinstance(fmt_info, dict) else fmt_info)
            if not fmt_name:
                continue
            lvl_text = fmt_info.get("lvl_text") if isinstance(fmt_info, dict) else None
            marker = format_list_marker(fmt_name, lvl_text, 1, int(level_key) + 1)
            try:
                level_int = int(level_key)
            except (TypeError, ValueError):
                level_int = 0
            preview.append({
                "level": level_key,
                "marker": marker,
                "text": f"Ordered list item level {level_key}",
                "indent_pt": base_indent_pt + (level_int * 18),
            })
        return preview

    def build_unordered_preview(formats):
        preview = []
        for level_key in sorted(formats.keys(), key=lambda x: int(x)):
            bullet_type = formats[level_key]
            if isinstance(bullet_type, dict):
                if bullet_type.get("disabled"):
                    continue
                bullet_type = bullet_type.get("type") or infer_bullet_type(bullet_type.get("lvl_text", ""))
            if not bullet_type:
                continue
            try:
                level_int = int(level_key)
            except (TypeError, ValueError):
                level_int = 0
            preview.append({
                "level": level_key,
                "marker": bullet_glyph(bullet_type),
                "text": f"Unordered list item level {level_key}",
                "indent_pt": base_indent_pt + (level_int * 18),
            })
        return preview


    ordered_preview = build_ordered_preview(ordered_formats)
    unordered_preview = build_unordered_preview(unordered_formats)
    original_ordered_preview = build_ordered_preview(original_ordered_formats)
    original_unordered_preview = build_unordered_preview(original_unordered_formats)

    return {
        "body": {"text": body_text, "style": body_preview_style},
        "headings": heading_previews,
        "mapping": mapping_preview,
        "ordered_list": ordered_preview,
        "unordered_list": unordered_preview,
        "original_ordered_list": original_ordered_preview,
        "original_unordered_list": original_unordered_preview,
        "heading_examples": heading_examples,
        "mapping_mode": config.get("conversion", {}).get("mapping_mode", "map_new"),
    }

def _get_paragraph_numpr(p):
    """Get numbering properties from paragraph or its style."""
    ppr = p._p.pPr
    if ppr is not None and ppr.numPr is not None:
        return ppr.numPr, None
    style = getattr(p, "style", None)
    try:
        style_ppr = style.element.pPr if style is not None else None
    except Exception:
        style_ppr = None
    if style_ppr is not None and style_ppr.numPr is not None:
        ilvl = None
        try:
            if style_ppr.numPr.ilvl is not None:
                ilvl = style_ppr.numPr.ilvl.val
            else:
                ilvl = 0
        except Exception:
            ilvl = 0
        return style_ppr.numPr, ilvl
    return None, None

def _format_list_number(value, fmt, level):
    """Format a number according to list format type."""
    fmt = (fmt or "decimal").lower()
    if fmt in {"decimal", "decimalzero"}:
        return str(value)
    elif fmt == "lowerroman":
        return _int_to_roman(value).lower()
    elif fmt == "upperroman":
        return _int_to_roman(value)
    elif fmt == "lowerletter":
        return _int_to_letters(value, upper=False)
    elif fmt == "upperletter":
        return _int_to_letters(value, upper=True)
    return str(value)

_CHAPTER_WORDS = (
    "One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve|"
    "Thirteen|Fourteen|Fifteen|Sixteen|Seventeen|Eighteen|Nineteen|Twenty"
)
# Deliberately extends the canonical core/html_processor._HEADING_PREFIX_RE:
# the generator also strips spelled-out chapter words ("Chapter One") and
# bare letter prefixes without dots ("A Title") for its style previews.
_HEADING_PREFIX_RE = re.compile(
    r"^\s*(?:"
    r"(?i:(?:Chapter|Section))\s+(?:[IVXLCDM\d]+|(?:" + _CHAPTER_WORDS + r"))(?:\.[A-Z\d]+)*(?:\s*[:.–—\-])?\s+|"
    r"(?:\d+|[IVXLCDM]{1,6}|[A-Z]{1,3}|[a-z]{1,3})(?:[.\s]+(?:\d+|[IVXLCDM]{1,6}|[A-Z]{1,3}|[a-z]{1,3})){1,5}\.?\s+(?:[:.–—\-]\s*)?|"
    r"(?i:[IVXLCDM]+)\.(?:[A-Z]{1,3}|[a-z]{1,3})(?:\.\d+){0,3}\.?\s+(?:[:.–—\-]\s*)?|"
    r"(?i:[IVXLCDM]+)(?:\.\d+){0,3}\.?\s+(?:[:.–—\-]\s*)?|"
    r"(?:[A-Z]{1,3}|[a-z]{1,3})(?:[.)])?\s+(?:[:.–—\-]\s*)?|"
    r"\d+(?:\.\d+){0,3}(?:[.)])?\s+(?:[:.–—\-]\s*)?"
    r")"
)
def _strip_heading_prefix_for_preview(text: str) -> str:
    if not text:
        return ""
    normalized = _normalized(text).lstrip()
    match = _HEADING_PREFIX_RE.match(normalized)
    if not match:
        return normalized.strip()
    return normalized[match.end():].lstrip()

def _apply_heading_list_prefix(text: str, prefix: str) -> str:
    if not prefix:
        return _strip_heading_prefix_for_preview(text)
    base_text = _strip_heading_prefix_for_preview(text)
    if not base_text:
        return prefix.strip()
    return f"{prefix} {base_text}".strip()

def _extract_style_map_tokens(text: str) -> list[str]:
    if not text:
        return []
    normalized = _normalized(text)
    match = _HEADING_PREFIX_RE.match(normalized)
    if not match:
        return []
    prefix = match.group(0)
    prefix = re.sub(r'^(?:Chapter|Section)\s+', '', prefix, flags=re.IGNORECASE).strip()
    prefix = re.sub(r'[:.\-\s]+$', '', prefix).strip()
    return [part for part in re.split(r'[.\s]+', prefix) if part]

def build_infer_maps_from_headings(headings: list[dict]) -> tuple[dict, dict]:
    style_map = {}
    sequence_map = {}
    for heading in headings or []:
        text = heading.get("text") or ""
        level = heading.get("level")
        if not text or not level:
            continue
        tokens = _extract_style_map_tokens(text) or _extract_heading_prefix_tokens(text)
        if not tokens:
            continue
        token_types = []
        for token in tokens:
            token_type = _classify_heading_token(token)
            if not token_type:
                token_types = []
                break
            token_types.append(token_type)
        if not token_types:
            continue
        last_token_type = token_types[-1]
        if last_token_type not in style_map:
            style_map[last_token_type] = level
        seq_key = tuple(token_types)
        if seq_key not in sequence_map:
            sequence_map[seq_key] = level
    return style_map, sequence_map

def analyze_docx(docx_path):
    """Analyze DOCX file and extract styles, headings, lists, and structure."""
    doc = Document(str(docx_path))
    
    analysis = {
        "document_info": {
            "filename": docx_path.name,
            "analyzed_at": datetime.now().isoformat(),
        },
        "styles": {},
        "headings": [],
        "lists": [],
        "list_examples": [],
        "detected_styles": {
            "body_styles": [],
            "heading_styles": [],
            "list_styles": [],
        }
    }

    style_objects = {}
    # Extract all styles
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style_info = extract_style_info(style)
            analysis["styles"][style.name] = style_info
            style_objects[style.name] = style
            
            # Categorize styles
            style_lower = style.name.lower()
            if "heading" in style_lower:
                analysis["detected_styles"]["heading_styles"].append(style.name)
            elif any(x in style_lower for x in ["normal", "body", "paragraph"]):
                analysis["detected_styles"]["body_styles"].append(style.name)
            elif any(x in style_lower for x in ["list", "bullet"]):
                analysis["detected_styles"]["list_styles"].append(style.name)

    # Determine the primary body style for samples
    body_style_name = "Normal"
    if analysis["detected_styles"]["body_styles"]:
        body_style_name = analysis["detected_styles"]["body_styles"][0]
    analysis["body_style_name"] = body_style_name

    # Build resolved style map (inheritance-aware)
    resolved_styles = {}
    for style_name, style_obj in style_objects.items():
        resolved_styles[style_name] = resolve_style_info(style_obj)
    analysis["resolved_styles"] = resolved_styles

    # Keep document info minimal to avoid noisy outputs.
    
    # Extract numbering definitions
    abstract_nums, num_to_abstract = _extract_numbering_defs(doc)
    
    # Extract headings and list items
    list_counters = {}  # Track counters per numbering scheme
    list_items_by_level = {}  # Collect examples by level
    list_format_by_level = {}
    ordered_formats = {}
    unordered_formats = {}
    body_sample = None
    heading_samples = {}
    heading_number_prefixes = {}
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        if body_sample is None and paragraph.style.name == body_style_name:
            body_sample = {
                "index": i,
                "style_name": paragraph.style.name,
                "text": text[:240],
            }
        
        # Check for heading
        level = is_heading_style(paragraph)
        if level:
            level_key = str(level)
            if level_key not in heading_samples:
                heading_samples[level_key] = {
                    "index": i,
                    "style_name": paragraph.style.name,
                    "text": text[:240],
                }
            # Extract style info for this heading
            style_info = analysis["styles"].get(paragraph.style.name, {})
            analysis["headings"].append({
                "index": i,
                "level": level,
                "text": text,
                "style_name": paragraph.style.name,
                "font_name": style_info.get("font_name"),
                "font_size": style_info.get("font_size"),
                "font_size_pt": style_info.get("font_size_pt"),
                "color": style_info.get("color"),
                "bold": style_info.get("bold", False),
            })
        
        # Check for list item
        numpr, ilvl_override = _get_paragraph_numpr(paragraph)
        if numpr is not None:
            try:
                num_id = numpr.numId.val if numpr.numId is not None else None
                if ilvl_override is not None:
                    ilvl = ilvl_override
                else:
                    ilvl = numpr.ilvl.val if numpr.ilvl is not None else None
                
                if num_id is not None and ilvl is not None:
                    abs_id = num_to_abstract.get(str(num_id))
                    if abs_id and abs_id in abstract_nums:
                        levels = abstract_nums[abs_id]
                        if ilvl not in levels and levels:
                            ilvl = min(levels.keys())
                        if ilvl in levels:
                            lvl_info = levels[ilvl]
                            
                            # Initialize counters for this numbering scheme
                            scheme_key = f"{abs_id}"
                            if scheme_key not in list_counters:
                                list_counters[scheme_key] = [0] * 10
                            
                            counters = list_counters[scheme_key]
                            
                            # Increment current level counter
                            counters[ilvl] = counters[ilvl] + 1 if counters[ilvl] else 1
                            
                            # Reset deeper levels
                            for j in range(ilvl + 1, 10):
                                counters[j] = 0
                            
                            # Build the actual list number string
                            lvl_text_template = lvl_info.get("lvlText", f"%{ilvl + 1}.")
                            list_number = ""
                            try:
                                # Replace %1, %2, etc. with actual counter values
                                def repl(match):
                                    idx = int(match.group(1)) - 1
                                    if 0 <= idx < len(counters) and counters[idx] is not None and counters[idx] > 0:
                                        fmt = levels.get(idx, {}).get("numFmt", "decimal")
                                        return _format_list_number(counters[idx], fmt, idx)
                                    return ""
                                list_number = re.sub(r"%(\d+)", repl, lvl_text_template).strip()
                                # Remove any remaining % markers (they become empty strings) - this handles cases where
                                # a template references a level that doesn't exist or hasn't been counted yet
                                list_number = re.sub(r'%\d+', '', list_number).strip()
                                # Clean up any trailing separators that might be left
                                list_number = re.sub(r'[\.\s]+$', '', list_number)
                                # If we still don't have a number, use the counter directly with proper formatting
                                if not list_number and counters[ilvl] and counters[ilvl] > 0:
                                    fmt = lvl_info.get("numFmt", "decimal")
                                    list_number = _format_list_number(counters[ilvl], fmt, ilvl)
                                    # Add a period if the template ends with one or if it's a standard format
                                    if not lvl_text_template.endswith('%') and '.' not in list_number:
                                        list_number += '.'
                            except Exception as e:
                                # Fallback to simple counter
                                if counters[ilvl]:
                                    fmt = lvl_info.get("numFmt", "decimal")
                                    list_number = _format_list_number(counters[ilvl], fmt, ilvl)
                                else:
                                    list_number = ""
                            
                            # Collect examples (first 3 items per level)
                            if ilvl not in list_items_by_level:
                                list_items_by_level[ilvl] = []
                            if len(list_items_by_level[ilvl]) < 3:
                                list_items_by_level[ilvl].append({
                                    "text": text[:150],  # First 150 chars
                                    "level": ilvl,
                                    "format": lvl_info.get("numFmt", "decimal"),
                                    "lvl_text": lvl_info.get("lvlText", ""),
                                    "counter": counters[ilvl],
                                    "number_display": list_number,  # The actual formatted number
                                })
                            
                            # Store list numbering info
                            fmt_value = (lvl_info.get("numFmt") or "decimal")
                            if fmt_value == "bullet":
                                if str(ilvl) not in unordered_formats:
                                    unordered_formats[str(ilvl)] = {
                                        "type": infer_bullet_type(lvl_info.get("lvlText", "")),
                                        "lvl_text": lvl_info.get("lvlText", ""),
                                    }
                            else:
                                if str(ilvl) not in ordered_formats:
                                    ordered_formats[str(ilvl)] = {
                                        "format": fmt_value,
                                        "lvl_text": lvl_info.get("lvlText", ""),
                                        "start": lvl_info.get("start", 1),
                                    }
                            if ilvl not in list_format_by_level:
                                analysis["lists"].append({
                                    "index": i,
                                    "level": ilvl,
                                    "format": fmt_value,
                                    "lvl_text": lvl_info.get("lvlText", ""),
                                    "start": lvl_info.get("start", 1),
                                    "text_sample": text[:80],
                                })
                                list_format_by_level[ilvl] = True

                            # Capture numbering for heading previews when numbering is stored in list metadata.
                            if level and list_number:
                                heading_number_prefixes[i] = list_number
            except Exception as e:
                # Skip if we can't extract list info
                pass
    
    # Store list examples by level
    for ilvl in sorted(list_items_by_level.keys()):
        analysis["list_examples"].extend(list_items_by_level[ilvl][:3])
    
    analysis["style_samples"] = {
        "body": body_sample,
        "headings": heading_samples,
    }

    # Prefix heading samples and heading list with detected list numbering.
    if heading_number_prefixes:
        for heading in analysis["headings"]:
            idx = heading.get("index")
            prefix = heading_number_prefixes.get(idx)
            if not prefix:
                continue
            text = heading.get("text") or ""
            heading["text"] = _apply_heading_list_prefix(text, prefix)
        for level_key, sample in heading_samples.items():
            idx = sample.get("index")
            prefix = heading_number_prefixes.get(idx)
            if not prefix:
                continue
            text = sample.get("text") or ""
            sample["text"] = _apply_heading_list_prefix(text, prefix)

    # Build resolved style samples for display
    resolved_samples = []
    if body_sample:
        resolved = resolved_styles.get(body_sample["style_name"], {})
        resolved_samples.append({
            "label": f"Body ({body_sample['style_name']})",
            "style_name": body_sample["style_name"],
            "text": body_sample["text"],
            "font_name": resolved.get("font_name"),
            "font_size": resolved.get("font_size"),
            "font_size_pt": resolved.get("font_size_pt"),
            "color": resolved.get("color"),
            "bold": resolved.get("bold"),
            "italic": resolved.get("italic"),
            "underline": resolved.get("underline"),
            "paragraph_format": resolved.get("paragraph_format", {}),
            "style_chain": resolved.get("style_chain", []),
        })
    for level_key in sorted(heading_samples.keys(), key=lambda x: int(x)):
        sample = heading_samples[level_key]
        resolved = resolved_styles.get(sample["style_name"], {})
        resolved_samples.append({
            "label": f"H{level_key} ({sample['style_name']})",
            "style_name": sample["style_name"],
            "text": sample["text"],
            "font_name": resolved.get("font_name"),
            "font_size": resolved.get("font_size"),
            "font_size_pt": resolved.get("font_size_pt"),
            "color": resolved.get("color"),
            "bold": resolved.get("bold"),
            "italic": resolved.get("italic"),
            "underline": resolved.get("underline"),
            "paragraph_format": resolved.get("paragraph_format", {}),
            "style_chain": resolved.get("style_chain", []),
        })
    analysis["resolved_style_samples"] = resolved_samples

    analysis["list_formats"] = {
        "ordered": ordered_formats,
        "unordered": unordered_formats,
    }

    # Accessibility preflight: heading order checks
    heading_order_issues = []
    prev_level = None
    for heading in analysis["headings"]:
        level = heading.get("level")
        if prev_level is None:
            if level and level != 1:
                heading_order_issues.append({
                    "index": heading.get("index"),
                    "text": heading.get("text"),
                    "level": level,
                    "previous_level": prev_level,
                    "issue": f"First heading is H{level} (expected H1).",
                })
        else:
            if level and prev_level and level > prev_level + 1:
                heading_order_issues.append({
                    "index": heading.get("index"),
                    "text": heading.get("text"),
                    "level": level,
                    "previous_level": prev_level,
                    "issue": f"Heading level skipped from H{prev_level} to H{level}.",
                })
        prev_level = level
    if analysis["headings"] and not any(h.get("level") == 1 for h in analysis["headings"]):
        heading_order_issues.append({
            "index": None,
            "text": None,
            "level": None,
            "previous_level": None,
            "issue": "No H1 headings detected.",
        })
    analysis["accessibility"] = {
        "heading_order_issues": heading_order_issues,
        "heading_levels_used": sorted({h.get("level") for h in analysis["headings"] if h.get("level")}),
    }

    style_map, sequence_map = build_infer_maps_from_headings(analysis["headings"])
    analysis["infer_style_map"] = style_map
    analysis["infer_sequence_map"] = serialize_sequence_map(sequence_map)
    levels_used = analysis["accessibility"]["heading_levels_used"]
    analysis["document_info"]["toc_depth_inferred"] = max(levels_used) if levels_used else 2

    return analysis

def create_default_config(analysis):
    """Create default JSON configuration from analysis."""
    
    # Determine manual type (guess from heading structure)
    manual_type = "chapter"  # default
    
    # Get first body style for defaults
    body_style_name = analysis.get("body_style_name", "Normal")
    resolved_styles = analysis.get("resolved_styles") or {}
    body_style = resolved_styles.get(body_style_name) or analysis["styles"].get(body_style_name, {})
    
    # Ensure body style has all required fields with defaults
    body_font = body_style.get("font_name") or "Calibri"
    body_size = body_style.get("font_size_pt") or body_style.get("font_size") or 11
    body_color = body_style.get("color") or "#000000"
    body_paragraph_format = body_style.get("paragraph_format", {}) if body_style else {}
    # Ensure body_size is a number
    try:
        body_size = float(body_size)
    except (TypeError, ValueError):
        body_size = 11.0
    body_line_height = body_paragraph_format.get("line_spacing")
    try:
        body_line_height = float(body_line_height)
    except (TypeError, ValueError):
        body_line_height = 1.15
    
    # Get heading styles
    heading_styles = {}
    for i in range(1, 7):
        heading_name = f"Heading {i}"
        resolved_heading = resolved_styles.get(heading_name)
        if resolved_heading:
            heading_size_pt = resolved_heading.get("font_size_pt") or resolved_heading.get("font_size") or (body_size * (1.5 - (i-1)*0.1))
            heading_styles[f"h{i}"] = {
                "name": heading_name,
                "font_name": resolved_heading.get("font_name") or body_font,
                "font_size": heading_size_pt,
                "font_size_pt": heading_size_pt,
                "color": resolved_heading.get("color") or body_color,
                "bold": resolved_heading.get("bold") if resolved_heading.get("bold") is not None else True,
                "italic": resolved_heading.get("italic") if resolved_heading.get("italic") is not None else False,
                "underline": resolved_heading.get("underline") if resolved_heading.get("underline") is not None else False,
                "paragraph_format": resolved_heading.get("paragraph_format", {}),
                "font_inherit": False if i == 1 else True,
            }
        elif heading_name in analysis["styles"]:
            raw_heading = analysis["styles"][heading_name]
            raw_size_pt = raw_heading.get("font_size_pt")
            if raw_size_pt is None and raw_heading.get("font_size"):
                try:
                    raw_size_pt = float(raw_heading.get("font_size")) / 1.33
                except (TypeError, ValueError):
                    raw_size_pt = None
            heading_styles[f"h{i}"] = {
                "name": heading_name,
                "font_name": raw_heading.get("font_name") or body_font,
                "font_size": raw_size_pt or (body_size * (1.5 - (i-1)*0.1)),
                "font_size_pt": raw_size_pt or (body_size * (1.5 - (i-1)*0.1)),
                "color": raw_heading.get("color") or body_color,
                "bold": raw_heading.get("bold", True),
                "italic": raw_heading.get("italic", False),
                "underline": raw_heading.get("underline", False),
                "paragraph_format": raw_heading.get("paragraph_format", {}),
                "font_inherit": False if i == 1 else True,
            }
        else:
            # Create default heading style
            heading_styles[f"h{i}"] = {
                "font_name": body_font,
                "font_size": body_size * (1.5 - (i-1)*0.1),
                "font_size_pt": body_size * (1.5 - (i-1)*0.1),
                "color": body_color,
                "bold": True,
                "font_inherit": False if i == 1 else True,
            }
    
    toc_depth_inferred = analysis.get("document_info", {}).get("toc_depth_inferred") or 3
    try:
        toc_depth_inferred = int(toc_depth_inferred)
    except (TypeError, ValueError):
        toc_depth_inferred = 3
    toc_depth_inferred = max(1, min(5, toc_depth_inferred))

    config = {
        "version": "1.0",
        "document_info": analysis["document_info"],
        "conversion": {
            "mapping_mode": "map_new",
            "preserve_numbers": False,
            "toc_depth": toc_depth_inferred,
            "manual_type": manual_type,
        },
        "styles": {
            "body": {
                "font": body_font,
                "size": float(body_size),
                "color": body_color,
                "line_height": body_line_height,
                "paragraph_format": body_paragraph_format,
            },
            "headings": heading_styles,
        },
        "numbering": {
            "mode": "css-counters",
            "scheme": manual_type,
            "levels": {}
        },
        "theme": {
            "theme_id": "manual",
            "body_font": f'"{body_font}", sans-serif',
            "body_size": float(body_size),
            "text_color": body_color,
            "heading_color": heading_styles.get("h1", {}).get("color") or "#981E32",
            "subheading_color": heading_styles.get("h2", {}).get("color") or "#4D4D4D",
            "h1_size": 2.0,
            "h2_size": 1.5,
            "h3_size": 1.25,
            "h4_size": 1.1,
            "h5_size": 1.0,
            "h6_size": 0.95,
        },
        "headings_structure": analysis["headings"],  # All headings for reference
        "heading_examples": [],  # Will be populated with most complex heading
        "lists": {
            "multilevel_formats": {},  # Will be populated from analysis
            "unordered_formats": {},
            "examples": analysis.get("list_examples", []),
        }
    }
    config["infer_style_map"] = analysis.get("infer_style_map", {}) or {}
    infer_sequence_raw = analysis.get("infer_sequence_map", {}) or {}
    if isinstance(infer_sequence_raw, list):
        config["infer_sequence_map"] = infer_sequence_raw
    else:
        config["infer_sequence_map"] = serialize_sequence_map(infer_sequence_raw)
    
    # Add list format information
    list_formats = analysis.get("list_formats", {})
    ordered_formats = list_formats.get("ordered") or {}
    unordered_formats = list_formats.get("unordered") or {}
    if ordered_formats:
        config["lists"]["multilevel_formats"] = ensure_ordered_format_levels(ordered_formats)
    elif analysis.get("lists"):
        for list_info in analysis["lists"][:10]:  # First 10 different list types
            level = list_info.get("level", 0)
            if level not in config["lists"]["multilevel_formats"]:
                config["lists"]["multilevel_formats"][str(level)] = {
                    "format": list_info.get("format", "decimal"),
                    "lvl_text": list_info.get("lvl_text", ""),
                    "start": list_info.get("start", 1),
                }
        config["lists"]["multilevel_formats"] = ensure_ordered_format_levels(config["lists"]["multilevel_formats"])
    if unordered_formats:
        normalized_unordered = {}
        for level_key, bullet_info in unordered_formats.items():
            if isinstance(bullet_info, dict):
                normalized_unordered[level_key] = bullet_info.get("type") or infer_bullet_type(bullet_info.get("lvl_text", ""))
            else:
                normalized_unordered[level_key] = bullet_info
        config["lists"]["unordered_formats"] = ensure_unordered_levels(normalized_unordered)
    else:
        config["lists"]["unordered_formats"] = ensure_unordered_levels({
            "0": "disc",
            "1": "circle",
            "2": "square",
        })
    
    # Find most complex heading (one with most styling attributes or longest text)
    if analysis["headings"]:
        most_complex = max(analysis["headings"], key=lambda h: (
            len(h.get("text", "")),
            h.get("level", 0),
            bool(h.get("font_name")),
            bool(h.get("color")),
        ))
        config["heading_examples"] = [most_complex]
    else:
        config["heading_examples"] = []
    
    return config

# ------------------------------ Flask Routes ------------------------------

@app.route("/", methods=["GET"])
def index():
    """Main upload page."""
    return render_template_string(UPLOAD_TEMPLATE)

@app.route("/upload", methods=["POST"])
def upload():
    """Handle DOCX upload and analyze."""
    if "docx" not in request.files:
        flash("No file uploaded.")
        return redirect(url_for("index"))
    
    file = request.files["docx"]
    if file.filename == "":
        flash("No file selected.")
        return redirect(url_for("index"))
    
    if not file.filename.lower().endswith(".docx"):
        flash("Please upload a .docx file.")
        return redirect(url_for("index"))
    
    # Save uploaded file
    session_id = str(uuid.uuid4())
    upload_path = PERSIST_DIR / f"{session_id}_upload.docx"
    file.save(str(upload_path))
    
    try:
        # Analyze DOCX
        analysis = analyze_docx(upload_path)
        
        # Create default config
        config = create_default_config(analysis)
        
        # Save analysis and config
        analysis_path = PERSIST_DIR / f"{session_id}_analysis.json"
        config_path = PERSIST_DIR / f"{session_id}_config.json"
        
        analysis_path.write_text(json.dumps(analysis, indent=2, default=str), encoding="utf-8")
        config_path.write_text(json.dumps(config, indent=2, default=str), encoding="utf-8")
        
        # Redirect to editor
        return redirect(url_for("editor", session_id=session_id))
    
    except Exception as e:
        flash(f"Error analyzing document: {str(e)}")
        return redirect(url_for("index"))

def empty_analysis() -> dict:
    """The analysis shape for a session with no source document.

    Keys the editor template reads must exist even when there was nothing to
    analyse; the template treats a missing key as a hard error, not a blank.
    """
    return {
        "headings": [],
        "style_samples": {},
        "list_formats": {},
        "resolved_styles": {},
        "body_style_name": "Normal",
        "infer_style_map": {},
        "infer_sequence_map": [],
        "document_info": {},
        "accessibility": {"heading_order_issues": [], "heading_levels_used": []},
    }


@app.route("/import-config", methods=["POST"])
def import_config():  # noqa: D401
    """Import a JSON configuration file."""
    if "config_json" not in request.files:
        flash("No configuration file uploaded.")
        return redirect(url_for("index"))

    file = request.files["config_json"]
    if file.filename == "":
        flash("No configuration file selected.")
        return redirect(url_for("index"))

    if not file.filename.lower().endswith(".json"):
        flash("Please upload a .json configuration file.")
        return redirect(url_for("index"))

    try:
        config = json.loads(file.read().decode("utf-8"))
    except Exception as e:
        flash(f"Invalid JSON file: {e}")
        return redirect(url_for("index"))

    session_id = str(uuid.uuid4())
    config.setdefault("styles", {})
    config["styles"].setdefault("body", {
        "font": "Calibri",
        "size": 11,
        "color": "#000000",
        "line_height": 1.15,
    })
    config["styles"].setdefault("headings", {})
    config.setdefault("conversion", {})
    config["conversion"].setdefault("manual_type", "chapter")
    config["conversion"].setdefault("mapping_mode", "map_new")
    config.setdefault("numbering", {})
    config["numbering"].setdefault("scheme", config["conversion"]["manual_type"])
    config["numbering"].setdefault("mode", "css-counters")
    config["numbering"].setdefault("levels", {})
    config.setdefault("lists", {})
    config["lists"].setdefault("multilevel_formats", {})
    config["lists"].setdefault("unordered_formats", {
        "0": "disc",
        "1": "circle",
        "2": "square",
    })
    config.setdefault("document_info", {
        "filename": file.filename,
        "analyzed_at": datetime.now().isoformat(),
    })

    config_path = PERSIST_DIR / f"{session_id}_config.json"
    analysis_path = PERSIST_DIR / f"{session_id}_analysis.json"
    config_path.write_text(json.dumps(config, indent=2, default=str), encoding="utf-8")
    # An imported config carries no source document, so there is nothing to
    # analyse — but the editor reads this file and expects the shape that
    # analyse_docx() produces. Writing a bare {} accepted the upload and then
    # failed with a 500 on the very next request, which is the worst of both.
    analysis_path.write_text(
        json.dumps(empty_analysis(), indent=2, default=str), encoding="utf-8"
    )

    flash("Configuration imported successfully.")
    return redirect(url_for("editor", session_id=session_id))

@app.route("/editor/<session_id>", methods=["GET", "POST"])
def editor(session_id):
    """Editor interface for configuring styles and headings."""

    config_path, analysis_path = _session_paths(session_id)
    
    if not config_path.exists():
        flash("Session expired or invalid.")
        return redirect(url_for("index"))
    
    config = json.loads(config_path.read_text(encoding="utf-8"))

    if request.method == "POST":
        # Update config from form data
        try:
            config.setdefault("numbering", {})
            
            # Update styles from form
            if "body_font" in request.form:
                config["styles"]["body"]["font"] = request.form["body_font"]
            if "body_size" in request.form:
                config["styles"]["body"]["size"] = round_half(request.form["body_size"])
            if "body_color" in request.form:
                config["styles"]["body"]["color"] = request.form["body_color"]
            if "line_height" in request.form:
                config["styles"]["body"]["line_height"] = float(request.form["line_height"])
            
            # Update heading styles
            previous_h1_font = (config.get("styles", {}).get("headings", {}).get("h1", {}) or {}).get("font_name")
            if not previous_h1_font:
                previous_h1_font = config.get("styles", {}).get("body", {}).get("font")
            for i in range(1, 7):
                h_key = f"h{i}"
                inherit_key = f"{h_key}_inherit_font"
                if f"{h_key}_font" in request.form:
                    if h_key not in config["styles"]["headings"]:
                        config["styles"]["headings"][h_key] = {}
                    if i == 1:
                        config["styles"]["headings"][h_key]["font_name"] = request.form[f"{h_key}_font"]
                        config["styles"]["headings"][h_key]["font_inherit"] = False
                    else:
                        inherit_font = inherit_key in request.form
                        config["styles"]["headings"][h_key]["font_inherit"] = inherit_font
                        if inherit_font:
                            config["styles"]["headings"][h_key]["font_name"] = request.form.get("h1_font") or previous_h1_font
                        else:
                            config["styles"]["headings"][h_key]["font_name"] = request.form[f"{h_key}_font"]
                if f"{h_key}_size" in request.form:
                    if h_key not in config["styles"]["headings"]:
                        config["styles"]["headings"][h_key] = {}
                    config["styles"]["headings"][h_key]["font_size"] = round_half(request.form[f"{h_key}_size"])
                if f"{h_key}_color" in request.form:
                    if h_key not in config["styles"]["headings"]:
                        config["styles"]["headings"][h_key] = {}
                    config["styles"]["headings"][h_key]["color"] = request.form[f"{h_key}_color"]
                if h_key not in config["styles"]["headings"]:
                    config["styles"]["headings"][h_key] = {}
                config["styles"]["headings"][h_key]["italic"] = bool(request.form.get(f"{h_key}_italic"))
            
            # Update heading structure if provided
            if "headings_structure" in request.form:
                try:
                    headings_data = json.loads(request.form["headings_structure"])
                    config["headings_structure"] = headings_data
                except:
                    pass  # Ignore invalid JSON
            
            # Update list formats directly from form controls
            list_formats_data = {}
            for key, value in request.form.items():
                if not key.startswith("list_format_"):
                    continue
                level = key.replace("list_format_", "")
                template = request.form.get(f"list_template_{level}", "")
                entry = {
                    "format": value,
                    "lvl_text": template if value else "",
                    "start": 1
                }
                if not value:
                    entry["disabled"] = True
                list_formats_data[level] = entry
            if list_formats_data:
                config.setdefault("lists", {})
                config["lists"]["multilevel_formats"] = ensure_ordered_format_levels(list_formats_data)

            unordered_formats_data = {}
            for key, value in request.form.items():
                if not key.startswith("unordered_format_"):
                    continue
                level = key.replace("unordered_format_", "")
                if value:
                    unordered_formats_data[level] = {"type": value}
                else:
                    unordered_formats_data[level] = {"type": "", "disabled": True}
            if unordered_formats_data:
                config.setdefault("lists", {})
                config["lists"]["unordered_formats"] = ensure_unordered_levels(unordered_formats_data)
            
            # Update conversion settings
            if "manual_type" in request.form:
                config["conversion"]["manual_type"] = request.form["manual_type"]
                config["numbering"]["scheme"] = request.form["manual_type"]
            if "mapping_mode" in request.form:
                config["conversion"]["mapping_mode"] = request.form["mapping_mode"]
            
            # Save updated config when requested
            if "save_config" in request.form:
                config_path.write_text(json.dumps(config, indent=2, default=str), encoding="utf-8")
                flash("Configuration saved successfully.")
            else:
                flash("Preview refreshed.")
            
        except Exception as e:
            flash(f"Error saving configuration: {str(e)}")

    # Load analysis
    analysis = json.loads(analysis_path.read_text(encoding="utf-8")) if analysis_path.exists() else {}
    
    preview = build_preview_data(config, analysis)

    return render_template_string(
        EDITOR_TEMPLATE,
        session_id=session_id,
        config=config,
        analysis=analysis,
        preview=preview
    )

@app.route("/export/<session_id>", methods=["GET"])
def export_config(session_id):
    """Export JSON configuration file."""
    config_path, _analysis_path = _session_paths(session_id)
    
    if not config_path.exists():
        flash("Configuration not found.")
        return redirect(url_for("index"))
    
    config = json.loads(config_path.read_text(encoding="utf-8"))
    filename = config.get("document_info", {}).get("filename", "manual")
    filename_base = filename.replace(".docx", "").replace(".DOCX", "")
    export_filename = f"{filename_base}_config.json"
    
    return send_file(
        str(config_path),
        as_attachment=True,
        download_name=export_filename,
        mimetype="application/json"
    )

@app.route("/example/<session_id>", methods=["GET"])
def export_example_doc(session_id):
    """Export an example DOCX showcasing resolved styles."""
    config_path, _analysis_path = _session_paths(session_id)

    if not config_path.exists():
        flash("Configuration not found.")
        return redirect(url_for("index"))

    config = json.loads(config_path.read_text(encoding="utf-8"))
    example_path = PERSIST_DIR / f"{session_id}_example.docx"
    build_example_docx(config, example_path)

    filename = config.get("document_info", {}).get("filename", "manual")
    filename_base = filename.replace(".docx", "").replace(".DOCX", "")
    export_filename = f"{filename_base}_example.docx"

    return send_file(
        str(example_path),
        as_attachment=True,
        download_name=export_filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ------------------------------ HTML Templates ------------------------------

UPLOAD_TEMPLATE = """
<!doctype html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>DOCX Configuration Generator</title>
    <style>
        :root { --brand: #981E32; --ring: #d9d9d9; }
        * { box-sizing: border-box; }
        body {
            margin: 0;
            font-family: system-ui, -apple-system, "Segoe UI", Roboto, Arial, sans-serif;
            background: #f7f7f8;
            color: #111;
        }
        .header {
            background: var(--brand);
            color: #fff;
            padding: 14px 16px;
            font-weight: 700;
        }
        .main {
            max-width: 1200px;
            margin: 24px auto;
            padding: 0 16px;
        }
        .card {
            background: #fff;
            border: 1px solid var(--ring);
            border-radius: 12px;
            padding: 18px;
            margin-bottom: 16px;
        }
        h1 {
            margin: 0 0 1rem 0;
            font-size: 1.5rem;
        }
        label {
            font-weight: 600;
            display: block;
            margin: 0.5rem 0;
        }
        input[type="file"] {
            display: block;
            margin: 8px 0 16px;
            padding: 8px;
            border: 1px solid var(--ring);
            border-radius: 6px;
            width: 100%;
            max-width: 400px;
        }
        button {
            background: var(--brand);
            color: #fff;
            border: 0;
            border-radius: 10px;
            padding: 12px 20px;
            cursor: pointer;
            font-size: 16px;
            font-weight: 600;
        }
        button:hover {
            background: #7a1728;
        }
        .msg {
            padding: 12px 16px;
            background: #fff7ed;
            border: 1px solid #fed7aa;
            border-radius: 8px;
            color: #7c2d12;
            margin-bottom: 16px;
        }
        .description {
            color: #666;
            line-height: 1.6;
            margin-top: 16px;
        }
    </style>
</head>
<body>
    <div class="header">DOCX Configuration Generator</div>
    <div class="main">
        {% with messages = get_flashed_messages() %}
            {% if messages %}
                {% for m in messages %}
                    <div class="msg">{{ m }}</div>
                {% endfor %}
            {% endif %}
        {% endwith %}
        
        <div class="card">
            <h1>Upload DOCX File</h1>
            <form method="POST" enctype="multipart/form-data" action="{{ url_for('upload') }}">
                <label for="docx">Select .docx file to analyze:</label>
                <input id="docx" type="file" name="docx" accept=".docx" required>
                <button type="submit">Analyze & Configure</button>
            </form>
            <div class="description">
                <p><strong>What this tool does:</strong></p>
                <ul>
                    <li>Analyzes your DOCX file to extract styles, fonts, colors, and heading structure</li>
                    <li>Creates a default configuration based on detected styles</li>
                    <li>Allows you to edit DOCX export styles and list formats</li>
                    <li>Shows a live preview panel for the exported DOCX look</li>
                    <li>Exports a JSON configuration file for use with the main conversion app</li>
                </ul>
            </div>
        </div>

        <div class="card">
            <h1>Import JSON Configuration</h1>
            <form method="POST" enctype="multipart/form-data" action="{{ url_for('import_config') }}">
                <label for="config_json">Select .json configuration file:</label>
                <input id="config_json" type="file" name="config_json" accept=".json" required>
                <button type="submit">Import Configuration</button>
            </form>
            <div class="description">
                <p>Use this to resume work from a previously exported configuration.</p>
            </div>
        </div>
    </div>
</body>
</html>
"""

EDITOR_TEMPLATE = """
<!doctype html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Configuration Editor</title>
    <style>
        :root { --brand: #981E32; --ring: #d9d9d9; }
        * { box-sizing: border-box; }
        body {
            margin: 0;
            font-family: system-ui, -apple-system, "Segoe UI", Roboto, Arial, sans-serif;
            background: #f7f7f8;
            color: #111;
        }
        .header {
            background: var(--brand);
            color: #fff;
            padding: 14px 16px;
            font-weight: 700;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        .header a {
            color: #fff;
            text-decoration: none;
            font-weight: 400;
            font-size: 14px;
        }
        .main {
            max-width: 1400px;
            margin: 24px auto;
            padding: 0 16px;
        }
        .editor-layout {
            display: grid;
            grid-template-columns: 360px minmax(0, 1fr);
            gap: 16px;
            align-items: start;
        }
        .editor-sidebar {
            position: sticky;
            top: 16px;
            align-self: start;
            max-height: calc(100vh - 32px);
            overflow: auto;
            padding-right: 4px;
        }
        .editor-preview {
            min-width: 0;
        }
        .card {
            background: #fff;
            border: 1px solid var(--ring);
            border-radius: 12px;
            padding: 24px;
            margin-bottom: 16px;
        }
        h1 { margin: 0 0 1.5rem 0; font-size: 1.5rem; }
        h2 { margin: 1.5rem 0 1rem 0; font-size: 1.25rem; border-bottom: 2px solid var(--ring); padding-bottom: 8px; }
        h3 { margin: 1rem 0 0.5rem 0; font-size: 1.1rem; }
        .form-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 12px;
            margin: 16px 0;
        }
        .form-group {
            display: flex;
            flex-direction: column;
        }
        label {
            font-weight: 600;
            display: block;
            margin-bottom: 4px;
            font-size: 13px;
        }
        input[type="text"],
        input[type="number"],
        select {
            padding: 6px 10px;
            border: 1px solid var(--ring);
            border-radius: 6px;
            font-size: 13px;
        }
        input[type="color"] {
            width: 100%;
            height: 34px;
            border: 1px solid var(--ring);
            border-radius: 6px;
            cursor: pointer;
        }
        button {
            background: var(--brand);
            color: #fff;
            border: 0;
            border-radius: 8px;
            padding: 10px 18px;
            cursor: pointer;
            font-size: 14px;
            font-weight: 600;
        }
        button:hover {
            background: #7a1728;
        }
        .btn-secondary {
            background: #6b7280;
        }
        .btn-secondary:hover {
            background: #4b5563;
        }
        .btn-export {
            background: #0f766e;
        }
        .btn-export:hover {
            background: #0b5d5a;
        }
        .btn-link {
            display: inline-block;
            padding: 10px 18px;
            border-radius: 8px;
            color: #fff;
            text-decoration: none;
            font-size: 14px;
            font-weight: 600;
        }
        .action-bar {
            display: flex;
            flex-wrap: wrap;
            align-items: center;
            gap: 10px;
            background: #fff;
            border: 1px solid var(--ring);
            border-radius: 12px;
            padding: 10px 12px;
            margin-bottom: 12px;
        }
        .checkbox-inline {
            display: flex;
            align-items: center;
            gap: 8px;
            font-weight: 600;
            font-size: 13px;
        }
        .checkbox-inline input {
            margin: 0;
        }
        .msg {
            padding: 12px 16px;
            background: #dcfce7;
            border: 1px solid #86efac;
            border-radius: 8px;
            color: #166534;
            margin-bottom: 16px;
        }
        .heading-section {
            margin: 24px 0;
            padding: 16px;
            background: #f9fafb;
            border-radius: 8px;
        }
        .actions {
            margin-top: 24px;
            padding-top: 24px;
            border-top: 2px solid var(--ring);
        }
        .info-box {
            background: #eff6ff;
            border: 1px solid #bfdbfe;
            border-radius: 8px;
            padding: 12px;
            margin: 16px 0;
            color: #1e40af;
            font-size: 14px;
        }
        .sample-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
            gap: 16px;
        }
        .sample-card {
            border: 1px solid var(--ring);
            border-radius: 10px;
            padding: 12px;
            background: #fafafa;
        }
        .sample-title {
            font-weight: 600;
            margin-bottom: 6px;
            font-size: 14px;
        }
        .sample-chain {
            font-size: 12px;
            color: #666;
            margin-bottom: 6px;
        }
        .sample-text {
            padding: 8px;
            background: #fff;
            border: 1px dashed #ddd;
            border-radius: 6px;
            margin-bottom: 8px;
        }
        .sample-meta {
            font-size: 12px;
            color: #444;
            margin-top: 4px;
        }
        .issue-list {
            list-style: disc;
            padding-left: 20px;
            margin: 12px 0;
        }
        .issue-item {
            margin-bottom: 6px;
        }
        .status-ok {
            background: #ecfeff;
            border: 1px solid #a5f3fc;
            color: #155e75;
            padding: 8px 12px;
            border-radius: 6px;
            font-size: 13px;
        }
        .preview-header {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 12px;
        }
        .preview-section {
            margin-top: 16px;
            padding-top: 16px;
            border-top: 1px solid #e5e7eb;
        }
        .preview-label {
            font-weight: 600;
            font-size: 13px;
            color: #555;
            margin-bottom: 8px;
        }
        .preview-block {
            margin-top: 16px;
            padding: 12px;
            border-radius: 10px;
            border: 1px solid #e5e7eb;
            background: #fff;
        }
        .preview-block.active {
            border-color: #0f766e;
            background: #f0fdfa;
        }
        .preview-list {
            list-style: none;
            padding: 0;
            margin: 8px 0 0 0;
        }
        .preview-list li {
            display: flex;
            gap: 10px;
            margin-bottom: 6px;
        }
        .preview-marker {
            min-width: 56px;
            font-family: monospace;
            font-weight: 600;
            color: #333;
        }
        .note-text {
            font-size: 12px;
            color: #666;
            margin-top: 8px;
        }
        .panel summary {
            cursor: pointer;
            font-weight: 600;
            font-size: 1.05rem;
        }
        .panel[open] summary {
            margin-bottom: 12px;
        }
        .panel-body {
            padding-top: 8px;
        }
    </style>
</head>
<body>
    <div class="header">
        <span>Configuration Editor</span>
        <a href="{{ url_for('index') }}">Start Over</a>
    </div>
    <div class="main">
        {% with messages = get_flashed_messages() %}
            {% if messages %}
                {% for m in messages %}
                    <div class="msg">{{ m }}</div>
                {% endfor %}
            {% endif %}
        {% endwith %}
        
        <form id="config-form" method="POST" action="{{ url_for('editor', session_id=session_id) }}">
            <div class="action-bar">
                <button type="submit" name="save_config">Save Configuration</button>
                <a href="{{ url_for('export_config', session_id=session_id) }}" class="btn-link btn-export">Download JSON</a>
                <a href="{{ url_for('export_example_doc', session_id=session_id) }}" class="btn-link btn-secondary">Download Example DOCX</a>
                <a href="{{ url_for('editor', session_id=session_id) }}" class="btn-link btn-secondary">Reset Form</a>
            </div>
            <div class="note-text">Preview refreshes automatically when you change a setting. Save Configuration keeps changes for this session.</div>
            <input type="hidden" name="preview_only" id="preview_only" value="">
            <button type="submit" name="preview_submit" id="preview_submit" style="display:none">Preview</button>
            {% set font_options = [
                "Aptos", "Calibri", "Arial", "Times New Roman", "Cambria",
                "Georgia", "Garamond", "Palatino Linotype", "Verdana", "Tahoma",
                "Trebuchet MS"
            ] %}
            <div class="editor-layout">
                <div class="editor-sidebar">
                    <details class="card panel" data-panel="numbering" open>
                        <summary>Numbering Settings</summary>
                        <div class="panel-body">
                            <div class="form-grid">
                                <div class="form-group">
                                    <label for="manual_type">Manual Type</label>
                                    <select id="manual_type" name="manual_type">
                                        <option value="chapter" {% if config.conversion.manual_type == "chapter" %}selected{% endif %}>Chapter</option>
                                        <option value="section" {% if config.conversion.manual_type == "section" %}selected{% endif %}>Section</option>
                                    </select>
                                </div>
                                <div class="form-group">
                                    <label for="mapping_mode">Mapping Mode</label>
                                    <select id="mapping_mode" name="mapping_mode">
                                        <option value="map_new" {% if config.conversion.mapping_mode == "map_new" %}selected{% endif %}>Map to New Numeric</option>
                                        <option value="keep_old" {% if config.conversion.mapping_mode == "keep_old" %}selected{% endif %}>Keep Original</option>
                                    </select>
                                </div>
                            </div>
                            <div class="note-text">Mapping preview is shown on the right.</div>
                        </div>
                    </details>

                    <details class="card panel" data-panel="heading-order" {% if analysis.accessibility.heading_order_issues %}open{% endif %}>
                        <summary>Heading Order (Original DOCX)</summary>
                        <div class="panel-body">
                            {% if analysis.accessibility.heading_order_issues %}
                                <ul class="issue-list">
                                    {% for issue in analysis.accessibility.heading_order_issues %}
                                        <li class="issue-item">{{ issue.issue }}</li>
                                    {% endfor %}
                                </ul>
                            {% else %}
                                <div class="status-ok">No heading order issues found in the original document.</div>
                            {% endif %}
                        </div>
                    </details>

                    <details class="card panel" data-panel="body">
                        <summary>Body Text</summary>
                        <div class="panel-body">
                            <div class="form-grid" style="margin-top: 12px;">
                                <div class="form-group">
                                    <label for="body_font">Font Family</label>
                                    <select id="body_font" name="body_font">
                                        {% set body_font = config.styles.body.font if config.styles.body.font else 'Calibri' %}
                                        {% if body_font not in font_options %}
                                            <option value="{{ body_font }}" selected>{{ body_font }} (current)</option>
                                        {% endif %}
                                        {% for font in font_options %}
                                            <option value="{{ font }}" {% if body_font == font %}selected{% endif %}>{{ font }}</option>
                                        {% endfor %}
                                    </select>
                                </div>
                                <div class="form-group">
                                    <label for="body_size">Font Size (pt)</label>
                                    <input type="number" id="body_size" name="body_size" step="0.5" value="{{ config.styles.body.size if config.styles.body.size else 11 }}">
                                </div>
                                <div class="form-group">
                                    <label for="body_color">Text Color</label>
                                    <input type="color" id="body_color" name="body_color" value="{{ config.styles.body.color if config.styles.body.color else '#000000' }}">
                                </div>
                                <div class="form-group">
                                    <label for="line_height">Line Height</label>
                                    <input type="number" id="line_height" name="line_height" step="0.1" value="{{ config.styles.body.line_height if config.styles.body.line_height else 1.15 }}">
                                </div>
                            </div>
                        </div>
                    </details>

                    <details class="card panel" data-panel="headings">
                        <summary>Heading Styles</summary>
                        <div class="panel-body" style="margin-top: 12px;">
                        {% for i in range(1, 7) %}
                            {% set h_key = "h" + i|string %}
                            {% set h_style = config.styles.headings.get(h_key, {}) %}
                            {% set h1_font = config.styles.headings.get('h1', {}).get('font_name') if config.styles.headings.get('h1') else (config.styles.body.font if config.styles.body.font else 'Calibri') %}
                            {% set inherit_h1 = h_style.get('font_inherit', True) if i != 1 else False %}
                            {% set default_body_size = config.styles.body.size if config.styles.body.size else 11 %}
                            {% set default_font_size = h_style.get('font_size') if h_style.get('font_size') else (default_body_size * (1.5 - (i-1)*0.1)) %}
                            {% set default_font = h1_font if inherit_h1 else (h_style.get('font_name') if h_style.get('font_name') else (config.styles.body.font if config.styles.body.font else 'Calibri')) %}
                            {% set default_color = h_style.get('color') if h_style.get('color') else '#000000' %}
                            {% set default_italic = h_style.get('italic') if h_style.get('italic') is not none else False %}
                            <div class="heading-section">
                                <h3>Heading {{ i }} (H{{ i }})</h3>
                                <div class="form-grid">
                                    <div class="form-group">
                                        <label for="{{ h_key }}_font">Font Family</label>
                                        <select id="{{ h_key }}_font" name="{{ h_key }}_font">
                                            {% set current_font = default_font %}
                                            {% if current_font not in font_options %}
                                                <option value="{{ current_font }}" selected>{{ current_font }} (current)</option>
                                            {% endif %}
                                            {% for font in font_options %}
                                                <option value="{{ font }}" {% if current_font == font %}selected{% endif %}>{{ font }}</option>
                                            {% endfor %}
                                        </select>
                                    </div>
                                    {% if i != 1 %}
                                    <div class="form-group">
                                        <label class="checkbox-inline">
                                            <input type="checkbox" name="{{ h_key }}_inherit_font" {% if inherit_h1 %}checked{% endif %}>
                                            Inherit H1 font
                                        </label>
                                    </div>
                                    {% endif %}
                                    <div class="form-group">
                                        <label for="{{ h_key }}_size">Font Size (pt)</label>
                                        <input type="number" id="{{ h_key }}_size" name="{{ h_key }}_size" step="0.5" value="{{ default_font_size }}">
                                    </div>
                                    <div class="form-group">
                                        <label for="{{ h_key }}_color">Color</label>
                                        <input type="color" id="{{ h_key }}_color" name="{{ h_key }}_color" value="{{ default_color }}">
                                    </div>
                                    <div class="form-group">
                                        <label class="checkbox-inline">
                                            <input type="checkbox" name="{{ h_key }}_italic" {% if default_italic %}checked{% endif %}>
                                            Italic
                                        </label>
                                    </div>
                                </div>
                            </div>
                        {% endfor %}
                        </div>
                    </details>

                    <details class="card panel" data-panel="lists">
                        <summary>List Formats</summary>
                        <div class="panel-body">
                            {% if config.lists.multilevel_formats %}
                            <div style="margin: 12px 0;">
                                <h3 style="font-size: 1.1rem; margin: 1rem 0 0.5rem;">Ordered List Formats</h3>
                                <table style="width: 100%; border-collapse: collapse; margin-top: 12px;">
                                    <thead>
                                        <tr style="background: #f3f4f6; border-bottom: 2px solid var(--ring);">
                                            <th style="padding: 8px; text-align: left; font-size: 12px; width: 80px;">Level</th>
                                            <th style="padding: 8px; text-align: left; font-size: 12px; width: 160px;">Format</th>
                                            <th style="padding: 8px; text-align: left; font-size: 12px;">Template</th>
                                        </tr>
                                    </thead>
                                    <tbody id="list-formats-tbody">
                                        {% for level_str, fmt_info in config.lists.multilevel_formats.items() %}
                                            <tr style="border-bottom: 1px solid #e5e7eb;" data-level="{{ level_str }}">
                                                <td style="padding: 8px; font-size: 12px; font-weight: 600;">Level {{ level_str }}</td>
                                                <td style="padding: 8px;">
                                                    <select name="list_format_{{ level_str }}" style="width: 100%; padding: 4px; font-size: 12px; font-family: monospace; border: 1px solid var(--ring); border-radius: 4px;">
                                                    <option value="" {% if not fmt_info.format %}selected{% endif %}>Not used</option>
                                                    <option value="decimal" {% if fmt_info.format == "decimal" %}selected{% endif %}>decimal (1, 2, 3)</option>
                                                    <option value="decimalZero" {% if fmt_info.format in ["decimalZero", "decimalzero"] %}selected{% endif %}>decimalZero (01, 02, 03)</option>
                                                    <option value="upperLetter" {% if fmt_info.format in ["upperLetter", "upperletter"] %}selected{% endif %}>upperLetter (A, B, C)</option>
                                                    <option value="lowerLetter" {% if fmt_info.format in ["lowerLetter", "lowerletter"] %}selected{% endif %}>lowerLetter (a, b, c)</option>
                                                    <option value="upperRoman" {% if fmt_info.format in ["upperRoman", "upperroman"] %}selected{% endif %}>upperRoman (I, II, III)</option>
                                                    <option value="lowerRoman" {% if fmt_info.format in ["lowerRoman", "lowerroman"] %}selected{% endif %}>lowerRoman (i, ii, iii)</option>
                                                    </select>
                                                </td>
                                                <td style="padding: 8px;">
                                                    <input type="text" name="list_template_{{ level_str }}" 
                                                           value="{{ fmt_info.lvl_text }}" 
                                                           style="width: 100%; padding: 4px; font-size: 12px; border: 1px solid var(--ring); border-radius: 4px; font-family: monospace;">
                                                </td>
                                            </tr>
                                        {% endfor %}
                                    </tbody>
                                </table>
                                <input type="hidden" id="list_formats_json" name="list_formats" value="">
                                <script>
                                    function updateListFormatsJSON() {
                                        const tbody = document.getElementById('list-formats-tbody');
                                        if (!tbody) return;
                                        const rows = tbody.querySelectorAll('tr[data-level]');
                                        const formats = {};
                                        rows.forEach(row => {
                                            const level = row.getAttribute('data-level');
                                            const formatSelect = row.querySelector('select[name^="list_format_"]');
                                            const templateInput = row.querySelector('input[name^="list_template_"]');
                                            if (level && formatSelect && templateInput) {
                                                formats[level] = {
                                                    format: formatSelect.value,
                                                    lvl_text: templateInput.value,
                                                    start: 1
                                                };
                                            }
                                        });
                                        document.getElementById('list_formats_json').value = JSON.stringify(formats);
                                    }
                                    document.addEventListener('DOMContentLoaded', function() {
                                        const tbody = document.getElementById('list-formats-tbody');
                                        if (tbody) {
                                            tbody.addEventListener('change', updateListFormatsJSON);
                                            tbody.addEventListener('input', updateListFormatsJSON);
                                            updateListFormatsJSON();
                                        }
                                    });
                                </script>
                            </div>
                            {% endif %}

                            {% if config.lists.unordered_formats %}
                            <div style="margin: 16px 0;">
                                <h3 style="font-size: 1.1rem; margin: 1rem 0 0.5rem;">Unordered List Bullet Styles</h3>
                                <table style="width: 100%; border-collapse: collapse; margin-top: 12px;">
                                    <thead>
                                        <tr style="background: #f3f4f6; border-bottom: 2px solid var(--ring);">
                                            <th style="padding: 8px; text-align: left; font-size: 12px; width: 80px;">Level</th>
                                            <th style="padding: 8px; text-align: left; font-size: 12px;">Bullet Type</th>
                                        </tr>
                                    </thead>
                                    <tbody id="unordered-formats-tbody">
                                        {% for level_str, bullet_type in config.lists.unordered_formats.items() %}
                                            {% if bullet_type is mapping %}
                                                {% set bullet_value = bullet_type.get('type') %}
                                            {% else %}
                                                {% set bullet_value = bullet_type %}
                                            {% endif %}
                                            <tr style="border-bottom: 1px solid #e5e7eb;" data-level="{{ level_str }}">
                                                <td style="padding: 8px; font-size: 12px; font-weight: 600;">Level {{ level_str }}</td>
                                                <td style="padding: 8px;">
                                                    <select name="unordered_format_{{ level_str }}" style="width: 100%; padding: 4px; font-size: 12px; font-family: monospace; border: 1px solid var(--ring); border-radius: 4px;">
                                                    <option value="" {% if not bullet_value %}selected{% endif %}>Not used</option>
                                                    <option value="disc" {% if bullet_value == "disc" %}selected{% endif %}>disc (&#8226;)</option>
                                                    <option value="circle" {% if bullet_value == "circle" %}selected{% endif %}>circle (&#9702;)</option>
                                                    <option value="square" {% if bullet_value == "square" %}selected{% endif %}>square (&#9642;)</option>
                                                    <option value="dash" {% if bullet_value == "dash" %}selected{% endif %}>dash (&#8211;)</option>
                                                    </select>
                                                </td>
                                            </tr>
                                        {% endfor %}
                                    </tbody>
                                </table>
                                <input type="hidden" id="unordered_formats_json" name="unordered_formats" value="">
                                <script>
                                    function updateUnorderedFormatsJSON() {
                                        const tbody = document.getElementById('unordered-formats-tbody');
                                        if (!tbody) return;
                                        const rows = tbody.querySelectorAll('tr[data-level]');
                                        const formats = {};
                                        rows.forEach(row => {
                                            const level = row.getAttribute('data-level');
                                            const formatSelect = row.querySelector('select[name^="unordered_format_"]');
                                            if (level && formatSelect) {
                                                formats[level] = formatSelect.value;
                                            }
                                        });
                                        document.getElementById('unordered_formats_json').value = JSON.stringify(formats);
                                    }
                                    document.addEventListener('DOMContentLoaded', function() {
                                        const tbody = document.getElementById('unordered-formats-tbody');
                                        if (tbody) {
                                            tbody.addEventListener('change', updateUnorderedFormatsJSON);
                                            updateUnorderedFormatsJSON();
                                        }
                                    });
                                </script>
                            </div>
                            {% endif %}
                        </div>
                    </details>
                </div>

                <div class="editor-preview">
                    <div class="card">
                        <div class="preview-header">
                            <h2>Preview</h2>
                        </div>
                        <div class="note-text">Preview updates when you change a setting. Current mapping mode: {{ "Keep Original" if preview.mapping_mode == "keep_old" else "Map to New Numeric" }}.</div>

                        <details class="preview-block" data-panel="preview-original">
                            <summary class="preview-label">Original (read-only)</summary>
                            <div class="preview-section">
                                <div class="preview-label">Body Text</div>
                                <div style="{{ preview.body.style }}">{{ preview.body.text }}</div>
                            </div>
                            <div class="preview-section">
                                <div class="preview-label">Headings (from upload)</div>
                                {% for heading in preview.heading_examples %}
                                    <div style="{{ heading.style }}">{{ heading.keep_text }}</div>
                                {% endfor %}
                            </div>
                            <div class="preview-section">
                                <div class="preview-label">Ordered Lists</div>
                                <ul class="preview-list">
                                    {% if preview.original_ordered_list %}
                                        {% for item in preview.original_ordered_list %}
                                            <li style="margin-left: {{ item.indent_pt if item.indent_pt is not none else 0 }}pt;"><span class="preview-marker">{{ item.marker }}</span><span style="{{ preview.body.style }}">{{ item.text }}</span></li>
                                        {% endfor %}
                                    {% else %}
                                        <li><span class="preview-marker">-</span><span style="{{ preview.body.style }}">No ordered lists detected.</span></li>
                                    {% endif %}
                                </ul>
                            </div>
                            <div class="preview-section">
                                <div class="preview-label">Unordered Lists</div>
                                <ul class="preview-list">
                                    {% if preview.original_unordered_list %}
                                        {% for item in preview.original_unordered_list %}
                                            <li style="margin-left: {{ item.indent_pt if item.indent_pt is not none else 0 }}pt;"><span class="preview-marker">{{ item.marker }}</span><span style="{{ preview.body.style }}">{{ item.text }}</span></li>
                                        {% endfor %}
                                    {% else %}
                                        <li><span class="preview-marker">-</span><span style="{{ preview.body.style }}">No unordered lists detected.</span></li>
                                    {% endif %}
                                </ul>
                            </div>
                        </details>
                        <div class="preview-block active">
                            <div class="preview-label">Output Preview</div>
                            <div class="note-text">Mapping mode: {{ "Keep Original" if preview.mapping_mode == "keep_old" else "Map to New Numeric" }}</div>
                            <div class="preview-section">
                                <div class="preview-label">Body Text</div>
                                <div style="{{ preview.body.style }}">{{ preview.body.text }}</div>
                            </div>
                            <div class="preview-section">
                                <div class="preview-label">Headings</div>
                                {% for heading in preview.heading_examples %}
                                    <div style="{{ heading.style }}">{{ heading.output_text }}</div>
                                {% endfor %}
                            </div>
                            <div class="preview-section">
                                <div class="preview-label">Ordered Lists</div>
                                <ul class="preview-list">
                                    {% if preview.ordered_list %}
                                        {% for item in preview.ordered_list %}
                                            <li style="margin-left: {{ item.indent_pt if item.indent_pt is not none else 0 }}pt;"><span class="preview-marker">{{ item.marker }}</span><span style="{{ preview.body.style }}">{{ item.text }}</span></li>
                                        {% endfor %}
                                    {% else %}
                                        <li><span class="preview-marker">-</span><span style="{{ preview.body.style }}">No ordered list formatting selected.</span></li>
                                    {% endif %}
                                </ul>
                            </div>
                            <div class="preview-section">
                                <div class="preview-label">Unordered Lists</div>
                                <ul class="preview-list">
                                    {% if preview.unordered_list %}
                                        {% for item in preview.unordered_list %}
                                            <li style="margin-left: {{ item.indent_pt if item.indent_pt is not none else 0 }}pt;"><span class="preview-marker">{{ item.marker }}</span><span style="{{ preview.body.style }}">{{ item.text }}</span></li>
                                        {% endfor %}
                                    {% else %}
                                        <li><span class="preview-marker">-</span><span style="{{ preview.body.style }}">No unordered list formatting selected.</span></li>
                                    {% endif %}
                                </ul>
                            </div>
                        </div>


                    </div>
                </div>
            </div>
            <script>
                document.addEventListener('DOMContentLoaded', function() {
                    const form = document.getElementById('config-form');
                    if (!form) return;
                    const previewOnly = document.getElementById('preview_only');
                    const previewSubmit = document.getElementById('preview_submit');
                    const panels = document.querySelectorAll('details.panel[data-panel]');
                    const stored = localStorage.getItem('docx-config-panels');
                    let panelState = {};
                    if (stored) {
                        try {
                            panelState = JSON.parse(stored) || {};
                        } catch (e) {
                            panelState = {};
                        }
                    }
                    panels.forEach(panel => {
                        const key = panel.dataset.panel;
                        if (key && panelState.hasOwnProperty(key)) {
                            panel.open = panelState[key];
                        }
                        panel.addEventListener('toggle', () => {
                            panelState[key] = panel.open;
                            localStorage.setItem('docx-config-panels', JSON.stringify(panelState));
                        });
                    });
                    const saveBtn = form.querySelector('button[name="save_config"]');
                    if (saveBtn && previewOnly) {
                        saveBtn.addEventListener('click', () => {
                            previewOnly.value = "";
                        });
                    }
                    let previewTimer = null;
                    function schedulePreview() {
                        if (!previewOnly) return;
                        previewOnly.value = "1";
                        if (previewTimer) {
                            clearTimeout(previewTimer);
                        }
                        previewTimer = setTimeout(() => {
                            if (previewSubmit) {
                                form.requestSubmit(previewSubmit);
                            } else {
                                form.requestSubmit();
                            }
                        }, 350);
                    }
                    const autoFields = form.querySelectorAll('select, input[type="number"], input[type="color"], input[type="text"], input[type="checkbox"]');
                    autoFields.forEach(field => {
                        field.addEventListener('change', schedulePreview);
                        if (field.type === "text") {
                            field.addEventListener('blur', schedulePreview);
                        }
                    });
                    form.addEventListener('submit', function() {
                        panels.forEach(panel => {
                            const key = panel.dataset.panel;
                            if (key) {
                                panelState[key] = panel.open;
                            }
                        });
                        localStorage.setItem('docx-config-panels', JSON.stringify(panelState));
                        if (typeof updateListFormatsJSON === 'function') {
                            updateListFormatsJSON();
                        }
                        if (typeof updateUnorderedFormatsJSON === 'function') {
                            updateUnorderedFormatsJSON();
                        }
                    });
                });
            </script>
        </form>
    </div>
</body>
</html>
"""

# ------------------------------ Main Execution ------------------------------

if __name__ == "__main__":
    print("DOCX Configuration Generator starting...")
    print(f"Persist directory: {PERSIST_DIR}")
    app.run(debug=True, port=5000, host="127.0.0.1")

