import logging
import re
import hashlib
import json
import uuid
import zipfile
import io
import shutil
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Tuple, List, Dict, Optional, Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .permalinks import normalize_heading_ref, is_section_style
from .reference_linking import is_non_reference_token
from utils.helpers import (
    roman_to_int, 
    _int_to_roman, 
    _int_to_letters, 
    _format_number, 
    _token_type_from_numfmt,
    normalize_hex_color
)

logger = logging.getLogger(__name__)

# --- Basic DOCX Utilities ---

def compute_sha256(path: Path) -> str:
    """Compute SHA256 for a file."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def is_heading_style(p: Paragraph) -> int:
    """Detect if a paragraph has a Word 'Heading N' style."""
    try:
        name = p.style.name if p.style is not None else ""
    except Exception:
        name = ""
    m = re.match(r"Heading\s+([1-9])", str(name), re.IGNORECASE)
    return int(m.group(1)) if m else 0

def get_outline_level(p: Paragraph) -> int:
    """Extract outline level from paragraph properties (w:outlineLvl)."""
    pPr = p._p.pPr
    if pPr is None:
        return 0
    node = pPr.find(qn("w:outlineLvl"))
    if node is not None and node.val is not None:
        try:
            return int(node.val) + 1
        except Exception:
            return 0
    return 0

def get_style_outline_level(style) -> int:
    """Extract outline level from a style definition."""
    try:
        ppr = style._element.pPr if style is not None else None
    except Exception:
        ppr = None
    if ppr is None:
        return 0
    node = ppr.find(qn("w:outlineLvl"))
    if node is not None and node.get(qn("w:val")) is not None:
        try:
            return int(node.get(qn("w:val"))) + 1
        except Exception:
            return 0
    return 0

def get_heading_level(p: Paragraph) -> int:
    """Unified check for heading level (style first, then outline level)."""
    return is_heading_style(p) or get_outline_level(p) or get_style_outline_level(p.style)

def guess_heading_level(text: str) -> int:
    """
    Guess heading level based on text patterns.
    CRITICAL: Must be conservative - only match clear heading patterns, not body text.
    """
    s = (text or "").strip()
    
    # Only match patterns that are clearly headings, not body text
    # Patterns must be followed by whitespace and actual text, or be standalone short phrases
    patterns = [
        (r"^Section\s+[IVXLCDM]+(?:(?:\s*[:\-–—]\s*)|\s+|\.\s+)\w", 1),
        (r"^Section\s+\d+(?:(?:\s*[:\-–—]\s*)|\s+|\.\s+)\w", 1),
        (r"^Chapter\s+(?:One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve|\d+|[IVXLCDM]+)(?:(?:\s*[:\-–—]\s*)|\s+|\.\s+)\w", 1),
        (r"^[IVXLCDM]+\.[A-Z]\.\s+\w", 2),
        (r"^[IVXLCDM]+\.[A-Z]\.\d+(\.\d+){0,2}\s+\w", 3),
        # Decimal manuals: 1.1 → H2, 1.1.1 → H3 (was off-by-one / Roman-era quirky)
        (r"^\d+\.\d+\.\d+(?:\.\d+)*\s+\w", 3),
        (r"^\d+\.\d+\s+\w", 2),
        (r"^[A-Z]\.\s+\w", 3),
    ]
    
    # Check for standalone short headings (all caps)
    if len(s) <= 60 and s.upper() == s and re.search(r"[A-Z]", s) and not re.search(r"[a-z]", s):
        if len(s) <= 4:
            return 0
        if re.match(r"^[A-Z\s]{2,60}$", s) and (len(s.split()) >= 2 or len(s) >= 6):
            return 2
    
    for pat, lvl in patterns:
        if re.match(pat, s):
            return lvl
    
    return 0

# --- Numbering & Sequence Mapping ---

def _extract_numbering_defs(doc: Document) -> tuple[dict, dict]:
    """
    Return (abstract_nums, num_to_abstract) from numbering.xml.
    abstract_nums: {abstractNumId: {ilvl: {"numFmt": str, "lvlText": str, "start": int}}}
    num_to_abstract: {numId: abstractNumId}
    """
    numbering_part = getattr(doc.part, "numbering_part", None)
    if numbering_part is None:
        return {}, {}

    try:
        root = numbering_part.element
    except Exception:
        return {}, {}

    abstract_nums = {}
    num_to_abstract = {}

    for abstract in root.findall(qn("w:abstractNum")):
        abs_id = abstract.get(qn("w:abstractNumId"))
        if abs_id is None:
            continue
        levels = {}
        for lvl in abstract.findall(qn("w:lvl")):
            ilvl_raw = lvl.get(qn("w:ilvl"))
            try:
                ilvl = int(ilvl_raw)
            except Exception:
                continue
            fmt_el = lvl.find(qn("w:numFmt"))
            fmt = fmt_el.get(qn("w:val")) if fmt_el is not None else "decimal"
            text_el = lvl.find(qn("w:lvlText"))
            lvl_text = text_el.get(qn("w:val")) if text_el is not None else f"%{ilvl + 1}."
            style_el = lvl.find(qn("w:pStyle"))
            p_style = style_el.get(qn("w:val")) if style_el is not None else None
            start_el = lvl.find(qn("w:start"))
            try:
                start_val = int(start_el.get(qn("w:val"))) if start_el is not None else 1
            except Exception:
                start_val = 1
            levels[ilvl] = {"numFmt": fmt, "lvlText": lvl_text, "start": start_val, "pStyle": p_style}
        abstract_nums[abs_id] = levels

    for num in root.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        abs_el = num.find(qn("w:abstractNumId"))
        abs_id = abs_el.get(qn("w:val")) if abs_el is not None else None
        if num_id is not None and abs_id is not None:
            num_to_abstract[num_id] = abs_id

    return abstract_nums, num_to_abstract

def _get_paragraph_numpr(p: Paragraph):
    """Extract numPr from a paragraph or its style."""
    ppr = p._p.pPr
    if ppr is not None and ppr.numPr is not None:
        return ppr.numPr
    try:
        style_ppr = p.style._element.pPr if p.style is not None else None
    except Exception:
        style_ppr = None
    if style_ppr is not None and style_ppr.numPr is not None:
        return style_ppr.numPr
    return None

def _numpr_val(el):
    """Safely extract the 'val' attribute from an XML element."""
    if el is None:
        return None
    val = getattr(el, "val", None)
    if val is None:
        try:
            val = el.get(qn("w:val"))
        except Exception:
            val = None
    try:
        return int(val)
    except Exception:
        return None

def _build_list_number(levels: dict, counters: list[int | None], ilvl: int) -> str:
    """Construct a list numbering string based on numbering levels and counters."""
    lvl_def = levels.get(ilvl)
    if not lvl_def:
        return ""
    lvl_text = lvl_def.get("lvlText") or ""
    if not lvl_text:
        return ""

    def repl(match):
        idx = int(match.group(1)) - 1
        if idx < 0 or idx >= len(counters):
            return ""
        value = counters[idx]
        if value is None:
            return ""
        fmt = levels.get(idx, {}).get("numFmt", "decimal")
        # Use helper for consistent formatting
        return _format_number(value, fmt)

    return re.sub(r"%(\d+)", repl, lvl_text).strip()

def _v4_token_type_from_numfmt(numfmt: str) -> str:
    """Specific token mapping for sequence maps, mimicking V4's internal logic."""
    fmt = (numfmt or "").replace(" ", "").lower()
    mapping = {
        "decimal": "decimal",
        "decimalzero": "decimal",
        "upperroman": "roman_upper",
        "lowerroman": "roman_lower",
        "upperletter": "alpha_upper",
        "lowerletter": "alpha_lower",
        "upperalpha": "alpha_upper",
        "loweralpha": "alpha_lower"
    }
    return mapping.get(fmt, "")

def _token_sequence_from_lvltext(levels: dict, ilvl: int) -> tuple[str, ...]:
    """Identify the sequence of numbering types for a given level."""
    lvl_def = levels.get(ilvl)
    if not lvl_def:
        return ()
    lvl_text = lvl_def.get("lvlText") or ""
    tokens = []
    for raw in re.findall(r"%(\d+)", lvl_text):
        try:
            idx = int(raw) - 1
        except Exception:
            continue
        fmt = levels.get(idx, {}).get("numFmt")
        token_type = _v4_token_type_from_numfmt(fmt)
        if token_type:
            tokens.append(token_type)
    if tokens:
        return tuple(tokens)
    fmt = lvl_def.get("numFmt")
    token_type = _v4_token_type_from_numfmt(fmt)
    return (token_type,) if token_type else ()

def extract_sequence_map_from_doc(doc: Document) -> dict:
    """Map numbering sequences (e.g., ('roman_upper', 'alpha_upper')) to heading levels."""
    seq_map: dict[tuple[str, ...], int] = {}
    abstract_nums, num_to_abstract = _extract_numbering_defs(doc)
    if not abstract_nums or not num_to_abstract:
        return seq_map

    for p in doc.paragraphs:
        current_level = get_heading_level(p)
        if current_level <= 0:
            continue

        numpr = _get_paragraph_numpr(p)
        if numpr is None:
            continue

        num_id = _numpr_val(numpr.numId)
        ilvl = _numpr_val(numpr.ilvl)
        if num_id is None:
            continue

        num_id_str = str(num_id)
        abs_id = num_to_abstract.get(num_id_str)
        levels = abstract_nums.get(abs_id or "")
        if not levels:
            continue

        if ilvl is None:
            style_id = None
            try:
                style_id = p.style.style_id if p.style is not None else None
            except Exception:
                style_id = None
            if style_id:
                for lvl_idx, lvl_def in levels.items():
                    if (lvl_def.get("pStyle") or "").lower() == style_id.lower():
                        ilvl = lvl_idx
                        break
        if ilvl is None:
            continue

        sequence = _token_sequence_from_lvltext(levels, ilvl)
        if sequence and sequence not in seq_map:
            seq_map[sequence] = current_level

    return seq_map

def merge_sequence_maps(template_map: dict | None, source_map: dict | None) -> dict:
    """Combine numbering sequence maps, giving precedence to key patterns."""
    combined = dict(template_map or {})
    if not source_map:
        return combined
    root_tokens = {"alpha_upper", "decimal", "roman_upper"}
    for seq, level in source_map.items():
        if not seq or level is None:
            continue
        if len(seq) == 1 and level <= 3:
            if seq[0] in root_tokens:
                combined[seq] = level
            continue
        if seq not in combined:
            combined[seq] = level
    return combined

def serialize_sequence_map(sequence_map: dict) -> list[dict]:
    """Convert sequence map to a serializable list format."""
    items = []
    for seq, level in (sequence_map or {}).items():
        if not seq:
            continue
        if not isinstance(seq, (tuple, list)):
            continue
        items.append({"sequence": list(seq), "level": int(level) if level is not None else None})
    return items

def deserialize_sequence_map(raw) -> dict:
    """Convert serialized sequence data back into a map."""
    seq_map: dict[tuple[str, ...], int] = {}
    if not raw:
        return seq_map
    if isinstance(raw, dict):
        for key, level in raw.items():
            if not isinstance(key, str):
                continue
            parts = [p for p in key.split(">") if p]
            if not parts:
                continue
            try:
                seq_map[tuple(parts)] = int(level)
            except Exception:
                continue
        return seq_map
    if isinstance(raw, list):
        for item in raw:
            if not isinstance(item, dict):
                continue
            seq = item.get("sequence")
            level = item.get("level")
            if not isinstance(seq, list) or not seq:
                continue
            try:
                seq_map[tuple(seq)] = int(level)
            except Exception:
                continue
    return seq_map

# --- Structural Transformations ---

_WHITESPACE_EQUIV = {
    '\u00a0': ' ',  # NBSP -> space
    '\u2009': ' ', '\u2002': ' ', '\u2003': ' ', '\u200a': ' ',
    '\u202f': ' ', '\u205f': ' ', '\u3000': ' '
}
_ZERO_WIDTH = {'\u200b', '\u200c', '\u200d'}

def _norm_char(c: str) -> str:
    if c in _ZERO_WIDTH:
        return ''
    return _WHITESPACE_EQUIV.get(c, c)

def _normalized(s: str) -> str:
    """Normalize whitespace and remove zero-width characters."""
    return ''.join(_norm_char(c) for c in s)

_HEADING_PREFIX_RE = re.compile(
    r"^\s*(?:"
    # Section/Chapter with multi-level numbering: Section I.1 | Section I.A.2 | Chapter 1.2.3
    r"(?i:(?:Chapter|Section))\s+[IVXLCDM\d]+(?:\.[A-Z\d]+)*(?:\s*[:.\u2013\u2014\-])?\s+|"
    r"(?:\d+|[IVXLCDM]{1,6}|[A-Z]{1,3}|[a-z]{1,3})(?:[.\s]+(?:\d+|[IVXLCDM]{1,6}|[A-Z]{1,3}|[a-z]{1,3})){1,5}\.?\s+(?:[:.\u2013\u2014\-]\s*)?|"
    r"(?i:[IVXLCDM]+)\.(?:[A-Z]{1,3}|[a-z]{1,3})(?:\.\d+){0,3}\.?\s+(?:[:.\u2013\u2014\-]\s*)?|"  # I.A. | I.AA.2. | I.A.2.1. : (requires space after)
    r"(?i:[IVXLCDM]+)(?:\.\d+){0,3}\.?\s+(?:[:.\u2013\u2014\-]\s*)?|"         # II.3. | IV.2.1. (requires space after)
    r"(?:[A-Z]{1,3}|[a-z]{1,3})\.\s+(?:[:.\u2013\u2014\-]\s*)?|"                          # A. or AA. (requires space after)
    r"\d+(?:\.\d+){0,3}(?:[.)])?\s+(?:[:.\u2013\u2014\-]\s*)?"          # 1. | 1.2 | 1.2.3 (requires space after)
    r")"
)

_HEADING_INFER_TOKEN = r'(?:\d+|[IVXLCDMivxlcdm]{1,6}|[A-Z]{1,3}|[a-z]{1,3})'
_HEADING_INFER_RE = re.compile(
    r'^\s*(?:(?i:Chapter|Section)\s+)?'
    r'(' + _HEADING_INFER_TOKEN + r'(?:\.\s*' + _HEADING_INFER_TOKEN + r'){1,5})\b'
)

def _is_numeric_or_roman(token: str) -> bool:
    if not token:
        return False
    if token.isdigit():
        return True
    return bool(re.fullmatch(r'[IVXLCDM]+', token, re.IGNORECASE))

def _extract_heading_prefix_tokens(text: str) -> list[str]:
    if not text:
        return []
    normalized = _normalized(text).strip()
    m = _HEADING_INFER_RE.match(normalized)
    if not m:
        return []
    seq = m.group(1)
    parts = [p.strip() for p in seq.split('.') if p.strip()]
    return parts

def _extract_style_map_tokens(text: str) -> list[str]:
    if not text:
        return []
    normalized = _normalized(text)
    m = _HEADING_PREFIX_RE.match(normalized)
    if not m:
        return []
    prefix = m.group(0)
    prefix = re.sub(r'^(?:Chapter|Section)\s+', '', prefix, flags=re.IGNORECASE).strip()
    prefix = re.sub(r'[:.\-\s]+$', '', prefix).strip()
    parts = [p for p in re.split(r'[.\s]+', prefix) if p]
    return parts

def _classify_heading_token(token: str) -> str:
    if not token:
        return ""
    if token.isdigit():
        return "decimal"
    if re.fullmatch(r'[ivxlcdm]+', token):
        return "roman_lower"
    if re.fullmatch(r'[IVXLCDM]+', token):
        return "roman_upper"
    if token.isalpha():
        return "alpha_upper" if token.isupper() else "alpha_lower"
    return ""

def _classify_token_for_style_map(token: str, parts_len: int, style_map: dict, prefer_roman_single: bool = False) -> str:
    base = _classify_heading_token(token)
    if not token or not style_map:
        return base
    if len(token) == 1 and token.lower() in "ivxlcdm":
        has_alpha = "alpha_lower" in style_map or "alpha_upper" in style_map
        has_roman = "roman_lower" in style_map or "roman_upper" in style_map
        if has_roman and not has_alpha:
            return "roman_lower" if token.islower() else "roman_upper"
        if has_alpha and not has_roman:
            return "alpha_lower" if token.islower() else "alpha_upper"
        if has_alpha and has_roman:
            # Treat single-letter roman as alpha unless depth suggests a roman level (e.g., A.1.a.i)
            if prefer_roman_single or parts_len >= 5:
                return "roman_lower" if token.islower() else "roman_upper"
            return "alpha_lower" if token.islower() else "alpha_upper"
    return base

def extract_style_map_from_reference(ref_path: Path) -> tuple[dict, dict]:
    """
    Read the "Converter Style Map" section in a reference DOCX and build:
      - style_map: last-token-type -> heading level
      - sequence_map: full token sequence -> heading level
    Expected format:
      - A heading titled "Converter Style Map"
      - Example headings underneath (e.g., "A. Example", "1. Example", "a. Example", "i. Example", "1.1 Example")
      - A paragraph "End Style Map" to terminate the section
    """
    style_map = {}
    sequence_map = {}
    if not ref_path or not Path(ref_path).exists():
        return style_map, sequence_map
    try:
        doc = Document(str(ref_path))
    except Exception:
        return style_map, sequence_map

    in_map = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        if re.search(r'converter\s+style\s+map', text, re.IGNORECASE):
            in_map = True
            continue
        if not in_map:
            continue
        if re.search(r'end\s+style\s+map', text, re.IGNORECASE):
            break
        tokens = _extract_style_map_tokens(text)
        if not tokens:
            continue
        token_types = []
        for token in tokens:
            token_type = _classify_heading_token(token)
            if not token_type:
                token_types = []
                break
            token_types.append(token_type)
        if not token_types:
            continue
        level = get_heading_level(p)
        if level <= 0:
            continue
        last_token_type = token_types[-1]
        if last_token_type not in style_map:
            style_map[last_token_type] = level
        seq_key = tuple(token_types)
        if seq_key not in sequence_map:
            sequence_map[seq_key] = level
    return style_map, sequence_map

def _has_number_prefix(text: str) -> bool:
    """Check if text starts with a numbering prefix."""
    if not text:
        return False
    normalized = _normalized(text).strip()
    if re.match(r"^(?:Chapter|Section)\s+[\w\dIVXLCDM]+", normalized, re.IGNORECASE):
        return True
    if re.match(r"^(?:[IVXLCDM]+|[A-Z]|\d+)(?:\.[A-Z\d]+)*[.)]\s+", normalized, re.IGNORECASE):
        return True
    return bool(re.match(r"^(?:[IVXLCDM]+|[A-Z]|\d+)(?:\.[A-Za-z\d]+){1,5}(?=\S)", normalized, re.IGNORECASE))

def inject_heading_list_numbering(doc: Document) -> None:
    """Add list numbering prefixes from Word definitions to heading text if missing."""
    abstract_nums, num_to_abstract = _extract_numbering_defs(doc)
    if not abstract_nums or not num_to_abstract:
        return

    counters: dict[str, list[int | None]] = {}

    for p in doc.paragraphs:
        numpr = _get_paragraph_numpr(p)
        if numpr is None:
            continue

        num_id = _numpr_val(numpr.numId)
        ilvl = _numpr_val(numpr.ilvl)
        if num_id is None or ilvl is None:
            heading_level = get_heading_level(p)
            if num_id is None or heading_level <= 0:
                continue

            num_id_str = str(num_id)
            abs_id = num_to_abstract.get(num_id_str)
            levels = abstract_nums.get(abs_id or "")

            style_id = None
            try:
                style_id = p.style.style_id if p.style is not None else None
            except Exception:
                style_id = None
            if style_id and levels:
                for lvl_idx, lvl_def in levels.items():
                    if (lvl_def.get("pStyle") or "").lower() == style_id.lower():
                        ilvl = lvl_idx
                        break

            if ilvl is None:
                ilvl = max(0, heading_level - 1)

        num_id_str = str(num_id)
        abs_id = num_to_abstract.get(num_id_str)
        levels = abstract_nums.get(abs_id or "")
        if not levels:
            continue

        if num_id_str not in counters:
            counters[num_id_str] = [None] * 9
        counter = counters[num_id_str]

        start_val = levels.get(ilvl, {}).get("start", 1)
        if counter[ilvl] is None:
            counter[ilvl] = start_val
        else:
            counter[ilvl] += 1
        for j in range(ilvl + 1, len(counter)):
            counter[j] = None

        num_text = _build_list_number(levels, counter, ilvl)
        if not num_text:
            continue

        if get_heading_level(p) <= 0:
            continue

        text = (p.text or "").strip()
        if not text:
            continue

        if _has_number_prefix(text) or text.lower().startswith(num_text.lower()):
            continue

        prefix = num_text
        if not prefix.endswith((" ", "\t")):
            prefix += " "

        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = prefix
        if prefix.startswith(" ") or prefix.endswith(" "):
            t.set(qn("xml:space"), "preserve")
        r.append(t)
        p._p.insert(0, r)

def apply_heading_levels_from_numbering(doc: Document, style_map: dict, sequence_map: dict) -> None:
    """Adjust Heading styles based on the identified numbering sequence."""
    if not style_map and not sequence_map:
        return
    abstract_nums, num_to_abstract = _extract_numbering_defs(doc)
    if not abstract_nums or not num_to_abstract:
        return

    for p in doc.paragraphs:
        current_level = get_heading_level(p)
        if current_level <= 0:
            continue

        numpr = _get_paragraph_numpr(p)
        if numpr is None:
            continue
        num_id = _numpr_val(numpr.numId)
        ilvl = _numpr_val(numpr.ilvl)
        if num_id is None:
            continue

        num_id_str = str(num_id)
        abs_id = num_to_abstract.get(num_id_str)
        levels = abstract_nums.get(abs_id or "")
        if not levels:
            continue

        if ilvl is None:
            style_id = None
            try:
                style_id = p.style.style_id if p.style is not None else None
            except Exception:
                style_id = None
            if style_id:
                for lvl_idx, lvl_def in levels.items():
                    if (lvl_def.get("pStyle") or "").lower() == style_id.lower():
                        ilvl = lvl_idx
                        break
        if ilvl is None:
            continue

        sequence = _token_sequence_from_lvltext(levels, ilvl)
        target_level = None
        if sequence_map and sequence in sequence_map:
            target_level = sequence_map[sequence]
        if not target_level:
            token_type = sequence[-1] if sequence else _v4_token_type_from_numfmt(levels.get(ilvl, {}).get("numFmt"))
            if token_type and token_type in style_map:
                target_level = style_map[token_type]

        if not target_level or target_level == current_level:
            continue
        if 1 <= target_level <= 6:
            try:
                p.style = f"Heading {target_level}"
            except Exception:
                pass

def promote_headings(doc: Document) -> None:
    """Ensure essential headings are properly styled as Heading 1-4."""
    for p in list(doc.paragraphs):
        txt = (p.text or "").strip()
        if not txt:
            continue
        lvl = is_heading_style(p) or get_outline_level(p) or guess_heading_level(txt)
        if lvl and 1 <= lvl <= 4:
            try:
                p.style = f"Heading {lvl}"
            except Exception:
                pass

def strip_manual_toc_paragraphs(doc: Document) -> None:
    """Remove Word table of contents and cover page material directly from the DOCX."""
    idxs, stripping = [], False
    toc_start_idx = None

    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()

        if re.fullmatch(r"\d{4}(-\d{4})?", t):
            idxs.append(i); continue

        if re.fullmatch(r"(?:Policies and Procedures|Faculty)\s+Manual", t, re.IGNORECASE):
            idxs.append(i); continue

        if re.search(r"Table of Contents", t, re.IGNORECASE):
            toc_start_idx = i
            idxs.append(i)
            stripping = True
            continue

        if stripping:
            matches_chapter_pattern = (
                re.match(r"^Chapter\s+(?:One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve|\d+|[IVXLCDM]+)\s+[-–—]\s+\w", t, re.IGNORECASE) or
                re.match(r"^Section\s+[IVXLCDM]+\s+[-–—]\s+\w", t, re.IGNORECASE) or
                re.match(r"^Section\s+\d+\s+[-–—]\s+\w", t, re.IGNORECASE)
            )
            
            if matches_chapter_pattern:
                has_trailing_page_number = bool(re.search(r"\d+\s*$", t))
                is_short = len(t) < 120
                if is_short and has_trailing_page_number:
                    idxs.append(i)
                    continue
                else:
                    stripping = False
            elif len(t) > 120 and not re.search(r"\d+\s*$", t):
                stripping = False
            else:
                is_toc_entry = (
                    t == "" or
                    re.fullmatch(r"(\.{3,}|\d+)", t) or
                    (re.match(r"^[A-Z0-9]+\.\s+", t) and not re.match(r"^(?:Chapter|Section)\s+", t, re.IGNORECASE)) or
                    (len(t) < 120 and re.search(r"\d+\s*$", t) and not re.match(r"^(?:Chapter|Section)\s+", t, re.IGNORECASE))
                )
                if is_toc_entry:
                    idxs.append(i)
                    continue
                else:
                    if toc_start_idx is not None and (i - toc_start_idx) > 10:
                        stripping = False

    for i in reversed(idxs):
        el = doc.paragraphs[i]._p
        el.getparent().remove(el)

def build_numbering_crosswalk(doc: Document, manual_type: str = "chapter") -> dict:
    """Maps old numbering strings (1.D.2.e) to new numbering formats (1.4.2.5)."""
    crosswalk = {}
    heading_pattern = re.compile(
        r'^(?:Chapter|Section)\s+((?:One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve|Thirteen|Fourteen|Fifteen|Sixteen|Seventeen|Eighteen|Nineteen|Twenty|[\dIVXLCDM]+)(?:\.[A-Za-z\d]+)*)|'
        r'^([\dIVXLCDM]+(?:\.[A-Za-z\d]+)+)',
        re.IGNORECASE
    )
    
    chapter_counter = 0
    section_counter = 0
    subsection_counter = 0
    subsubsection_counter = 0
    subsubsubsection_counter = 0
    
    def letter_to_number(letter):
        if len(letter) != 1:
            return None
        if letter.isupper():
            return ord(letter) - ord('A') + 1
        elif letter.islower():
            return ord(letter) - ord('a') + 1
        return None
    
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        
        is_heading = is_heading_style(p) or get_outline_level(p) or guess_heading_level(text)
        if is_heading:
            match = heading_pattern.match(text)
            if match:
                numbering = match.group(1) or match.group(2)
                if numbering:
                    full_old = match.group(0)
                    level = is_heading if isinstance(is_heading, int) else 1
                    
                    if level == 1:
                        chapter_counter += 1
                        section_counter = subsection_counter = subsubsection_counter = subsubsubsection_counter = 0
                        # Match apply_css_counter_numbering: decimal "Section N" / "Chapter N"
                        pref = "Section" if is_section_style(manual_type) else "Chapter"
                        new_numbering = f"{pref} {chapter_counter}"
                    elif level == 2:
                        section_counter += 1
                        subsection_counter = subsubsection_counter = subsubsubsection_counter = 0
                        new_numbering = f"{chapter_counter}.{section_counter}"
                    elif level == 3:
                        subsection_counter += 1
                        subsubsection_counter = subsubsubsection_counter = 0
                        new_numbering = f"{chapter_counter}.{section_counter}.{subsection_counter}"
                    elif level == 4:
                        subsubsection_counter += 1
                        subsubsubsection_counter = 0
                        new_numbering = f"{chapter_counter}.{section_counter}.{subsection_counter}.{subsubsection_counter}"
                    elif level == 5:
                        subsubsubsection_counter += 1
                        new_numbering = f"{chapter_counter}.{section_counter}.{subsection_counter}.{subsubsection_counter}.{subsubsubsection_counter}"
                    else:
                        continue
                    
                    crosswalk[full_old] = new_numbering
                    crosswalk[numbering] = new_numbering.split()[-1] if ' ' in new_numbering else new_numbering
    
    return crosswalk

def generate_stable_ref_id(para_idx: int, start_pos: int, ref_text: str) -> str:
    """Generate a stable, content-based reference ID."""
    ref_hash = hashlib.md5(ref_text.encode('utf-8')).hexdigest()[:8]
    return f"ref_{para_idx}_{start_pos}_{ref_hash}"

def extract_heading_structure_and_references(doc: Document) -> tuple[dict, list]:
    """Analyzes DOCX to extract heading structure and all internal references."""
    heading_map = {}
    references = []
    
    heading_pattern = re.compile(
        r'^(?:Chapter|Section)\s+((?:One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve|Thirteen|Fourteen|Fifteen|Sixteen|Seventeen|Eighteen|Nineteen|Twenty|[\dIVXLCDM]+)(?:\.[A-Za-z\d]+)*)|'
        r'^([\dIVXLCDM]+(?:\.[A-Za-z\d]+)+)',
        re.IGNORECASE
    )
    
    reference_pattern = re.compile(
        r'(?:Section|Chapter)\s+[\dIVXLCDM]+(?:\.[A-Za-z\d]+)+|'
        r'(?<!\w)[\dIVXLCDM]+(?:\.[A-Za-z\d]+){2,}(?!\w)',
        re.IGNORECASE
    )
    
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text: continue

        has_style = is_heading_style(p)
        has_outline = get_outline_level(p)
        has_guess = guess_heading_level(text)
        is_heading = has_style or has_outline or has_guess

        heading_match = heading_pattern.match(text)
        looks_like_heading = bool(heading_match) and (len(text.split()) <= 20)

        if is_heading or looks_like_heading:
            if heading_match:
                numbering = heading_match.group(1) or heading_match.group(2)
                if numbering:
                    numbering = normalize_heading_ref(numbering)
                    if not numbering: continue
                    title = text[heading_match.end():].strip()
                    title = re.sub(r'^[\s\-:\.–—]+', '', title)
                    sentence_end = re.search(r'\.\s+[A-Z]', title)
                    if sentence_end: title = title[:sentence_end.start() + 1]
                    title = title[:200].strip()
                    if not title: continue

                    if heading_match.group(0).startswith(('Chapter', 'Section')):
                        prefix = heading_match.group(0).split()[0]
                        heading_map[f"{prefix} {numbering}"] = title
                    else:
                        heading_map[numbering] = title
            elif is_heading:
                synthetic_key = f"_h{i}_{text[:30]}"
                title = text[:200].strip()
                sentence_end = re.search(r'\.\s+[A-Z]', title)
                if sentence_end: title = title[:sentence_end.start() + 1]
                if title and len(title.split()) <= 20:
                    heading_map[synthetic_key] = title

    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text: continue
        if is_heading_style(p) or get_outline_level(p) or guess_heading_level(text):
            continue
        if re.search(r'\.{3,}|\t+\d+$', text):
            continue
        ref_matches = reference_pattern.finditer(text)
        for ref_match in ref_matches:
            ref_text = ref_match.group(0)
            if is_non_reference_token(ref_text):
                continue
            start_pos = ref_match.start()
            end_pos = ref_match.end()
            if start_pos == 0:
                rest_of_text = text[end_pos:].strip()
                rest_of_text = re.sub(r'^[\s\.\-:–—]+', '', rest_of_text)
                if rest_of_text and rest_of_text[0].isupper():
                    before_period = re.split(r'\.\s+[A-Z]', rest_of_text)[0]
                    if len(before_period.split()) <= 15 and len(before_period) < 100:
                        continue
            references.append((i, text, ref_text, start_pos, end_pos))
    
    return heading_map, references

# --- Hyperlinks & Tables ---

def is_bad_docx_link(href: str) -> bool:
    """Identify broken, local, or unsafe links in DOCX (only http(s), mailto, and # anchors allowed)."""
    if not href:
        return True
    value = href.strip()
    lower = value.lower()
    if lower.startswith("#x") or lower.startswith("file:"):
        return True
    if re.match(r'^[a-zA-Z]:\\\\', value) or re.match(r'^[a-zA-Z]:/', value) or value.startswith("\\\\"):
        return True
    if lower.startswith("#"):
        return False
    if lower.startswith("mailto:"):
        return "@" not in value[7:80]
    if lower.startswith("http://") or lower.startswith("https://"):
        return False
    if lower.startswith(("javascript:", "vbscript:", "data:")):
        return True
    return True

def extract_docx_hyperlinks(doc: Document) -> dict:
    """Extract all hyperlinks from a DOCX file, mapped by paragraph index."""
    links_by_para = {}
    for idx, p in enumerate(doc.paragraphs):
        try:
            hyperlink_elems = p._p.xpath('.//w:hyperlink')
        except Exception:
            hyperlink_elems = []
        if not hyperlink_elems:
            continue
        for link in hyperlink_elems:
            href = ""
            try:
                rel_id = link.get(qn('r:id'))
            except Exception:
                rel_id = None
            try:
                anchor = link.get(qn('w:anchor'))
            except Exception:
                anchor = None
            if rel_id and rel_id in doc.part.rels:
                href = doc.part.rels[rel_id].target_ref
            elif anchor:
                href = f"#{anchor}"
            try:
                texts = [node.text for node in link.xpath('.//w:t', namespaces=link.nsmap) if node.text]
                text = "".join(texts).strip()
            except Exception:
                text = ""
            if not href and not text:
                continue
            links_by_para.setdefault(idx, []).append({
                "text": text,
                "href": href,
                "bad": is_bad_docx_link(href)
            })
    return links_by_para

# Parts of a DOCX that can carry revision marks.
_REVISION_PARTS = (
    "word/document.xml", "word/footnotes.xml", "word/endnotes.xml",
)
_REVISION_MARK_RE = re.compile(rb"<w:(?:ins|del)\b")


def count_tracked_changes(path: Path) -> int:
    """Number of unresolved revision marks in a DOCX.

    The converter reads the document twice and the two readers disagree about
    tracked changes: python-docx (heading and reference extraction) sees neither
    inserted nor deleted text, while Pandoc accepts changes when producing the
    HTML. A manual returned from an editor with changes still pending therefore
    converts with headings the extractor reads as empty and references it never
    sees — measured on the real Faculty Manual, 14 headings read as empty and two
    curated internal links disappeared, with no error. Callers refuse the upload
    rather than let that happen quietly.
    """
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            return sum(
                len(_REVISION_MARK_RE.findall(archive.read(part)))
                for part in _REVISION_PARTS if part in names
            )
    except (zipfile.BadZipFile, OSError, KeyError):
        # Not readable as a DOCX; the conversion itself will report that.
        return 0


def has_tables_in_docx(path: Path) -> bool:
    """Check if a DOCX file contains any tables."""
    try:
        doc = Document(path)
        return bool(doc.tables)
    except Exception:
        return False

# --- Bookmarks & XML Sanitization ---

def sanitize_docx_bookmark_id(value: str) -> str:
    """Ensure a bookmark ID is Word-safe."""
    if not value:
        return "b_1"
    cleaned = re.sub(r'[^A-Za-z0-9_]', '_', value)
    if not re.match(r'^[A-Za-z]', cleaned):
        cleaned = f"b_{cleaned}"
    return cleaned

def sanitize_docx_styles(path: Path) -> None:
    """Fix missing style references in document.xml by replacing with defaults."""
    if not path.exists():
        return
    try:
        with zipfile.ZipFile(path, 'r') as zf:
            styles_xml = zf.read('word/styles.xml').decode('utf-8', errors='ignore')
            doc_xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
            other_files = {name: zf.read(name) for name in zf.namelist()
                          if name not in ('word/styles.xml', 'word/document.xml')}
    except Exception as e:
        logger.error(f"sanitize_docx_styles failed to read: {e}")
        return

    style_ids = set(re.findall(r'<w:style[^>]*w:styleId="([^"]+)"', styles_xml))
    p_styles = set(re.findall(r'<w:pStyle w:val="([^"]+)"', doc_xml))
    r_styles = set(re.findall(r'<w:rStyle w:val="([^"]+)"', doc_xml))
    missing_p = p_styles - style_ids
    missing_r = r_styles - style_ids

    if missing_p or missing_r:
        fixed_xml = doc_xml
        for s in missing_p:
            fixed_xml = fixed_xml.replace(f'<w:pStyle w:val="{s}"', '<w:pStyle w:val="Normal"')
        for s in missing_r:
            fixed_xml = fixed_xml.replace(f'<w:rStyle w:val="{s}"', '<w:rStyle w:val="DefaultParagraphFont"')
        try:
            with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zf:
                zf.writestr('word/styles.xml', styles_xml.encode('utf-8'))
                zf.writestr('word/document.xml', fixed_xml.encode('utf-8'))
                for name, data in other_files.items():
                    zf.writestr(name, data)
        except Exception as e:
            logger.error(f"sanitize_docx_styles failed to write: {e}")

def fix_numbering_xml(docx_path: Path) -> None:
    """Remove orphan w:num entries that reference non-existent abstractNum definitions."""
    if not docx_path.exists():
        return
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            if 'word/numbering.xml' not in zf.namelist(): return
            numbering = zf.read('word/numbering.xml').decode('utf-8')
            other_files = {n: zf.read(n) for n in zf.namelist() if n != 'word/numbering.xml'}
    except Exception as e:
        logger.error(f"fix_numbering_xml failed to read: {e}")
        return

    defined_abstract = set(re.findall(r'<w:abstractNum[^>]*w:abstractNumId="(\d+)"', numbering))
    def remove_orphan_num(match):
        num_block = match.group(0)
        ref_match = re.search(r'<w:abstractNumId w:val="(\d+)"', num_block)
        if ref_match and ref_match.group(1) not in defined_abstract:
            return ''
        return num_block

    fixed = re.sub(r'<w:num\s[^>]*w:numId="[^"]*"[^>]*>.*?</w:num>', remove_orphan_num, numbering, flags=re.DOTALL)
    try:
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('word/numbering.xml', fixed.encode('utf-8'))
            for name, data in other_files.items():
                zf.writestr(name, data)
    except Exception as e:
        logger.error(f"fix_numbering_xml failed to write: {e}")

def relocate_body_level_bookmarks(path: Path) -> None:
    """Move body-level bookmarks into the following paragraph for better Word compatibility."""
    if not path.exists(): return
    try:
        with zipfile.ZipFile(path, "r") as zf:
            doc_xml = zf.read("word/document.xml")
            other_files = {name: zf.read(name) for name in zf.namelist() if name != "word/document.xml"}
    except Exception: return

    try:
        ns_map = {}
        for _, elem in ET.iterparse(io.BytesIO(doc_xml), events=("start-ns",)):
            prefix, uri = elem
            if prefix not in ns_map: ns_map[prefix] = uri
        if "w" not in ns_map: ns_map["w"] = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for prefix, uri in ns_map.items(): ET.register_namespace(prefix, uri)
        w_ns = ns_map.get("w")
        ns = {"w": w_ns}
        root = ET.fromstring(doc_xml)
        body = root.find("w:body", ns)
        if body is None: return

        children = list(body)
        starts, ends = {}, {}
        for idx, child in enumerate(children):
            tag = child.tag.split("}", 1)[1] if "}" in child.tag else child.tag
            bid = child.attrib.get(f"{{{w_ns}}}id")
            if tag == "bookmarkStart" and bid: starts[bid] = (idx, child)
            elif tag == "bookmarkEnd" and bid: ends[bid] = (idx, child)

        if not starts: return
        for bid, (start_idx, start_el) in starts.items():
            if bid not in ends: continue
            end_idx, end_el = ends[bid]
            target_p = None
            for j in range(start_idx + 1, len(children)):
                candidate = children[j]
                tag = candidate.tag.split("}", 1)[1] if "}" in candidate.tag else candidate.tag
                if tag == "p":
                    target_p = candidate; break
                if tag == "tbl":
                    p = candidate.find(".//w:p", ns)
                    if p is not None: target_p = p; break
            if target_p is not None:
                body.remove(start_el)
                body.remove(end_el)
                pPr = target_p.find("w:pPr", ns)
                target_p.insert(1 if pPr is not None else 0, start_el)
                target_p.append(end_el)

        updated = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("word/document.xml", updated)
            for name, data in other_files.items(): zf.writestr(name, data)
    except Exception as exc: logger.error(f"relocate_body_level_bookmarks failed: {exc}")

# --- Reference Document Support ---

def _parse_twips(value: str | None, fallback: float | None = None) -> float | None:
    if not value: return fallback
    try: return float(value) / 1440.0
    except (TypeError, ValueError): return fallback

def _to_twips(value: float | None) -> str | None:
    if value is None: return None
    try: return str(int(round(float(value) * 1440)))
    except (TypeError, ValueError): return None

def extract_reference_doc_summary(path: Path) -> dict:
    """Extract styling and layout summary from a reference DOCX."""
    summary = {
        "body_font": "", "body_size_pt": None, "heading_font": "",
        "heading_sizes_pt": {}, "link_color": "", "margins_in": {},
        "orientation": "portrait", "include_header_footer": True
    }
    if not path or not path.exists(): return summary
    try:
        with zipfile.ZipFile(path, "r") as zf:
            styles_xml = zf.read("word/styles.xml")
            doc_xml = zf.read("word/document.xml")
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        styles_root = ET.fromstring(styles_xml)
        doc_root = ET.fromstring(doc_xml)

        def _style_font_info(style_id):
            for style in styles_root.findall("w:style", ns):
                if (style.get(f"{{{ns['w']}}}styleId") or "").lower() == style_id.lower():
                    rpr = style.find("w:rPr", ns)
                    if rpr is None: return None, None
                    rfonts = rpr.find("w:rFonts", ns)
                    font = None
                    if rfonts is not None:
                        font = rfonts.get(f"{{{ns['w']}}}ascii") or rfonts.get(f"{{{ns['w']}}}hAnsi")
                    size_el = rpr.find("w:sz", ns)
                    size_pt = int(size_el.get(f"{{{ns['w']}}}val")) / 2.0 if size_el is not None else None
                    return font, size_pt
            return None, None

        summary["body_font"], summary["body_size_pt"] = _style_font_info("Normal")
        h_font = None
        for l in range(1, 7):
            f, s = _style_font_info(f"Heading{l}")
            if f and not h_font: h_font = f
            if s: summary["heading_sizes_pt"][str(l)] = s
        summary["heading_font"] = h_font or ""

        body = doc_root.find("w:body", ns)
        if body is not None:
            sect = body.find("w:sectPr", ns)
            if sect is not None:
                pg_mar = sect.find("w:pgMar", ns)
                if pg_mar is not None:
                    summary["margins_in"] = {k: _parse_twips(pg_mar.get(f"{{{ns['w']}}}{k}")) for k in ("top", "right", "bottom", "left")}
                pg_sz = sect.find("w:pgSz", ns)
                if pg_sz is not None and pg_sz.get(f"{{{ns['w']}}}orient") == "landscape":
                    summary["orientation"] = "landscape"
    except Exception as exc: logger.error(f"extract_reference_doc_summary failed: {exc}")
    return summary

def _remove_child_elements(parent, tag, ns):
    for child in list(parent):
        local = child.tag.split("}", 1)[1] if "}" in child.tag else child.tag
        if local == tag: parent.remove(child)

def _update_style_font(style, ns, font_name: str | None, size_pt: float | None):
    if style is None: return
    rpr = style.find("w:rPr", ns)
    if rpr is None: rpr = ET.SubElement(style, f"{{{ns['w']}}}rPr")
    if font_name:
        rfonts = rpr.find("w:rFonts", ns)
        if rfonts is None: rfonts = ET.SubElement(rpr, f"{{{ns['w']}}}rFonts")
        for k in ("ascii", "hAnsi", "cs"): rfonts.set(f"{{{ns['w']}}}{k}", font_name)
    if size_pt is not None:
        size_val = str(int(round(size_pt * 2)))
        sz = rpr.find("w:sz", ns)
        if sz is None: sz = ET.SubElement(rpr, f"{{{ns['w']}}}sz")
        sz.set(f"{{{ns['w']}}}val", size_val)

def build_clean_reference_doc(src_path: Path, output_path: Path, overrides: dict | None = None) -> Path | None:
    """Build a minimal reference DOCX with specified style overrides."""
    if not src_path.exists(): return None
    overrides = overrides or {}
    try:
        with zipfile.ZipFile(src_path, "r") as zin: src_files = {n: zin.read(n) for n in zin.namelist()}
        styles_xml = src_files.get("word/styles.xml", b"")
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main", "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
        styles_root = ET.fromstring(styles_xml)
        
        def find_style(sid):
            for s in styles_root.findall("w:style", ns):
                if (s.get(f"{{{ns['w']}}}styleId") or "").lower() == sid.lower(): return s
            return None

        _update_style_font(find_style("Normal"), ns, overrides.get("body_font"), overrides.get("body_size_pt"))
        h_sizes = overrides.get("heading_sizes_pt") or {}
        for l in range(1, 7):
            _update_style_font(find_style(f"Heading{l}"), ns, overrides.get("heading_font"), h_sizes.get(str(l)))

        src_files["word/styles.xml"] = ET.tostring(styles_root, encoding="utf-8", xml_declaration=True)
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for n, d in src_files.items(): zout.writestr(n, d)
        return output_path
    except Exception as exc: logger.error(f"build_clean_reference_doc failed: {exc}"); return None

# --- Reference Configuration & Styles ---

def _rgbcolor_from_hex(value: str):
    if not value: return None
    raw = str(value).strip().lstrip("#")
    if len(raw) == 3: raw = "".join(ch * 2 for ch in raw)
    if len(raw) != 6: return None
    try: return RGBColor.from_string(raw.upper())
    except Exception: return None

def _config_font_size_pt(style_info):
    if not style_info: return None
    for key in ("font_size_pt", "font_size", "size"):
        if key in style_info and style_info.get(key) is not None:
            try: return float(style_info.get(key))
            except (TypeError, ValueError): return None
    return None

def _apply_paragraph_format_to_style(style_obj, format_info):
    if not format_info: return
    pf = style_obj.paragraph_format
    if format_info.get("left_indent_pt") is not None: pf.left_indent = Pt(format_info["left_indent_pt"])
    if format_info.get("right_indent_pt") is not None: pf.right_indent = Pt(format_info["right_indent_pt"])
    if format_info.get("first_line_indent_pt") is not None: pf.first_line_indent = Pt(format_info["first_line_indent_pt"])
    if format_info.get("space_before_pt") is not None: pf.space_before = Pt(format_info["space_before_pt"])
    if format_info.get("space_after_pt") is not None: pf.space_after = Pt(format_info["space_after_pt"])
    if format_info.get("line_spacing") is not None: pf.line_spacing = format_info["line_spacing"]

def _apply_style_to_docx_style(style_obj, style_info):
    if not style_info: return
    font_name = style_info.get("font_name") or style_info.get("font")
    font_size_pt = _config_font_size_pt(style_info)
    font_color = _rgbcolor_from_hex(style_info.get("color"))
    if font_name: style_obj.font.name = font_name
    if font_size_pt: style_obj.font.size = Pt(font_size_pt)
    if font_color: style_obj.font.color.rgb = font_color
    if style_info.get("bold") is not None: style_obj.font.bold = bool(style_info.get("bold"))
    if style_info.get("italic") is not None: style_obj.font.italic = bool(style_info.get("italic"))
    if style_info.get("underline") is not None: style_obj.font.underline = bool(style_info.get("underline"))
    _apply_paragraph_format_to_style(style_obj, style_info.get("paragraph_format") or {})

def load_reference_config(path: Path) -> dict:
    try: return json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc: logger.error(f"Failed to load reference config: {exc}"); return {}

def build_reference_doc_from_config(config: dict, out_path: Path) -> None:
    """Generate a DOCX file from a style configuration dictionary."""
    doc = Document()
    styles = doc.styles
    body_cfg = config.get("styles", {}).get("body", {}) or {}
    _apply_style_to_docx_style(styles["Normal"], body_cfg)
    for level in range(1, 7):
        heading_cfg = config.get("styles", {}).get("headings", {}).get(f"h{level}", {}) or {}
        if not heading_cfg: continue
        try: _apply_style_to_docx_style(styles[f"Heading {level}"], heading_cfg)
        except Exception: continue
    doc.save(str(out_path))

# --- Main Entry Point ---

def detect_manual_type_from_docx(doc: Document) -> str:
    """Infer chapter vs section from heading text, not cover-page keywords.

    Prefers explicit ``Chapter`` / ``Section`` labels. Defaults to ``chapter``.
    """
    chapter_hits = 0
    section_hits = 0
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        m = re.match(r"^(Chapter|Section)\s+", text, re.IGNORECASE)
        if not m:
            continue
        lvl = get_heading_level(p) or guess_heading_level(text)
        # Styled/outline H1, guessed H1, or a short title-like Chapter/Section line
        if not (lvl == 1 or len(text) <= 120):
            continue
        if m.group(1).lower() == "chapter":
            chapter_hits += 1
        else:
            section_hits += 1
    if chapter_hits or section_hits:
        return "section" if section_hits > chapter_hits else "chapter"
    return "chapter"


def preprocess_docx(
    input_path: Path,
    output_path: Path,
    style_map: dict | None = None,
    sequence_map: dict | None = None,
    manual_type_override: str | None = None,
) -> tuple:
    """
    Clean and prepare the source DOCX for Pandoc conversion.
    This is the main entry point for DOCX-to-DOCX structural normalization.

    ``manual_type_override`` may be ``chapter``, ``section``, or ``policy``
    (policy collapses to section-style). ``auto`` / empty / None → detect.
    """
    doc = Document(input_path)

    # 1. Detect Manual Type (H1 Chapter/Section), optional user override
    detected = detect_manual_type_from_docx(doc)
    override = (manual_type_override or "").strip().lower()
    if override in ("chapter", "section", "policy"):
        manual_type = override
        logger.info(f"DOCX: Manual type override={manual_type} (detected={detected})")
    else:
        manual_type = detected
        logger.info(f"DOCX: Detected manual type: {manual_type}")

    # 2. Structural normalization
    promote_headings(doc)
    strip_manual_toc_paragraphs(doc)
    
    # 3. Apply numbering logic if maps provided
    if style_map or sequence_map:
        apply_heading_levels_from_numbering(doc, style_map or {}, sequence_map or {})
    
    # 4. Inject list numbering for Pandoc preservation
    inject_heading_list_numbering(doc)

    # 5. Extract metadata for the response
    heading_map, references = extract_heading_structure_and_references(doc)
    old_crosswalk = build_numbering_crosswalk(doc, manual_type)
    
    doc.save(output_path)
    return heading_map, old_crosswalk, references, manual_type
