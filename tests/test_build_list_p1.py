"""Tests for build_list items 1–3: manual type, acronyms, numbering/level fixes."""
import io
from pathlib import Path

import pytest
from docx import Document

from core.docx_processor import (
    build_numbering_crosswalk,
    detect_manual_type_from_docx,
    guess_heading_level,
    preprocess_docx,
)
from core.html_processor import apply_css_counter_numbering
from core.reference_linking import extract_references_from_html, is_non_reference_token


def _docx_with_paras(paras: list[tuple[str, str | None]], path: Path) -> Path:
    """paras: list of (text, style_name or None)."""
    doc = Document()
    for text, style in paras:
        p = doc.add_paragraph(text)
        if style:
            p.style = style
    doc.save(path)
    return path


# --- guess_heading_level (must_fix #2) ---------------------------------

@pytest.mark.parametrize(
    "text, expected",
    [
        ("Section 1. Organization", 1),
        ("Section I. Organization", 1),
        ("1.1 Academic Units", 2),
        ("1.1.1 Designation of Academic Unit", 3),
        ("I.A. Something", 2),
        ("I.A.1 Something", 3),
        ("See Chapter 4.2 for details about leave.", 0),
    ],
)
def test_guess_heading_level_decimal_and_roman(text, expected):
    assert guess_heading_level(text) == expected


# --- crosswalk L1 decimal (must_fix #1) --------------------------------

def test_section_crosswalk_level1_matches_css_decimal(tmp_path):
    path = tmp_path / "s.docx"
    _docx_with_paras(
        [
            ("Section I. Organization", "Heading 1"),
            ("Section II. Duties", "Heading 1"),
        ],
        path,
    )
    doc = Document(path)
    cw = build_numbering_crosswalk(doc, "section")
    # Keys may be "Section I" style match; values must be decimal like CSS
    vals = set(cw.values())
    assert "Section 1" in vals
    assert "Section 2" in vals
    assert "Section I" not in vals

    css = apply_css_counter_numbering(
        "<h1>Organization</h1><h1>Duties</h1>", manual_type="section"
    )
    assert "Section 1 - Organization" in css
    assert "Section 2 - Duties" in css


# --- manual type detection (build_list #1) -------------------------------

def test_detect_chapter_from_h1_not_policy_keyword(tmp_path):
    path = tmp_path / "gspp.docx"
    _docx_with_paras(
        [
            ("Graduate School Policies and Procedures", None),
            ("Chapter 1. Administration", "Heading 1"),
            ("Chapter 2. Admissions", "Heading 1"),
        ],
        path,
    )
    assert detect_manual_type_from_docx(Document(path)) == "chapter"


def test_detect_section_from_h1(tmp_path):
    path = tmp_path / "fac.docx"
    _docx_with_paras(
        [
            ("Faculty Manual", None),
            ("Section I. Organization", "Heading 1"),
        ],
        path,
    )
    assert detect_manual_type_from_docx(Document(path)) == "section"


def test_preprocess_manual_type_override(tmp_path):
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _docx_with_paras(
        [("Chapter 1. Intro", "Heading 1")],
        src,
    )
    _, _, _, mt = preprocess_docx(src, out, manual_type_override="section")
    assert mt == "section"


# --- degree acronym denylist (build_list #2) ----------------------------

@pytest.mark.parametrize(
    "token, ignored",
    [
        ("D.V.M.", True),
        ("D.N.P.", True),
        ("Ph.D.", True),
        ("M.S.", True),
        ("Ed.D.", True),
        ("Chapter 4.2", False),
        ("1.2.3", False),
        ("Section I.A.2", False),
    ],
)
def test_is_non_reference_token(token, ignored):
    assert is_non_reference_token(token) is ignored


def test_html_extract_skips_degree_acronyms():
    html = (
        "<p>The D.V.M. credential is separate.</p>"
        "<p>See Chapter 4.2.1 for NADC rules.</p>"
    )
    refs = extract_references_from_html(html)
    texts = [r[2] for r in refs]
    assert "D.V.M." not in texts
    assert any("4.2.1" in t or "Chapter 4.2.1" in t for t in texts)
