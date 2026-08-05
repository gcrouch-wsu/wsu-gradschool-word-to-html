"""Importing a session bundle together with a newer Word document.

This is the return leg of the editing cycle: a bundle is exported, the DOCX goes
to whoever edits the manual, and an edited DOCX comes back. Uploading that file
on its own starts the reference review from nothing, because the link targets,
external URLs and validations live in the bundle's ``edits.json`` rather than in
the Word file.

Repackaging the zip by hand was the alternative, and it had two silent traps:
zipping the extracted *folder* rather than its contents ("Invalid bundle:
manifest.json missing") and renaming the document ("Security error: Malicious
bundle detected"). Neither message describes what actually went wrong. Accepting
the revised document alongside the bundle removes both.
"""

import io
import json
import zipfile

import pytest
from docx import Document

from config import SessionDir

from tests.conftest import pandoc_required


def _bundle_with(docx_bytes, edits=None, manifest_extra=None):
    """A minimal but valid session bundle."""
    manifest = {
        "document": "source.docx",
        "doc_hash": "0" * 64,
        "manual_type": "chapter",
        "toc_depth": 2,
        "mapping_mode": "map_new",
        "files": {"docx": "source.docx", "edits": "edits.json"},
    }
    manifest.update(manifest_extra or {})
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("source.docx", docx_bytes)
        z.writestr("edits.json", json.dumps(edits or {"reference_edits": {}}))
        z.writestr("manifest.json", json.dumps(manifest))
    return buf.getvalue()


def _docx(paragraphs, heading="Chapter One - Administration"):
    doc = Document()
    doc.add_heading(heading, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


ORIGINAL = ["Governance is described in Chapter 1.A.1 of this manual."]
REVISED = [
    "An editorial note inserted during review.",
    "Governance is described in Chapter 1.A.1 of this manual.",
]


@pandoc_required
def test_the_revised_document_is_used_instead_of_the_bundled_one(client_nocsrf):
    bundle = _bundle_with(_docx(ORIGINAL))
    r = client_nocsrf.post(
        "/import_bundle",
        data={
            "bundle": (io.BytesIO(bundle), "b.zip"),
            "revised_docx": (io.BytesIO(_docx(REVISED)), "edited.docx"),
            "skip_review": "1",
        },
        content_type="multipart/form-data",
    )
    assert r.status_code == 302
    sid = r.headers["Location"].rstrip("/").split("/")[-1]
    body = client_nocsrf.get(f"/convert/{sid}").get_data(as_text=True)
    assert "editorial note inserted during review" in body.lower()


@pandoc_required
def test_the_document_may_be_named_anything(client_nocsrf):
    """Repackaging by hand rejected a renamed file; an upload has no such rule."""
    bundle = _bundle_with(_docx(ORIGINAL))
    r = client_nocsrf.post(
        "/import_bundle",
        data={
            "bundle": (io.BytesIO(bundle), "b.zip"),
            "revised_docx": (io.BytesIO(_docx(REVISED)), "Faculty Manual FINAL v2 (edited).docx"),
            "skip_review": "1",
        },
        content_type="multipart/form-data",
    )
    assert r.status_code == 302, "the uploaded filename must not matter"


@pandoc_required
def test_the_hash_warning_is_suppressed_when_a_revision_is_supplied(client_nocsrf):
    """A mismatch is expected here, so warning about it is noise."""
    bundle = _bundle_with(_docx(ORIGINAL))
    body = client_nocsrf.post(
        "/import_bundle",
        data={
            "bundle": (io.BytesIO(bundle), "b.zip"),
            "revised_docx": (io.BytesIO(_docx(REVISED)), "edited.docx"),
            "skip_review": "1",
        },
        content_type="multipart/form-data",
        follow_redirects=True,
    ).get_data(as_text=True)
    assert "hash does not match" not in body
    assert "Using the revised document" in body


@pandoc_required
def test_the_hash_warning_still_fires_without_a_revision(client_nocsrf):
    """Swapping a document in by hand should still be flagged."""
    bundle = _bundle_with(_docx(REVISED))  # manifest hash is deliberately wrong
    body = client_nocsrf.post(
        "/import_bundle",
        data={"bundle": (io.BytesIO(bundle), "b.zip"), "skip_review": "1"},
        content_type="multipart/form-data",
        follow_redirects=True,
    ).get_data(as_text=True)
    assert "hash does not match" in body


@pytest.mark.parametrize("filename", ["notes.txt", "manual.pdf", "archive.zip"])
def test_a_non_docx_revision_is_refused(client_nocsrf, filename):
    bundle = _bundle_with(b"PK\x03\x04not-a-real-docx")
    body = client_nocsrf.post(
        "/import_bundle",
        data={
            "bundle": (io.BytesIO(bundle), "b.zip"),
            "revised_docx": (io.BytesIO(b"x"), filename),
        },
        content_type="multipart/form-data",
        follow_redirects=True,
    ).get_data(as_text=True)
    assert "must be a .docx" in body


@pandoc_required
def test_an_empty_revision_field_behaves_like_a_plain_import(client_nocsrf):
    """Submitting the form without choosing a file must not change anything."""
    bundle = _bundle_with(_docx(ORIGINAL))
    r = client_nocsrf.post(
        "/import_bundle",
        data={
            "bundle": (io.BytesIO(bundle), "b.zip"),
            "revised_docx": (io.BytesIO(b""), ""),
            "skip_review": "1",
        },
        content_type="multipart/form-data",
    )
    assert r.status_code == 302
    sid = r.headers["Location"].rstrip("/").split("/")[-1]
    body = client_nocsrf.get(f"/convert/{sid}").get_data(as_text=True)
    assert "editorial note" not in body.lower(), "the bundled document should be used"


@pandoc_required
def test_curated_reference_edits_survive_the_revision(client_nocsrf):
    """The whole point: the editor's changes land, the reference review is kept."""
    from core.docx_processor import generate_stable_ref_id

    original = _docx(ORIGINAL)
    # Find the reference the pipeline will detect, and curate a URL for it.
    r = client_nocsrf.post(
        "/import_bundle",
        data={"bundle": (io.BytesIO(_bundle_with(original)), "b.zip"), "skip_review": "1"},
        content_type="multipart/form-data",
    )
    sid = r.headers["Location"].rstrip("/").split("/")[-1]
    references = json.loads(SessionDir(sid).session_json.read_text())["references"]
    assert references, "the fixture should contain a reference"
    ref = references[0]
    rid = generate_stable_ref_id(ref[0], ref[3], ref[2])
    edits = {
        "reference_validations": {rid: True},
        "reference_external_urls": {rid: "https://policies.wsu.edu/curated"},
        "reference_edits": {rid: ref[2]},
    }

    # Now import that bundle with an edited document that shifts the reference.
    r2 = client_nocsrf.post(
        "/import_bundle",
        data={
            "bundle": (io.BytesIO(_bundle_with(original, edits=edits)), "b.zip"),
            "revised_docx": (io.BytesIO(_docx(REVISED)), "edited.docx"),
            "skip_review": "1",
        },
        content_type="multipart/form-data",
    )
    sid2 = r2.headers["Location"].rstrip("/").split("/")[-1]
    body = client_nocsrf.get(f"/convert/{sid2}").get_data(as_text=True)
    assert "editorial note inserted during review" in body.lower(), "the edit landed"
    assert "https://policies.wsu.edu/curated" in body, "the curated URL survived"
