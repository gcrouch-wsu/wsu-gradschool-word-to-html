"""Uploads carrying unresolved tracked changes are refused.

The converter reads a DOCX twice and the two readers disagree about revision
marks: python-docx, which does heading and reference extraction, sees neither
inserted nor deleted text, while Pandoc accepts changes when producing the HTML.

Measured on the real Faculty Manual returned from an editor with 43 pending
changes: 14 headings read as empty to the extractor, and two curated internal
links disappeared from the output. No error, no warning — the manual simply
published with two references that were no longer links. Refusing the upload
turns that into an instruction.
"""

import io
import re
import zipfile

import pytest
from docx import Document

from core.docx_processor import count_tracked_changes


def _docx_bytes(text="Governance is described in Chapter 1.A.1."):
    doc = Document()
    doc.add_heading("Chapter One - Administration", level=1)
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _with_revision_marks(docx_bytes, insertions=1, deletions=1):
    """Wrap a run in <w:ins>/<w:del> the way Word records a tracked edit."""
    src = io.BytesIO(docx_bytes)
    out = io.BytesIO()
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(out, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8")
                marks = ""
                for i in range(insertions):
                    marks += (f'<w:ins w:id="{900+i}" w:author="Editor" w:date="2026-01-01T00:00:00Z">'
                              f'<w:r><w:t>added {i}</w:t></w:r></w:ins>')
                for i in range(deletions):
                    marks += (f'<w:del w:id="{950+i}" w:author="Editor" w:date="2026-01-01T00:00:00Z">'
                              f'<w:r><w:delText>removed {i}</w:delText></w:r></w:del>')
                xml = xml.replace("</w:body>", f"<w:p>{marks}</w:p></w:body>")
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    return out.getvalue()


def _bundle(docx_bytes):
    import json
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("source.docx", docx_bytes)
        z.writestr("edits.json", json.dumps({"reference_edits": {}}))
        z.writestr("manifest.json", json.dumps({
            "document": "source.docx", "doc_hash": "0" * 64, "manual_type": "chapter",
            "toc_depth": 2, "mapping_mode": "map_new",
            "files": {"docx": "source.docx", "edits": "edits.json"},
        }))
    return buf.getvalue()


# --- the detector ---------------------------------------------------------

def test_a_clean_document_reports_no_changes():
    assert count_tracked_changes_from_bytes(_docx_bytes()) == 0


def test_revision_marks_are_counted():
    marked = _with_revision_marks(_docx_bytes(), insertions=3, deletions=2)
    assert count_tracked_changes_from_bytes(marked) == 5


def test_something_that_is_not_a_docx_does_not_raise(tmp_path):
    junk = tmp_path / "not.docx"
    junk.write_bytes(b"this is not a zip")
    assert count_tracked_changes(junk) == 0


def count_tracked_changes_from_bytes(data, tmp=None):
    import tempfile
    from pathlib import Path
    path = Path(tempfile.mkdtemp()) / "d.docx"
    path.write_bytes(data)
    return count_tracked_changes(path)


# --- the upload paths -----------------------------------------------------

def _refusal(response):
    return re.search(r"has \d+ unresolved tracked change", response.get_data(as_text=True))


def test_convert_refuses_a_document_with_pending_changes(client_nocsrf):
    response = client_nocsrf.post(
        "/convert",
        data={"docx": (io.BytesIO(_with_revision_marks(_docx_bytes())), "m.docx"),
              "mapping_mode": "map_new"},
        content_type="multipart/form-data", follow_redirects=True,
    )
    assert _refusal(response), "a pending change must stop the conversion"


def test_convert_names_the_count_and_says_what_to_do(client_nocsrf):
    response = client_nocsrf.post(
        "/convert",
        data={"docx": (io.BytesIO(_with_revision_marks(_docx_bytes(), 4, 3)), "m.docx"),
              "mapping_mode": "map_new"},
        content_type="multipart/form-data", follow_redirects=True,
    )
    body = response.get_data(as_text=True)
    assert "7 unresolved tracked change" in body
    assert "Accept or reject them in Word" in body


def test_a_revised_docx_with_pending_changes_is_refused(client_nocsrf):
    response = client_nocsrf.post(
        "/import_bundle",
        data={"bundle": (io.BytesIO(_bundle(_docx_bytes())), "b.zip"),
              "revised_docx": (io.BytesIO(_with_revision_marks(_docx_bytes())), "edited.docx"),
              "skip_review": "1"},
        content_type="multipart/form-data", follow_redirects=True,
    )
    assert _refusal(response)


def test_a_bundle_whose_own_document_has_pending_changes_is_refused(client_nocsrf):
    response = client_nocsrf.post(
        "/import_bundle",
        data={"bundle": (io.BytesIO(_bundle(_with_revision_marks(_docx_bytes()))), "b.zip"),
              "skip_review": "1"},
        content_type="multipart/form-data", follow_redirects=True,
    )
    assert _refusal(response)


@pytest.mark.parametrize("route,data", [
    ("/convert", None),
    ("/import_bundle", None),
])
def test_clean_documents_are_unaffected(client_nocsrf, route, data):
    """The guard must not stand between an ordinary manual and the converter."""
    if route == "/convert":
        payload = {"docx": (io.BytesIO(_docx_bytes()), "m.docx"), "mapping_mode": "map_new"}
    else:
        payload = {"bundle": (io.BytesIO(_bundle(_docx_bytes())), "b.zip"), "skip_review": "1"}
    response = client_nocsrf.post(route, data=payload, content_type="multipart/form-data")
    assert response.status_code == 302
    assert not _refusal(client_nocsrf.get(response.headers["Location"]))
