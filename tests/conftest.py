"""Shared fixtures for the test suite.

The fixture DOCX is generated with python-docx at test time rather than
committed as a binary (the repo .gitignore deliberately excludes *.docx).
"""
import io

import pytest

from word_to_wordpressV4 import app
from config import PANDOC_PINNED_VERSION
from core.pandoc_wrapper import get_pandoc_version, check_min_version

# e2e/parity tests need a Pandoc at least as new as the pinned version (the
# app's own policy is warn-but-run on mismatch, so minimum — not exact). The
# detected version is surfaced in the skip reason for diagnosability.
_installed_pandoc = get_pandoc_version()
pandoc_required = pytest.mark.skipif(
    not check_min_version(_installed_pandoc, PANDOC_PINNED_VERSION),
    reason=(
        f"Pandoc >= {PANDOC_PINNED_VERSION} required for e2e parity tests "
        f"(found {_installed_pandoc or 'none on PATH'})"
    ),
)


@pytest.fixture
def client_nocsrf():
    """Test client with CSRF disabled, for exercising route logic directly.

    (CSRF enforcement itself is covered by tests/test_csrf_forms.py.)
    """
    app.config["TESTING"] = True
    app.config["WTF_CSRF_ENABLED"] = False
    try:
        with app.test_client() as c:
            yield c
    finally:
        app.config["WTF_CSRF_ENABLED"] = True
        app.config["TESTING"] = False


@pytest.fixture
def fixture_docx_bytes() -> bytes:
    """A small manual-shaped DOCX: two chapters, sub-headings, internal references."""
    from docx import Document

    doc = Document()
    doc.add_heading("Chapter One - Administration of Graduate Programs", level=1)
    doc.add_paragraph("Introductory text. See Chapter 1.A for governance details.")
    doc.add_heading("Governance", level=2)
    doc.add_paragraph("Body text about governance and oversight.")
    doc.add_heading("Faculty Duties", level=2)
    doc.add_paragraph("As described in Chapter 1.A, duties are defined here.")
    doc.add_heading("Chapter Two - Degree Requirements", level=1)
    doc.add_paragraph("Degree requirements are listed below. See also Chapter 1.B.")
    doc.add_heading("Credit Hours", level=2)
    doc.add_paragraph("A minimum number of graded credit hours is required.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def fixture_docx_with_table_bytes() -> bytes:
    """Like fixture_docx_bytes but with a table, to exercise the table-review step."""
    from docx import Document

    doc = Document()
    doc.add_heading("Chapter One - Administration", level=1)
    doc.add_paragraph("Intro text referencing Chapter 1.A.")
    doc.add_heading("Deadlines", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Deadline"
    table.rows[1].cells[0].text = "Fall"
    table.rows[1].cells[1].text = "August 1"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
