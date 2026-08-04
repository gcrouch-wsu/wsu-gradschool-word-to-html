"""External URL entry in the reference review editor.

The output gate (``is_safe_href``) requires an explicit scheme, which is
correct — but the review form discarded anything it refused without saying so,
so an operator who pasted "policies.wsu.edu/bppm-10-65" watched the value
disappear from the field and only found the missing link later in the output.

Input is now normalized (scheme-less hosts are promoted to https) and whatever
is still refused is reported back.
"""

import io
import json
import re

import pytest

from config import SessionDir
from utils.url_policy import is_safe_href, normalize_external_href

from tests.conftest import pandoc_required


@pytest.mark.parametrize(
    "raw,expected",
    [
        ("policies.wsu.edu/prf/bppm-10-65", "https://policies.wsu.edu/prf/bppm-10-65"),
        ("www.wsu.edu", "https://www.wsu.edu"),
        ("wsu.edu:8080/a?b=1#c", "https://wsu.edu:8080/a?b=1#c"),
        ("//policies.wsu.edu/x", "https://policies.wsu.edu/x"),
        ("  wsu.edu/x  ", "https://wsu.edu/x"),
    ],
)
def test_scheme_less_hosts_are_promoted_to_https(raw, expected):
    assert normalize_external_href(raw) == expected


@pytest.mark.parametrize(
    "raw",
    ["https://wsu.edu/x", "http://wsu.edu/x", "mailto:a@wsu.edu", "#anchor-id"],
)
def test_already_valid_values_pass_through_unchanged(raw):
    assert normalize_external_href(raw) == raw


@pytest.mark.parametrize(
    "raw",
    [
        "javascript:alert(1)",
        "data:text/html,<script>alert(1)</script>",
        "vbscript:msgbox",
        "file://c/x.txt",
    ],
)
def test_dangerous_schemes_are_never_rewritten(raw):
    """A value carrying its own scheme must be refused, not turned into a URL."""
    assert normalize_external_href(raw) == ""


@pytest.mark.parametrize(
    "raw",
    ["42.52.040", "Section 6.0", "Chapter 42.52", "not a url", "", "   ", "wsu"],
)
def test_reference_like_text_is_not_mistaken_for_a_host(raw):
    """Dotted reference labels must not become links."""
    assert normalize_external_href(raw) == ""


def test_normalized_output_always_passes_the_strict_output_gate():
    for raw in ("policies.wsu.edu/x", "//wsu.edu", "WSU.EDU/Path", "wsu.edu:8080"):
        normalized = normalize_external_href(raw)
        assert normalized and is_safe_href(normalized)


# --- editor behavior -----------------------------------------------------

DOCX_PARAGRAPHS = [
    "Governance is described in Chapter 1.A.1 of this manual.",
    "See Chapter 1.A.2 and Chapter 1.A.3 for details.",
]


@pytest.fixture
def review_session(client_nocsrf):
    """A converted session parked on the reference review page."""
    from docx import Document

    doc = Document()
    doc.add_heading("Chapter One - Administration", level=1)
    doc.add_paragraph(DOCX_PARAGRAPHS[0])
    doc.add_heading("Governance", level=2)
    doc.add_paragraph(DOCX_PARAGRAPHS[1])
    buf = io.BytesIO()
    doc.save(buf)

    r = client_nocsrf.post(
        "/convert",
        data={"docx": (io.BytesIO(buf.getvalue()), "m.docx"), "mapping_mode": "map_new"},
        content_type="multipart/form-data",
    )
    sid = r.headers["Location"].rstrip("/").split("/")[-1]
    client_nocsrf.post(f"/heading_review/{sid}", data={})
    page = client_nocsrf.get(f"/review/{sid}").get_data(as_text=True)
    ref_ids = sorted(set(re.findall(r'name="ref_external_([^"]+)"', page)))
    assert ref_ids, "the fixture document should produce references"
    return sid, ref_ids


def _submit(client, sid, ref_ids, extra):
    """Submit the review form the way a browser does: every field on the page."""
    page = client.get(f"/review/{sid}").get_data(as_text=True)
    form = {}
    for rid in ref_ids:
        form[f"ref_valid_{rid}"] = "on"
        m = re.search(
            r'name="ref_external_' + re.escape(rid) + r'"[^>]*?value="([^"]*)"', page, re.S
        )
        form[f"ref_external_{rid}"] = m.group(1) if m else ""
    form.update(extra)
    return client.post(f"/review/{sid}", data=form, follow_redirects=True)


def _saved_urls(sid):
    return json.loads(SessionDir(sid).edits_json.read_text()).get(
        "reference_external_urls", {}
    )


@pandoc_required
def test_editor_promotes_a_scheme_less_url(client_nocsrf, review_session):
    sid, ref_ids = review_session
    _submit(client_nocsrf, sid, ref_ids,
            {f"ref_external_{ref_ids[0]}": "policies.wsu.edu/prf/bppm-10-65", "save_edits": "1"})
    assert list(_saved_urls(sid).values()) == ["https://policies.wsu.edu/prf/bppm-10-65"]


@pandoc_required
def test_editor_reports_a_value_it_cannot_save(client_nocsrf, review_session):
    sid, ref_ids = review_session
    body = _submit(client_nocsrf, sid, ref_ids,
                   {f"ref_external_{ref_ids[0]}": "Section 6.0", "save_edits": "1"}).get_data(as_text=True)
    assert "were NOT saved" in body, "the operator must be told the value was dropped"
    assert "Section 6.0" in body, "the message must name the offending value"
    assert _saved_urls(sid) == {}


@pandoc_required
def test_rejected_value_does_not_disturb_a_good_one(client_nocsrf, review_session):
    sid, ref_ids = review_session
    if len(ref_ids) < 2:
        pytest.skip("needs at least two references")
    _submit(client_nocsrf, sid, ref_ids, {
        f"ref_external_{ref_ids[0]}": "policies.wsu.edu/x",
        f"ref_external_{ref_ids[1]}": "Section 6.0",
        "save_edits": "1",
    })
    assert list(_saved_urls(sid).values()) == ["https://policies.wsu.edu/x"]


@pandoc_required
def test_promoted_url_survives_resubmit_and_reaches_the_output(client_nocsrf, review_session):
    """The normalized value re-renders in the field and links in the conversion."""
    from bs4 import BeautifulSoup

    sid, ref_ids = review_session
    _submit(client_nocsrf, sid, ref_ids,
            {f"ref_external_{ref_ids[0]}": "policies.wsu.edu/prf/bppm-10-65", "save_edits": "1"})

    page = client_nocsrf.get(f"/review/{sid}").get_data(as_text=True)
    assert "https://policies.wsu.edu/prf/bppm-10-65" in page, "field must show the normalized URL"

    # A second save must not lose it (the field posts back the normalized value).
    _submit(client_nocsrf, sid, ref_ids, {"proceed": "1"})
    assert list(_saved_urls(sid).values()) == ["https://policies.wsu.edu/prf/bppm-10-65"]

    main = BeautifulSoup(
        client_nocsrf.get(f"/convert/{sid}").get_data(as_text=True), "html.parser"
    ).find("main")
    anchor = main.find("a", href="https://policies.wsu.edu/prf/bppm-10-65")
    assert anchor is not None
    assert anchor.get("target") == "_blank"
    assert "external-link" in (anchor.get("class") or [])


@pandoc_required
def test_dangerous_value_is_reported_and_never_stored(client_nocsrf, review_session):
    sid, ref_ids = review_session
    body = _submit(client_nocsrf, sid, ref_ids,
                   {f"ref_external_{ref_ids[0]}": "javascript:alert(1)", "save_edits": "1"}).get_data(as_text=True)
    assert "were NOT saved" in body
    assert _saved_urls(sid) == {}
